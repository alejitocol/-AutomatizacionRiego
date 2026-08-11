# --- INTERFAZ Y GRÁFICOS ---
import streamlit as st
import plotly.express as px
import plotly.io as pio
import folium
from streamlit_folium import st_folium

# --- MANEJO DE DATOS Y CÁLCULOS ---
import pandas as pd
import numpy as np
import math
import json
import datetime

# --- GEOPROCESAMIENTO (Sin GDAL pesado) ---
import rasterio
from rasterio.mask import mask
from rasterio.io import MemoryFile
import geopandas as gpd
from shapely.geometry import Point, mapping
from shapely.ops import transform

# --- UTILIDADES DEL SISTEMA ---
import os
import shutil
import tempfile
import zipfile
import io
import requests

# --- MATPLOTLIB (exportación de gráficas a bytes) ---
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

# --- AÑO DINÁMICO PARA FUENTES ---
import datetime as _dt
ANO_ACTUAL = _dt.datetime.now().year

# --- GENERACIÓN DE INFORMES (ADR) ---
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- GENERADORES DE ANEXOS (Anexo 3a, 6, 7 Word + 7a Excel) ---
# Módulo externo con las cuatro funciones generadoras. Debe estar en la
# misma carpeta que app.py (generadores_anexos.py).
import generadores_anexos as gax

# --- PRUEBAS (ADR) ---
import streamlit as st
import pandas as pd
import io

# --- CACHÉ / RED ---
REQUEST_TIMEOUT = 30

@st.cache_data(show_spinner=False)
def fetch_nasa_data(lat, lon, fecha_inicio_str, fecha_fin_str):
    url_nasa = (
        f"https://power.larc.nasa.gov/api/temporal/daily/point"
        f"?parameters=PRECTOTCORR,T2M_MAX,T2M_MIN,EVPTRNS"
        f"&community=AG&longitude={lon}&latitude={lat}"
        f"&start={fecha_inicio_str}&end={fecha_fin_str}&format=JSON"
    )
    response = requests.get(url_nasa, timeout=REQUEST_TIMEOUT)
    response.raise_for_status()
    data = response.json()
    params = data['properties']['parameter']
    df_nasa = pd.DataFrame({
        'Fecha': list(params['PRECTOTCORR'].keys()),
        'Precipitacion': list(params['PRECTOTCORR'].values()),
        'T_Max': list(params['T2M_MAX'].values()),
        'T_Min': list(params['T2M_MIN'].values()),
        'Evaporacion': list(params['EVPTRNS'].values())
    })
    df_nasa.replace(-999.0, np.nan, inplace=True)
    df_nasa.fillna(0, inplace=True)
    df_nasa['Fecha'] = pd.to_datetime(df_nasa['Fecha'], format='%Y%m%d')
    return df_nasa

def calcular_ret_vectorizado(df_nasa, lat_input, krs=0.0023):
    """
    Evapotranspiración de referencia por Hargreaves-Samani (1985).

    ET0 = kRS · Ra_mm · (Tmean + 17.8) · sqrt(Tmax − Tmin)

    kRS = coeficiente de Hargreaves-Samani:
      0.0023  → valor original (zonas áridas / semiáridas)
      0.00185 → corrección FAO-56 para zonas húmedas (~0.85 × 0.0023)
      Rango recomendado FAO-56: 0.0019 – 0.0025

    El usuario puede ajustar kRS en la interfaz según la clasificación
    climática de la zona (índice de aridez = P/ET0).
    """
    df = df_nasa.copy()
    df['DOY'] = df['Fecha'].dt.dayofyear
    lat_rad = math.radians(lat_input)

    t_max = df['T_Max'].to_numpy(dtype=float)
    t_min = df['T_Min'].to_numpy(dtype=float)
    doy = df['DOY'].to_numpy(dtype=float)
    t_mean = (t_max + t_min) / 2.0

    dr = 1.0 + 0.033 * np.cos(2.0 * np.pi * doy / 365.0)
    delta = 0.409 * np.sin((2.0 * np.pi * doy / 365.0) - 1.39)
    ws_arg = -np.tan(lat_rad) * np.tan(delta)
    ws_arg = np.clip(ws_arg, -1.0, 1.0)
    ws = np.arccos(ws_arg)

    ra = (24.0 * 60.0 / np.pi) * 0.0820 * dr * (
        ws * np.sin(lat_rad) * np.sin(delta) +
        np.cos(lat_rad) * np.cos(delta) * np.sin(ws)
    )
    ra_mm = ra * 0.408
    delta_t = np.maximum(t_max - t_min, 0.0)

    ret = np.where(
        t_max > t_min,
        krs * ra_mm * (t_mean + 17.8) * np.sqrt(delta_t),
        0.0
    )
    df['RET'] = ret
    return df

def preparar_base_nasa(lat_input, lon_input, fecha_inicio, fecha_fin, krs=0.0023):
    f_inicio_nasa = fecha_inicio.strftime("%Y%m%d")
    f_fin_nasa = fecha_fin.strftime("%Y%m%d")
    df_nasa = fetch_nasa_data(lat_input, lon_input, f_inicio_nasa, f_fin_nasa)
    if df_nasa is None or df_nasa.empty:
        st.warning(
            f"⚠️ NASA POWER no retornó información para el periodo solicitado "
            f"({fecha_inicio} – {fecha_fin}) en este punto. El servicio puede no tener "
            f"cobertura completa para años más antiguos en esta ubicación; "
            f"se recomienda ajustar la fecha de inicio (por ejemplo, a partir de 2010 o 2018)."
        )
        return df_nasa
    primer_dato = df_nasa['Fecha'].min()
    if primer_dato.date() > fecha_inicio:
        st.info(
            f"ℹ️ Se solicitaron datos desde {fecha_inicio}, pero el primer registro disponible "
            f"en NASA POWER para este punto es {primer_dato.date()}. Los años previos no cuentan "
            f"con datos para esta ubicación."
        )
    df_nasa = calcular_ret_vectorizado(df_nasa, lat_input, krs=krs)
    return df_nasa

def agregar_decadas(df_base_diario):
    df = df_base_diario.copy()
    df['Año'] = df['Fecha'].dt.year
    df['Mes'] = df['Fecha'].dt.month
    df['Día'] = df['Fecha'].dt.day
    df['Decada_Mes'] = np.select(
        [df['Día'] <= 10, df['Día'] <= 20],
        [1, 2],
        default=3
    )
    df['Decada_Año'] = (df['Mes'] - 1) * 3 + df['Decada_Mes']
    return df

def _agregar_parrafo_justificado(doc, texto, negrita=False, italica=False, tamanio=None):
    """Agrega un párrafo justificado con formato opcional."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texto)
    run.bold = negrita
    run.italic = italica
    if tamanio:
        run.font.size = Pt(tamanio)
    return p

def _agregar_ecuacion(doc, texto_ecuacion, descripcion=None):
    """Agrega un bloque de ecuación con borde y descripción opcional."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto_ecuacion)
    run.bold = True
    run.font.size = Pt(11)
    if descripcion:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run2 = p2.add_run(descripcion)
        run2.italic = True
        run2.font.size = Pt(10)

def crear_memoria_hidrologia(
    datos_clima, coordenadas, df_simulacion=None,
    tipo_almacenamiento="No definido", vol_max=0,
    # --- Datos del Proyecto (inputs del usuario) ---
    nombre_proyecto="",
    departamento="",
    municipio="",
    lat_coord=0.0, lon_coord=0.0,
    nombre_beneficiario="",
    id_predio="",
    nombre_cultivo="",
    sistema_riego="",
    # --- Parámetros del reservorio ---
    radio_tanque=0.0,
    altura_tanque=0.0,
    diametro_tanque=0.0,
    # Cosecha lluvias
    habilitar_cosecha=False,
    largo_tejado=0.0,
    ancho_tejado=0.0,
    area_tejado_fisica=0.0,
    # Excavado
    es_excavado=False,
    df_batimetria=None,
    h_vals=None, a_vals=None, v_vals=None,
    # Gráficas
    imagen_clima_bytes=None,
    imagen_simulacion_bytes=None,
    imagen_area_volumen_bytes=None,
    imagen_esquema_bytes=None,
    # Área cultivo
    area_cultivo_ha=0.0,
    # Resumen predio
    fuente_datos="N/D",
    num_anios_serie=0,
    num_sectores=1,
):
    from docx.shared import RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # =========================================================
    # ESTILO GLOBAL
    # =========================================================
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # =========================================================
    # PORTADA / TÍTULO
    # =========================================================
    titulo = doc.add_heading('Anexo 3. Hidrología Básica e Información de Reservorios', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph()

    # =========================================================
    # SECCIÓN 1: DATOS DEL PROYECTO
    # =========================================================
    h1 = doc.add_heading('1. DATOS DEL PROYECTO', level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Texto introductorio
    _agregar_parrafo_justificado(doc,
        "A continuación, se relaciona la información del proyecto para el cual se desarrolla el presente análisis de hidrología básica e información del reservorio de almacenamiento hídrico.")

    doc.add_paragraph()

    # Tabla de datos del proyecto
    tabla_datos = doc.add_table(rows=9, cols=2)
    tabla_datos.style = 'Table Grid'
    tabla_datos.columns[0].width = Inches(2.5)
    tabla_datos.columns[1].width = Inches(4.5)

    datos_proyecto = [
        ("Nombre del Proyecto", nombre_proyecto if nombre_proyecto else "N/D"),
        ("Departamento", departamento if departamento else "N/D"),
        ("Municipio", municipio if municipio else "N/D"),
        ("Coordenadas (Lat, Lon)", f"{lat_coord:.6f}°N, {lon_coord:.6f}°W"),
        ("Nombre del Potencial Beneficiario", nombre_beneficiario if nombre_beneficiario else "N/D"),
        ("ID del Predio", id_predio if id_predio else "N/D"),
        ("Cultivo Analizado", nombre_cultivo if nombre_cultivo else "N/D"),
        ("Sistema de Riego", sistema_riego if sistema_riego else "N/D"),
        ("Tipo de Almacenamiento", tipo_almacenamiento),
    ]

    for i, (campo, valor) in enumerate(datos_proyecto):
        celda_campo = tabla_datos.rows[i].cells[0]
        celda_valor = tabla_datos.rows[i].cells[1]
        celda_campo.text = campo
        celda_valor.text = str(valor)
        # Negrita en el campo
        for run in celda_campo.paragraphs[0].runs:
            run.bold = True
        # Color de fondo en celda de campo
        tc_pr = celda_campo._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E2F3')
        tc_pr.append(shd)

    doc.add_paragraph()

    # =========================================================
    # SECCIÓN 2: RESUMEN DEL PREDIO
    # =========================================================
    doc.add_heading('2. RESUMEN DEL PREDIO', level=1)

    _agregar_parrafo_justificado(doc,
        "A continuación se presenta un resumen ejecutivo de los parámetros de diseño adoptados para el predio, "
        "consolidando la fuente de información climática utilizada, la serie temporal analizada y las "
        "características principales del sistema de riego y almacenamiento hídrico.")

    doc.add_paragraph()

    tabla_resumen = doc.add_table(rows=7, cols=2)
    tabla_resumen.style = 'Table Grid'
    tabla_resumen.columns[0].width = Inches(3.0)
    tabla_resumen.columns[1].width = Inches(4.0)

    area_m2_resumen = area_cultivo_ha * 10000
    area_str_resumen = f"{area_m2_resumen:,.0f} m²".replace(",", ".")

    datos_resumen = [
        ("Fuente de datos climáticos",   fuente_datos if fuente_datos else "N/D"),
        ("Años en la serie de datos",     str(num_anios_serie) if num_anios_serie else "N/D"),
        ("Área determinada de riego",     area_str_resumen),
        ("Tipo de riego",                 sistema_riego if sistema_riego else "N/D"),
        ("Cultivo seleccionado",          nombre_cultivo if nombre_cultivo else "N/D"),
        ("Número de sectores de riego",   str(num_sectores)),
        ("Volumen del reservorio",        f"{vol_max:.2f} m³ ({tipo_almacenamiento})"),
    ]

    for i, (campo, valor) in enumerate(datos_resumen):
        celda_campo = tabla_resumen.rows[i].cells[0]
        celda_valor = tabla_resumen.rows[i].cells[1]
        celda_campo.text = campo
        celda_valor.text = valor
        for run in celda_campo.paragraphs[0].runs:
            run.bold = True
        tc_pr_r = celda_campo._tc.get_or_add_tcPr()
        shd_r = OxmlElement('w:shd')
        shd_r.set(qn('w:val'),   'clear')
        shd_r.set(qn('w:color'), 'auto')
        shd_r.set(qn('w:fill'),  'D9E2F3')
        tc_pr_r.append(shd_r)

    p_fuente_resumen = doc.add_paragraph(f"Fuente: Elaboración propia ADR {ANO_ACTUAL}.")
    p_fuente_resumen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_fuente_resumen.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph()

    # =========================================================
    # SECCIÓN 3: CONTENIDO TÉCNICO - PRECIPITACIÓN DIARIA Y DECADAL
    # =========================================================
    doc.add_heading('3. METODOLOGÍA APLICADA – CLIMATOLOGÍA', level=1)

    _agregar_parrafo_justificado(doc,
        "A continuación, se relaciona la información para los predios con los siguientes datos:")

    items_lista = [
        "Precipitación Diaria y Decadal.",
        "Representación del Cuerpo de Agua.",
        "Representación Área y Volumen en Función de la Elevación.",
        "Altura y Volumen de Sedimentos del Reservorio.",
    ]
    for item in items_lista:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.add_run(item)

    doc.add_paragraph()

    # ------- 2.1 Precipitación Decadal -------
    doc.add_heading('3.1 Precipitación Decadal', level=2)
    _agregar_parrafo_justificado(doc,
        "Según lo analizado en el numeral de Climatología del Informe General, para la revisión actual se debe "
        "partir de la información obtenida en WaPOR referente a la precipitación diaria, para posteriormente "
        "agrupar las precipitaciones de forma total decadal (8, 10 u 11 días) según corresponda a cada década "
        "de estudio. Se realiza un ordenamiento de la precipitación decadal, de mayor a menor, para todos los "
        "años de información disponible garantizando un periodo mínimo de 8 años de datos analizados.")

    doc.add_paragraph()
    _agregar_parrafo_justificado(doc,
        "Se aplica el método Critchley y Siegert (1996) y Veenhuizen (2000) (ver Ecuación 1) con el cual es "
        "posible definir la probabilidad de que ocurran las precipitaciones decadales en el rango de tiempo "
        "de los años evaluados. De esta forma, se estima la probabilidad de excedencia del 75% para cada década.")

    _agregar_ecuacion(doc,
        "P (%) = (m – 0.375) / (N + 0.25) × 100    [Ecuación 1]",
        "Donde:\n  m: Número de orden\n  N: Número total de observaciones")

    doc.add_paragraph()

    # Gráfica climática si existe
    if imagen_clima_bytes:
        _agregar_parrafo_justificado(doc, "Ilustración 1. Precipitación Diaria y Decadal – Comportamiento Hídrico Decadal (P75%, Evaporación y RET).", italica=True)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(io.BytesIO(imagen_clima_bytes), width=Inches(5.5))
        p_caption = doc.add_paragraph(f"Fuente: Elaboración propia ADR {ANO_ACTUAL}.")
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_caption.runs:
            run.italic = True
            run.font.size = Pt(9)
        doc.add_paragraph()

    # ------- 2.2 Precipitación Efectiva -------
    doc.add_heading('3.2 Precipitación Efectiva', level=2)
    _agregar_parrafo_justificado(doc,
        "Una vez definida la precipitación decadal, se estima la precipitación efectiva, que se entiende como "
        "la precipitación disponible para el crecimiento de las plantas y para ser utilizada por los cultivos, "
        "es decir, la que queda almacenada en el suelo y dentro de la zona radicular. Para ello, se hace uso "
        "del método del Servicio de Conservación de Suelos (SCS) del Departamento de Agricultura de los Estados "
        "Unidos (USDA) que consiste en una función escalonada ajustada para la precipitación total decadal "
        "(ver Ecuación 2).")

    _agregar_ecuacion(doc,
        "Pe = P × (125 – 0.2P) / 125        para P ≤ 250/3 mm\n"
        "Pe = (125/3) + 0.1P                   para P > 250/3 mm    [Ecuación 2]",
        "Donde:\n  Pe: Precipitación efectiva (mm)\n  P: Precipitación total decadal (mm)")

    doc.add_paragraph()
    _agregar_parrafo_justificado(doc,
        "Tanto la precipitación decadal como la precipitación efectiva son componentes relevantes en la "
        "modelación del balance hídrico del cultivo, por lo que esta información se verá ampliada en el "
        "Anexo 6. Disponibilidad y demandas de agua.")

    doc.add_paragraph()

    # =========================================================
    # SECCIÓN 4: METODOLOGÍA APLICADA AL RESERVORIO
    # =========================================================
    doc.add_heading('4. METODOLOGÍA APLICADA AL RESERVORIO', level=1)

    _agregar_parrafo_justificado(doc,
        "Para garantizar la continuidad del cultivo, se requirió disponer de un cuerpo de almacenamiento de "
        "agua que capture y disponga el recurso hídrico en los periodos donde la precipitación no satisface "
        "la demanda agrícola. Para ello, se hace uso de estructuras hidráulicas que aprovechen la precipitación "
        "y su posterior conducción al llenado de un reservorio.")

    doc.add_paragraph()

    # -------- Descripción según tipo de reservorio --------
    if es_excavado:
        # Reservorio excavado (Vaso Irregular)
        doc.add_heading('4.1 Reservorio Excavado (Vaso Irregular)', level=2)
        _agregar_parrafo_justificado(doc,
            "El reservorio utilizado en el presente proyecto corresponde a un vaso de almacenamiento de "
            "geometría irregular, dimensionado a partir de los datos batimétricos levantados en campo. "
            f"La capacidad máxima de almacenamiento proyectada es de {vol_max:.2f} m³, determinada a "
            "partir de las curvas Cota-Área-Volumen obtenidas en la batimetría del sitio.")

        doc.add_paragraph()
        _agregar_parrafo_justificado(doc,
            "A diferencia de un reservorio cilíndrico, en el vaso irregular tanto el área del espejo de "
            "agua como el volumen almacenado varían en función de la altura de la lámina de agua, siguiendo "
            "la geometría natural del terreno. Para el cálculo riguroso de las pérdidas por evaporación y "
            "los aportes por precipitación directa, se emplearon funciones de interpolación derivadas de la "
            "batimetría, aplicando regresión polinómica sobre las curvas Cota-Área-Volumen.")

        # Tabla de batimetría si existe
        if df_batimetria is not None and not df_batimetria.empty:
            doc.add_paragraph()
            p_tab_bat = doc.add_paragraph("Tabla B1. Datos de Batimetría del Reservorio Excavado")
            p_tab_bat.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p_tab_bat.runs:
                run.bold = True

            columnas_bat = list(df_batimetria.columns)
            n_cols_bat = len(columnas_bat)
            tabla_bat = doc.add_table(rows=1, cols=n_cols_bat)
            tabla_bat.style = 'Table Grid'

            hdr_bat = tabla_bat.rows[0].cells
            for i, col in enumerate(columnas_bat):
                hdr_bat[i].text = col
                for run in hdr_bat[i].paragraphs[0].runs:
                    run.bold = True
                tc_pr_b = hdr_bat[i]._tc.get_or_add_tcPr()
                shd_b = OxmlElement('w:shd')
                shd_b.set(qn('w:val'), 'clear')
                shd_b.set(qn('w:color'), 'auto')
                shd_b.set(qn('w:fill'), 'BDD7EEf')
                tc_pr_b.append(shd_b)

            for _, fila_bat in df_batimetria.iterrows():
                celdas = tabla_bat.add_row().cells
                for i, col in enumerate(columnas_bat):
                    val = fila_bat[col]
                    celdas[i].text = f"{val:.3f}" if isinstance(val, float) else str(val)

            p_fuente_bat = doc.add_paragraph(f"Fuente: Elaboración propia ADR {ANO_ACTUAL}.")
            p_fuente_bat.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p_fuente_bat.runs:
                run.italic = True
                run.font.size = Pt(9)

    else:
        # Tanque Australiano
        doc.add_heading('4.1 Tanque Australiano (Armo-Tanque Circular)', level=2)
        diametro_usado = diametro_tanque if diametro_tanque > 0 else (radio_tanque * 2)
        _agregar_parrafo_justificado(doc,
            f"El reservorio utilizado en el presente proyecto corresponde a un armo-tanque (reservorio tipo "
            f"australiano) circular de {vol_max:.2f} m³ de capacidad nominal, diámetro de {diametro_usado:.2f} m "
            f"y altura de la pared plana del tanque de {altura_tanque:.2f} m.")

        doc.add_paragraph()
        _agregar_parrafo_justificado(doc,
            "Al tratarse de un tanque con pared vertical (perfil cilíndrico), el área del espejo de agua se "
            "mantiene constante para cualquier nivel de llenado, dado que la sección transversal no varía con "
            "la altura. En consecuencia, el volumen almacenado es directamente proporcional a la altura de la "
            "lámina de agua en el interior del tanque.")

        # Cosecha de aguas lluvias
        if habilitar_cosecha and area_tejado_fisica > 0:
            doc.add_paragraph()
            _agregar_parrafo_justificado(doc,
                f"Se requiere disponer de una cubierta (tejado) de {largo_tejado:.1f} m × {ancho_tejado:.1f} m "
                f"(área física = {area_tejado_fisica:.2f} m²) contigua al cuerpo de almacenamiento, que ayude "
                "a captar la precipitación y su posterior derivación para el llenado del tanque mediante "
                "cosecha de aguas lluvias.")

        # Sedimentos
        doc.add_paragraph()
        _agregar_parrafo_justificado(doc,
            "Adicionalmente, la capacidad del tanque se ve reducida por la acumulación de sedimentos, los "
            "cuales se consideran como parte del volumen total de llenado. En este sentido, se estima un 3% "
            "de sedimentos respecto a la altura del tanque. Es importante tener en cuenta estas consideraciones "
            "durante el dimensionamiento del reservorio, ya que permiten garantizar su funcionamiento óptimo "
            "y sostenible a largo plazo, evitando problemas asociados a la disminución de la capacidad de "
            "almacenamiento, tal como se indica en la Tabla 1.")

        doc.add_paragraph()

        # Tabla 1: Dimensiones y capacidad del tanque
        p_tab1 = doc.add_paragraph("Tabla 1. Dimensiones y capacidad del tanque")
        p_tab1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_tab1.runs:
            run.bold = True

        col_widths_t1 = [Inches(2.2), Inches(1.5), Inches(1.5), Inches(1.8), Inches(1.2)]
        tabla1 = doc.add_table(rows=3, cols=5)
        tabla1.style = 'Table Grid'

        encabezados_t1 = [
            "Condiciones del tanque",
            "Altura Total (m)",
            "Diámetro (m)",
            "Capacidad (m³)",
            "% Sedimentos"
        ]
        hdr_t1 = tabla1.rows[0].cells
        for i, enc in enumerate(encabezados_t1):
            hdr_t1[i].text = enc
            for run in hdr_t1[i].paragraphs[0].runs:
                run.bold = True
            tc_pr_t1 = hdr_t1[i]._tc.get_or_add_tcPr()
            shd_t1 = OxmlElement('w:shd')
            shd_t1.set(qn('w:val'), 'clear')
            shd_t1.set(qn('w:color'), 'auto')
            shd_t1.set(qn('w:fill'), 'BDD7EEf')
            tc_pr_t1.append(shd_t1)

        # Fila de dimensiones del tanque adoptado
        diametro_mostrar = diametro_tanque if diametro_tanque > 0 else (radio_tanque * 2)
        vol_neto = math.pi * (radio_tanque ** 2) * altura_tanque if radio_tanque > 0 else vol_max
        altura_sed = round(altura_tanque * 0.03, 3)

        datos_t1_fila1 = [
            "Tanque australiano",
            f"{altura_tanque:.2f}",
            f"{diametro_mostrar:.2f}",
            f"{vol_neto:.2f}",
            "3"
        ]
        datos_t1_fila2 = [
            "Altura de sedimentos (m)",
            f"{altura_sed:.3f}",
            "—", "—", "—"
        ]
        for i, val in enumerate(datos_t1_fila1):
            tabla1.rows[1].cells[i].text = val
        for i, val in enumerate(datos_t1_fila2):
            tabla1.rows[2].cells[i].text = val

        p_fuente_t1 = doc.add_paragraph(f"Fuente: Elaboración propia ADR {ANO_ACTUAL}.")
        p_fuente_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_fuente_t1.runs:
            run.italic = True
            run.font.size = Pt(9)

    # =========================================================
    # SECCIÓN 5: ÁREA Y VOLUMEN EN FUNCIÓN DE LA ELEVACIÓN
    # =========================================================
    doc.add_paragraph()
    doc.add_heading('5. ÁREA Y VOLUMEN EN FUNCIÓN DE LA ELEVACIÓN', level=1)

    _agregar_parrafo_justificado(doc,
        "El modelamiento del reservorio contempla la simulación del área y el volumen en función de la altura "
        "de la lámina de agua. ")

    if es_excavado:
        _agregar_parrafo_justificado(doc,
            "En el caso del reservorio excavado (vaso irregular), tanto el área como el volumen varían con la "
            "elevación de acuerdo con la geometría irregular del terreno. Para interpretar adecuadamente ambas "
            "simulaciones, se aplica una regresión polinómica sobre las curvas batimétrica, tal como se muestra "
            "en la Ilustración 1.")
    else:
        _agregar_parrafo_justificado(doc,
            "En el caso de los armo-tanques (reservorios tipo australiano), el área se mantiene constante para "
            "cualquier elevación, dado que se trata de una pared vertical; mientras que el volumen está "
            "directamente determinado por la altura de la lámina de agua en el tanque. Para interpretar "
            "adecuadamente ambas simulaciones, es pertinente aplicar una regresión lineal sobre este "
            "comportamiento, tal como se muestra en la Ilustración 1.")

    doc.add_paragraph()

    if imagen_area_volumen_bytes:
        p_il1 = doc.add_paragraph("Ilustración 1. Ecuaciones de representación Área y Volumen en función de la elevación del reservorio.")
        p_il1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_il1.runs:
            run.bold = True
        p_img1 = doc.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img1 = p_img1.add_run()
        run_img1.add_picture(io.BytesIO(imagen_area_volumen_bytes), width=Inches(5.5))
        p_fuente_il1 = doc.add_paragraph(f"Fuente: Elaboración propia ADR {ANO_ACTUAL}.")
        p_fuente_il1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_fuente_il1.runs:
            run.italic = True
            run.font.size = Pt(9)
    else:
        _agregar_parrafo_justificado(doc,
            "[Ilustración 1. Ecuaciones de representación Área y Volumen en función de la elevación del reservorio. "
            f"Fuente: Elaboración propia ADR {ANO_ACTUAL}. — Gráfica no disponible, genere la simulación en Pestaña 3 primero.]",
            italica=True)

    doc.add_paragraph()

    # =========================================================
    # SECCIÓN 6: SIMULACIÓN DEL COMPORTAMIENTO DEL RESERVORIO
    # =========================================================
    doc.add_heading('6. SIMULACIÓN DEL COMPORTAMIENTO DEL RESERVORIO', level=1)

    _agregar_parrafo_justificado(doc,
        "Para analizar si un reservorio puede cumplir con las necesidades hídricas y satisfacer la demanda, "
        "se parte de un volumen inicial supuesto y se simula su variación considerando las entradas y salidas "
        "de agua en un periodo determinado, cuya explicación se detalla en el Anexo 6. Disponibilidad y "
        "demandas de agua.")

    doc.add_paragraph()
    _agregar_parrafo_justificado(doc,
        "Si durante este periodo el volumen del tanque no desciende a un nivel crítico que comprometa su "
        "capacidad de atender la demanda, se considera adecuado para su propósito. En caso contrario, se "
        "deben implementar medidas de gestión hídrica que garanticen un suministro sostenible. La Ilustración 2 "
        "muestra el comportamiento del reservorio en este proyecto a lo largo del tiempo, evidenciando tanto "
        "la recarga como el uso del agua en el transcurso de las décadas.")

    doc.add_paragraph()

    if imagen_simulacion_bytes:
        p_il2 = doc.add_paragraph("Ilustración 2. Comportamiento Volumen del Reservorio")
        p_il2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_il2.runs:
            run.bold = True
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img2 = p_img2.add_run()
        run_img2.add_picture(io.BytesIO(imagen_simulacion_bytes), width=Inches(5.5))
        p_fuente_il2 = doc.add_paragraph(f"Fuente: Elaboración propia ADR {ANO_ACTUAL}.")
        p_fuente_il2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_fuente_il2.runs:
            run.italic = True
            run.font.size = Pt(9)
    else:
        _agregar_parrafo_justificado(doc,
            f"[Ilustración 2. Comportamiento Volumen del Reservorio. Fuente: Elaboración propia ADR {ANO_ACTUAL}. "
            "— Gráfica no disponible, genere la simulación en Pestaña 3 primero.]",
            italica=True)

    doc.add_paragraph()

    # =========================================================
    # SECCIÓN 7: TABLA DE FUNCIONAMIENTO DEL VASO
    # =========================================================
    doc.add_heading('7. FUNCIONAMIENTO DEL VASO DE ALMACENAMIENTO (TABLA 2)', level=1)

    _agregar_parrafo_justificado(doc,
        "La simulación del volumen del reservorio debe realizarse para cada década, de manera que se obtengan "
        "resultados equivalentes al número de décadas con información disponible. Para este proceso, es "
        "necesario diligenciar un formato similar al presentado en la Tabla 2.")

    doc.add_paragraph()
    p_tab2 = doc.add_paragraph("Tabla 2. Funcionamiento del vaso de almacenamiento")
    p_tab2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_tab2.runs:
        run.bold = True

    columnas_tabla = ['AÑO', 'MES', 'DÉCADA', 'VOL. INICIAL (m³)', 'ENTRADAS (+) (m³)', 'SALIDAS (−) (m³)', 'VOL. FINAL (m³)', 'EXCEDENTE (m³)']

    if df_simulacion is not None and not df_simulacion.empty:
        tabla2 = doc.add_table(rows=1, cols=len(columnas_tabla))
        tabla2.style = 'Table Grid'

        hdr2 = tabla2.rows[0].cells
        for i, col_name in enumerate(columnas_tabla):
            hdr2[i].text = col_name
            for run in hdr2[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
            tc_pr2 = hdr2[i]._tc.get_or_add_tcPr()
            shd2 = OxmlElement('w:shd')
            shd2.set(qn('w:val'), 'clear')
            shd2.set(qn('w:color'), 'auto')
            shd2.set(qn('w:fill'), 'BDD7EE')
            tc_pr2.append(shd2)

        # Mapeo de meses
        meses_map = {1:"Ene",2:"Feb",3:"Mar",4:"Abr",5:"May",6:"Jun",
                     7:"Jul",8:"Ago",9:"Sep",10:"Oct",11:"Nov",12:"Dic"}

        for index, row in df_simulacion.iterrows():
            entradas = (row.get('Entrada Concesion (m3)', 0) +
                        row.get('Entrada Lluvia (m3)', 0) +
                        row.get('Entrada Escorrentia (m3)', 0))
            salidas = (row.get('Salida Riego (m3)', 0) +
                       row.get('Salida Evaporación (m3)', 0) +
                       row.get('Salida Infiltración (m3)', 0))

            decada_num = int(row.get('Decada', 0))
            mes_num = math.ceil(decada_num / 3) if decada_num > 0 else 1
            mes_num = min(mes_num, 12)
            mes_str = meses_map.get(mes_num, "N/D")

            row_cells = tabla2.add_row().cells
            row_cells[0].text = str(int(row.get('Año', 0)))
            row_cells[1].text = mes_str
            row_cells[2].text = str(decada_num)
            row_cells[3].text = f"{row.get('Volumen Inicial (m3)', 0):.2f}"
            row_cells[4].text = f"{entradas:.2f}"
            row_cells[5].text = f"{salidas:.2f}"
            row_cells[6].text = f"{row.get('Volumen Final (m3)', 0):.2f}"
            row_cells[7].text = f"{row.get('Volumen Derramado (m3)f', 0):.2f}"

            for cell in row_cells:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)

        p_nota2 = doc.add_paragraph()
        run_nota2 = p_nota2.add_run(
            "*Nota: La tabla presenta el balance completo para toda la serie climática analizada. "
            "El balance histórico completo reposa en los archivos digitales del proyecto.")
        run_nota2.italic = True
        run_nota2.font.size = Pt(9)
    else:
        doc.add_paragraph("⚠️ Error: No se encontraron datos de simulación. Ejecute la Pestaña 3 primero.")

    p_fuente_tab2 = doc.add_paragraph(f"Fuente: Elaboración propia ADR {ANO_ACTUAL}.")
    p_fuente_tab2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_fuente_tab2.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph()
    _agregar_parrafo_justificado(doc,
        "Con el diligenciamiento del formato anterior es posible identificar el comportamiento del reservorio "
        "a lo largo del tiempo, conociendo la cantidad de agua almacenada al inicio y al final de cada década. "
        "En este análisis deben evaluarse tanto las entradas que aportan agua al llenado del tanque como las "
        "salidas que representan su vaciado. Asimismo, es necesario estimar el volumen excedente (derramado) "
        "en las décadas en las que no sea posible continuar con el almacenamiento. El concepto de entradas y "
        "salidas se explica con mayor detalle en el Anexo 6. Disponibilidad y demandas de agua.")

    doc.add_paragraph()
    _agregar_parrafo_justificado(doc,
        "Finalmente, se concluye con el balance del reservorio y la evaluación de sostenibilidad en el tiempo "
        "entre la oferta y la demanda hídrica.")

    doc.add_paragraph()

    # =========================================================
    # SECCIÓN 8: ESQUEMA REPRESENTATIVO DE ÁREAS
    # =========================================================
    doc.add_heading('8. ESQUEMA REPRESENTATIVO DE ÁREAS', level=1)

    desc_esquema = (
        "La Ilustración 3 presenta el esquema a escala real de las obras principales del proyecto. "
        "Se muestra la distribución espacial de: (i) el área de cultivo determinada por los parámetros "
        "agronómicos ingresados; (ii) el cuerpo de agua (reservorio) con sus dimensiones adoptadas de diseño"
    )
    if habilitar_cosecha and area_tejado_fisica > 0:
        desc_esquema += (
            f"; y (iii) el área de cubierta o tejado ({area_tejado_fisica:.2f} m²) diseñado para la "
            "cosecha de aguas lluvias y posterior conducción al reservorio de almacenamiento"
        )
    desc_esquema += (
        ". La disposición gráfica permite dimensionar la magnitud de las obras civiles frente a la "
        "extensión agrícola del predio intervenido."
    )

    _agregar_parrafo_justificado(doc, desc_esquema)
    doc.add_paragraph()

    if imagen_esquema_bytes:
        p_il3 = doc.add_paragraph("Ilustración 3. Esquema representativo de áreas (Cuerpo de agua, cubierta y cultivo)")
        p_il3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_il3.runs:
            run.bold = True
        p_img3 = doc.add_paragraph()
        p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img3 = p_img3.add_run()
        run_img3.add_picture(io.BytesIO(imagen_esquema_bytes), width=Inches(5.5))
        p_fuente_il3 = doc.add_paragraph(f"Fuente: Elaboración propia ADR {ANO_ACTUAL}.")
        p_fuente_il3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_fuente_il3.runs:
            run.italic = True
            run.font.size = Pt(9)
    else:
        _agregar_parrafo_justificado(doc,
            f"[Ilustración 3. Esquema representativo de áreas. Fuente: Elaboración propia ADR {ANO_ACTUAL}. "
            "— Gráfica no disponible, genere la simulación en Pestaña 3 primero.]",
            italica=True)

    return doc

def crear_memoria_demandas_legacy(datos_cultivo):
    """[OBSOLETA] Reemplazada por gax.crear_memoria_demandas() en el módulo
    generadores_anexos. Se conserva por compatibilidad."""
    doc = Document()
    doc.add_heading('Memoria de Cálculo: Disponibilidad y Demandas hídricas', 0)
    
    doc.add_heading('1. Demanda Hídrica del Cultivo (ETc)', level=1)
    doc.add_paragraph(
        "El cálculo de la demanda se basa en la interacción entre la Evapotranspiración de Referencia (ET0) "
        "y el Coeficiente de Cultivo (Kc) específico para cada etapa fenológica."
    )
    
    doc.add_paragraph("ETc = ET0 * Kc", style='Intense Quote')
    
    doc.add_heading('2. Requerimiento de Riego', level=2)
    doc.add_paragraph(
        "Considerando la precipitación efectiva calculada en el Anexo 3, el requerimiento neto (Rn) se define como:"
    )
    doc.add_paragraph("Rn = ETc - P_efectiva", style='Intense Quote')
    
    return doc

# Inicialización de variables de estado (Session State)
# Coloca esto en la parte superior de tu app.py, fuera de cualquier pestaña
if 'zip_buffer' not in st.session_state:
    st.session_state.zip_buffer = None
if 'process_complete' not in st.session_state:
    st.session_state.process_complete = False

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Herramienta Computacional ADR", page_icon="💧", layout="wide")
st.title("💧 Herramienta Computacional: Análisis y Descarga de Datos Climáticos")

def extraer_fecha_de_nombre(nombre_archivo):
    """
    Extrae fecha desde nombres con formatos:
    - YYYYMMDD
    - YYYY-MM-DD
    - YYYY_MM_DD
    """
    import re
    from datetime import datetime

    match = re.search(r'(\d{4})[-_]?(\d{2})[-_]?(\d{2})', nombre_archivo)
    if not match:
        return None

    year, month, day = match.groups()
    try:
        return datetime(int(year), int(month), int(day))
    except ValueError:
        return None

def procesar_zip_wapor(archivo_zip, lon, lat, nombre_columna):
    """
    Lee un archivo ZIP con rasters de WaPOR (diarios o decadales) en memoria
    y extrae el valor para una coordenada puntual.

    Correcciones respecto a la versión anterior:
    - Verifica que la coordenada cae dentro del extent del TIF antes de muestrear.
    - Usa el nodata real del TIF (src.nodata) en lugar de un umbral hardcodeado.
    - Los píxeles con nodata o NaN se omiten (continue) en lugar de convertirse
      a 0.0, lo que evitaba que el P75% colapsara artificialmente a cero.
    - Aplica un límite físico de precipitación razonable (0–500 mm) para descartar
      artefactos numéricos del raster sin afectar lluvias extremas reales.
    - Los TIFs se procesan en orden cronológico (sorted) para consistencia.

    Retorna un DataFrame con columnas ['Fecha', nombre_columna].
    """
    resultados = []
    archivos_tif = 0
    archivos_sin_fecha = 0
    archivos_con_error = 0
    archivos_fuera_extent = 0

    with zipfile.ZipFile(archivo_zip) as z:
        nombres_ordenados = sorted(z.namelist())
        for filename in nombres_ordenados:
            if not filename.lower().endswith((".tif", ".tiff")):
                continue

            archivos_tif += 1

            fecha_extraida = extraer_fecha_de_nombre(os.path.basename(filename))
            if fecha_extraida is None:
                archivos_sin_fecha += 1
                continue

            fecha_obj = pd.to_datetime(fecha_extraida)

            try:
                with z.open(filename) as f:
                    with MemoryFile(f.read()) as memfile:
                        with memfile.open() as src:

                            # ── 1. Verificar que el punto cae dentro del extent del TIF ──
                            b = src.bounds
                            if not (b.left <= lon <= b.right and b.bottom <= lat <= b.top):
                                archivos_fuera_extent += 1
                                continue

                            # ── 2. Leer valor del píxel ──
                            muestra = list(src.sample([(lon, lat)]))
                            if not muestra or len(muestra[0]) == 0:
                                continue  # sin dato → omitir fila

                            raw = float(muestra[0][0])

                            # ── 3. Validar nodata usando el valor real del TIF ──
                            nodata_val = src.nodata
                            es_nodata = False
                            if nodata_val is not None:
                                try:
                                    if math.isnan(float(nodata_val)):
                                        es_nodata = math.isnan(raw)
                                    else:
                                        es_nodata = abs(raw - float(nodata_val)) < 0.01
                                except (TypeError, ValueError):
                                    pass

                            if es_nodata or math.isnan(raw):
                                continue  # omitir → NO convertir a 0.0

                            # ── 4. Límite físico según variable ──
                            # Precipitación: 0–500 mm/día o 0–500 mm/década (ambos rangos válidos)
                            # Evaporación/RET: 0–25 mm/día o 0–150 mm/década
                            # Usamos un tope conservador de 500 mm para todas las variables
                            if raw < 0.0 or raw > 500.0:
                                continue  # artefacto numérico → omitir

                            resultados.append({'Fecha': fecha_obj, nombre_columna: round(raw, 4)})

            except Exception:
                archivos_con_error += 1
                continue

    # Retorno consistente para evitar KeyError al hacer merge por 'Fecha'
    if resultados:
        df_resultado = pd.DataFrame(resultados)
        df_resultado = df_resultado.sort_values('Fecha').drop_duplicates(
            subset=['Fecha']).reset_index(drop=True)
    else:
        df_resultado = pd.DataFrame(columns=['Fecha', nombre_columna])

    return df_resultado, {
        'archivos_tif': archivos_tif,
        'archivos_validos': len(resultados),
        'archivos_sin_fecha': archivos_sin_fecha,
        'archivos_con_error': archivos_con_error,
        'archivos_fuera_extent': archivos_fuera_extent,
    }

# --- CREACIÓN DE PESTAÑAS ---
tab1, tab2, tab3, tab4 = st.tabs([
    "📊 Datos Agroclimáticos", 
    "💧 Balance Hídrico", 
    "📈 Volúmenes de Riego", 
    "⚙️ Generación WaPOR"
])
# =====================================================================
# --- PESTAÑA 1: ANÁLISIS DEL EXCEL / CSV --- (Sin cambios)
# =====================================================================
with tab1:
    st.markdown("Descarga la serie de la NASA o extrae información climática de repositorios Raster (WaPOR v3) para analizar la **Precipitación Confiable al 75%** y promedios decadales.")
    
    fuente_datos = st.radio("📡 Seleccione la fuente de datos:", ["NASA POWER (API Online)", "WaPOR v3 (Archivos Raster .ZIP)"], horizontal=True)
    
    col1, col2 = st.columns(2)
    with col1:
        lat_input = st.number_input("Latitud (Ej: 8.8488)", value=8.848795, format="%.6f", key="lat_t1")
    with col2:
        lon_input = st.number_input("Longitud (Ej: -73.6090)", value=-73.609039, format="%.6f", key="lon_t1")

    # Variable global de la pestaña que almacenará la serie diaria sin importar el origen
    df_base_diario = None 

    # ==========================================
    # RAMA 1: DESCARGA DESDE NASA POWER
    # ==========================================
    if fuente_datos == "NASA POWER (API Online)":
        col3, col4 = st.columns(2)
        with col3:
            fecha_inicio = st.date_input(
                "Fecha Inicio", pd.to_datetime("2018-01-01"),
                min_value=pd.to_datetime("2000-01-01"), key="fi_nasa"
            )
        with col4:
            fecha_fin = st.date_input("Fecha Fin", pd.to_datetime("2025-12-31"), key="ff_nasa")
        st.caption(
            "ℹ️ NASA POWER permite seleccionar fechas desde el año 2000. Si para el punto y periodo "
            "consultado el servicio no retorna información (años más antiguos con baja cobertura "
            "satelital/reanálisis), la aplicación lo indicará y se recomienda ajustar la fecha de inicio."
        )

        # ── Selector de coeficiente Hargreaves-Samani (kRS) ─────────────
        with st.expander("⚙️ Parámetro de Hargreaves-Samani (kRS) — ajuste por zona climática"):
            st.markdown(
                "La ETo se calcula con **Hargreaves-Samani (1985)**: `ETo = kRS · Ra · (Tmean+17.8) · √ΔT`  \n"
                "El coeficiente **kRS original es 0.0023** pero FAO-56 (Allen et al. 1998) recomienda "
                "reducirlo en zonas húmedas para evitar sobreestimación. "
                "La sobreestimación de ETo hace que ET > P aparezca en más décadas de las reales."
            )
            col_krs1, col_krs2 = st.columns([2, 1])
            with col_krs1:
                krs_opcion = st.radio(
                    "Zona climática (índice de aridez P/ET0):",
                    [
                        "Árida / Semiárida   (P/ET0 < 0.5) — kRS = 0.0023  [original HS]",
                        "Subhúmeda seca       (P/ET0 0.5–0.65) — kRS = 0.0021",
                        "Subhúmeda húmeda    (P/ET0 0.65–1.0) — kRS = 0.00195",
                        "Húmeda / Muy húmeda (P/ET0 > 1.0)   — kRS = 0.00185 [FAO-56]",
                        "Personalizado",
                    ],
                    key="krs_opcion",
                    index=0,
                )
            with col_krs2:
                _krs_map = {
                    "Árida / Semiárida": 0.0023,
                    "Subhúmeda seca": 0.0021,
                    "Subhúmeda húmeda": 0.00195,
                    "Húmeda / Muy húmeda": 0.00185,
                    "Personalizado": None,
                }
                _krs_key = krs_opcion.split("(")[0].strip()
                krs_auto = _krs_map.get(_krs_key, 0.0023)
                if krs_auto is None:
                    krs_valor = st.number_input(
                        "kRS personalizado:", value=0.0023,
                        min_value=0.0015, max_value=0.0030,
                        step=0.00005, format="%.5f", key="krs_custom"
                    )
                else:
                    krs_valor = krs_auto
                    st.metric("kRS seleccionado", f"{krs_valor:.5f}")

            # Referencia rápida de zonas colombianas
            st.caption(
                "**Referencia zonas Colombia:** "
                "La Guajira / Tatacoa → Árida (0.0023) · "
                "Tolima / Huila valle → Semiárida (0.0023) · "
                "Cundinamarca / Boyacá → Subhúmeda (0.00195–0.0021) · "
                "Eje cafetero / Antioquia → Húmeda (0.00185)"
            )

        if st.button("Obtener Datos NASA y Calcular Balance", type="primary"):
            with st.spinner('Consultando a la NASA y calculando variables hídricas... 🚀'):
                try:
                    df_base_diario = preparar_base_nasa(lat_input, lon_input, fecha_inicio, fecha_fin, krs=krs_valor)
                    st.session_state['krs_usado_tab1'] = krs_valor
                    st.success(f"✅ ¡Datos de NASA descargados con éxito! kRS aplicado = {krs_valor:.5f}")
                except Exception as e:
                    st.error(f"Error técnico con NASA: {e}")

    # ==========================================
    # RAMA 2: EXTRACCIÓN WAPOR v3 (.ZIP)
    # ==========================================
    elif fuente_datos == "WaPOR v3 (Archivos Raster .ZIP)":
        st.info("Sube los archivos .zip que contienen los TIFs de cada variable. El sistema cruzará la información con la coordenada ingresada.")
        
        st.caption("📦 Archivos ZIP de hasta **2 GB** soportados. Si el servidor rechaza archivos grandes, añade `server.maxUploadSize = 2048` a `.streamlit/config.toml`.")
        col_w1, col_w2, col_w3 = st.columns(3)
        with col_w1:
            zip_precip = st.file_uploader("ZIP Precipitación (hasta 2 GB)", type="zip", help="Contiene GeoTIFFs decadales de precipitación WaPOR v3")
        with col_w2:
            zip_evap = st.file_uploader("ZIP Evaporación (hasta 2 GB)", type="zip", help="Contiene GeoTIFFs decadales de evaporación WaPOR v3")
        with col_w3:
            zip_ret = st.file_uploader("ZIP Evapotranspiración RET (hasta 2 GB)", type="zip", help="Contiene GeoTIFFs decadales de RET WaPOR v3")
            
        if st.button("Procesar Datos WaPOR y Calcular Balance", type="primary"):
            if zip_precip and zip_evap and zip_ret:
                with st.spinner('Extrayendo píxeles de los archivos Raster. Esto puede tardar unos segundos... 🛰️'):
                    try:
                        # ── Paso 1: extraer series crudas de cada ZIP ──────────────────────
                        # P viene en TIFs DIARIOS (1 archivo/día → mm/día)
                        # E y RET vienen en TIFs DECADALES (1 archivo/década → mm/década)
                        df_p, rep_p = procesar_zip_wapor(zip_precip, lon_input, lat_input, 'Precipitacion')
                        df_e, rep_e = procesar_zip_wapor(zip_evap,   lon_input, lat_input, 'Evaporacion')
                        df_r, rep_r = procesar_zip_wapor(zip_ret,    lon_input, lat_input, 'RET')

                        # ── Paso 2: diagnóstico de extracción ─────────────────────────────
                        for nombre_var, rep_var in [
                            ("Precipitación (diaria)", rep_p),
                            ("Evaporación (decadal)",  rep_e),
                            ("RET (decadal)",          rep_r),
                        ]:
                            fuera = rep_var.get('archivos_fuera_extent', 0)
                            st.caption(
                                f"📋 {nombre_var}: {rep_var['archivos_tif']} TIFs · "
                                f"{rep_var['archivos_validos']} válidos · "
                                f"{rep_var['archivos_sin_fecha']} sin fecha · "
                                f"{rep_var['archivos_con_error']} error lectura · "
                                f"{fuera} fuera del extent"
                            )

                        # ── Paso 3: validar que hay datos en cada variable ─────────────────
                        faltantes = []
                        if df_p.empty:
                            faltantes.append(
                                f"Precipitación: 0 registros válidos — "
                                f"verifica que la coordenada ({lat_input:.4f}, {lon_input:.4f}) "
                                f"cae dentro del extent de los TIF de precipitación."
                            )
                        if df_e.empty:
                            faltantes.append(
                                f"Evaporación: 0 registros válidos — "
                                f"verifica coordenada y extent del ZIP de evaporación."
                            )
                        if df_r.empty:
                            faltantes.append(
                                f"RET: 0 registros válidos — "
                                f"verifica coordenada y extent del ZIP de RET."
                            )

                        if faltantes:
                            st.error("No se pudieron extraer fechas/datos válidos de uno o más ZIP.")
                            for msg in faltantes:
                                st.warning(msg)

                        else:
                            # ── Paso 4: agregar P diaria → acumulado decadal ──────────────
                            # P es diaria → hay que sumar todos los días de cada década
                            # antes de unir con E y RET (que ya son decadales).
                            # Se asigna la década usando la misma lógica que agregar_decadas():
                            #   Día 1-10  → Decada_Mes 1
                            #   Día 11-20 → Decada_Mes 2
                            #   Día 21-31 → Decada_Mes 3
                            # La "fecha representativa" de cada década se fija al día 1, 11 o 21
                            # para que el merge posterior con E y RET funcione correctamente.

                            df_p_dec = df_p.copy()
                            df_p_dec['Año']  = df_p_dec['Fecha'].dt.year
                            df_p_dec['Mes']  = df_p_dec['Fecha'].dt.month
                            df_p_dec['Dia']  = df_p_dec['Fecha'].dt.day
                            df_p_dec['Decada_Mes'] = np.select(
                                [df_p_dec['Dia'] <= 10, df_p_dec['Dia'] <= 20],
                                [1, 2], default=3
                            )
                            # Fecha canónica de inicio de cada década (día 1, 11 o 21)
                            dia_inicio_dec = {1: 1, 2: 11, 3: 21}
                            df_p_dec['Fecha_Decada'] = df_p_dec.apply(
                                lambda r: pd.Timestamp(
                                    int(r['Año']), int(r['Mes']),
                                    dia_inicio_dec[int(r['Decada_Mes'])]
                                ), axis=1
                            )

                            # Sumar precipitación diaria por (Año, Mes, Decada_Mes)
                            df_p_sum = (
                                df_p_dec
                                .groupby(['Fecha_Decada', 'Año', 'Mes', 'Decada_Mes'])['Precipitacion']
                                .sum()
                                .reset_index()
                                .rename(columns={'Fecha_Decada': 'Fecha'})
                            )

                            n_decadas_p = len(df_p_sum)
                            n_dias_p    = len(df_p)
                            st.info(
                                f"🌧️ **Precipitación:** {n_dias_p} registros diarios "
                                f"→ sumados en {n_decadas_p} décadas · "
                                f"P máx decadal = {df_p_sum['Precipitacion'].max():.1f} mm · "
                                f"P media decadal = {df_p_sum['Precipitacion'].mean():.1f} mm"
                            )

                            # ── Paso 5: E y RET decadales ya tienen fecha de inicio de década.
                            # Normalizar su fecha al día 1/11/21 para garantizar el merge.
                            def _normalizar_fecha_decadal(df_var):
                                """Ajusta la fecha de TIFs decadales al día 1, 11 o 21 del mes."""
                                df_out = df_var.copy()
                                dia = df_out['Fecha'].dt.day
                                decada_mes = np.select(
                                    [dia <= 10, dia <= 20], [1, 2], default=3
                                )
                                dia_canon = np.select(
                                    [decada_mes == 1, decada_mes == 2], [1, 11], default=21
                                )
                                df_out['Fecha'] = df_out.apply(
                                    lambda r: pd.Timestamp(r['Fecha'].year, r['Fecha'].month,
                                                           int(dia_canon[r.name])), axis=1
                                )
                                return df_out

                            df_e_norm = _normalizar_fecha_decadal(df_e)
                            df_r_norm = _normalizar_fecha_decadal(df_r)

                            # ── Paso 6: merge final P_decadal + E_decadal + RET_decadal ────
                            # Usar outer en P→E para conservar décadas aunque E no tenga dato,
                            # luego inner con RET (ambas deben existir para el balance).
                            df_clima = pd.merge(
                                df_p_sum[['Fecha', 'Precipitacion']],
                                df_e_norm[['Fecha', 'Evaporacion']],
                                on='Fecha', how='inner'
                            )
                            df_clima = pd.merge(
                                df_clima,
                                df_r_norm[['Fecha', 'RET']],
                                on='Fecha', how='inner'
                            )

                            if df_clima.empty:
                                st.warning(
                                    "No se encontraron décadas coincidentes entre los tres ZIP. "
                                    "Verifica que los periodos de las series se solapan."
                                )
                            else:
                                # df_clima tiene una fila por DÉCADA con P acumulada + E y RET
                                # El bloque común lo tratará igual que los datos NASA/WaPOR anteriores.
                                df_base_diario = df_clima
                                # Derivar fechas reales de la serie WaPOR desde los TIFs procesados
                                fecha_inicio = df_clima['Fecha'].min().date()
                                fecha_fin    = df_clima['Fecha'].max().date()
                                st.success(
                                    f"✅ ¡Datos WaPOR procesados correctamente! "
                                    f"{len(df_clima)} décadas con P acumulada + E + RET · "
                                    f"Serie: {fecha_inicio} → {fecha_fin}."
                                )

                    except Exception as e:
                        st.error(f"Error procesando los archivos Raster: {e}")
            else:
                st.warning("Por favor, sube los 3 archivos ZIP para continuar.")


    # ==========================================
    # BLOQUE COMÚN: PROCESAMIENTO DECADAL Y GRÁFICAS
    # (Se ejecuta si df_base_diario fue llenado por NASA o por WaPOR)
    # ==========================================
    if df_base_diario is not None and not df_base_diario.empty:

        # ── Cuando la fuente es WaPOR, df_base_diario ya contiene datos DECADALES
        # (P acumulada + E + RET por década) con Fecha = día 1/11/21 de cada mes.
        # En ese caso agregar_decadas() es suficiente para recuperar Año/Mes/Decada_Año
        # y el groupby().sum() sobre 1 registro/grupo es idempotente (correcto).
        #
        # Cuando la fuente es NASA, df_base_diario contiene datos DIARIOS y el
        # groupby().sum() acumula ~10 días reales por década (también correcto).
        # Ambas rutas convergen en df_decadal_anual con la misma estructura.

        df_base_diario = agregar_decadas(df_base_diario)

        # 1. Acumular por (Año, Decada_Año)
        df_decadal_anual = (
            df_base_diario
            .groupby(['Año', 'Decada_Año'])[['Precipitacion', 'Evaporacion', 'RET']]
            .sum()
            .reset_index()
        )

        # ── Diagnóstico WaPOR (mostrar antes del cálculo P75%) ──
        if 'WaPOR' in fuente_datos:
            n_anios = df_base_diario['Año'].nunique()
            n_filas = len(df_base_diario)
            pct_ceros = (df_base_diario['Precipitacion'] == 0.0).mean() * 100
            p_max = df_base_diario['Precipitacion'].max()
            p_media = df_base_diario['Precipitacion'].mean()
            temporal = "diaria" if n_filas > n_anios * 40 else "decadal"
            st.info(
                f"📡 **Diagnóstico WaPOR ({temporal}):** "
                f"{n_filas} registros · {n_anios} años · "
                f"P máx={p_max:.2f} mm · P media={p_media:.2f} mm · "
                f"{pct_ceros:.1f}% de registros con P=0 mm (excluidos del P75%)"
            )

        # 2. Promediar Evaporación y RET a lo largo del periodo
        df_promedio_decadal = (
            df_decadal_anual
            .groupby('Decada_Año')[['Evaporacion', 'RET']]
            .mean()
            .reset_index()
        )

        # 3. Precipitación P75% — método Blom (Critchley & Siegert 1996)
        # Se excluyen décadas con P=0 mm de la distribución de probabilidad,
        # porque el 0 puede representar tanto lluvia nula real como nodata residual.
        # El resultado es el valor decadal con 75% de probabilidad de ser excedido.
        def calcular_p75_blom(serie):
            """
            P75% por posición de Blom:
            P(%) = (m - 0.375) / (N + 0.25) × 100
            Ordenando de mayor a menor, el valor donde P cruza 75%.
            Si todas las observaciones son 0 (temporada seca), retorna 0.
            """
            s = np.array([v for v in serie if v > 0.0], dtype=float)
            N = len(s)
            if N == 0:
                return 0.0
            s_desc = np.sort(s)[::-1]           # mayor → menor
            m_arr = np.arange(1, N + 1)
            prob  = (m_arr - 0.375) / (N + 0.25) * 100.0
            idx   = np.searchsorted(prob, 75.0)
            idx   = min(idx, N - 1)
            return float(s_desc[idx])

        precip_75_serie = (
            df_decadal_anual
            .groupby('Decada_Año')['Precipitacion']
            .apply(calcular_p75_blom)
            .reset_index()
        )
        precip_75_serie.rename(columns={'Precipitacion': 'Prec_75%'}, inplace=True)

        # Unir precipitación P75% con promedios de Evaporación y RET
        df_promedio_decadal = pd.merge(df_promedio_decadal, precip_75_serie, on='Decada_Año')
        
        # --- EVIDENCIA DEL ANÁLISIS DE PROBABILIDAD (BLOM) ---
        st.subheader("🌧️ Análisis de Probabilidad (Precipitación ordenada)")
        try:
            df_prob_global = pd.DataFrame()
            for decada in range(1, 37):
                serie_decada = df_decadal_anual[df_decadal_anual['Decada_Año'] == decada]['Precipitacion'].values
                df_prob_global[f'D{decada}'] = np.sort(serie_decada)[::-1]
                
            N_anios = len(df_prob_global)
            m_arr = np.arange(1, N_anios + 1)
            prob_arr = ((m_arr - 0.375) / (N_anios + 0.25)) * 100
            
            df_prob_global.insert(0, 'Probabilidad Blom (%)', prob_arr)
            df_prob_global.insert(0, 'Orden (m)', m_arr)
            
            with st.expander("Ver matriz completa de ordenamiento y probabilidades"):
                st.dataframe(df_prob_global.round(2).style.format("{:.2f}", subset=['Probabilidad Blom (%)'] + [f'D{i}' for i in range(1, 37)]))
        except Exception as e:
            st.warning("No se pudo generar la tabla de visualización de probabilidad.")

        # --- GRÁFICA INTERACTIVA ---
        st.subheader("📈 Comportamiento Hídrico Decadal")
        df_melted = df_promedio_decadal.melt(
            id_vars=['Decada_Año'], value_vars=['Prec_75%', 'Evaporacion', 'RET'],
            var_name='Variable', value_name='Volumen (mm/década)'
        )
        
        import plotly.express as px
        fig = px.line(
            df_melted, x='Decada_Año', y='Volumen (mm/década)', color='Variable', 
            markers=True, color_discrete_map={'Prec_75%': '#1f77b4', 'Evaporacion': '#ff7f0e', 'RET': '#d62728'},
            labels={'Decada_Año': 'Década del Año (1 al 36)'}
        )
        fig.update_layout(xaxis=dict(tickmode='linear', tick0=1, dtick=1), hovermode="x unified")
        st.plotly_chart(fig)
        
        # --- GENERAR IMAGEN ESTÁTICA SEGURA (Matplotlib) ---
        try:
            import matplotlib.pyplot as plt
            import io
            fig_static, ax = plt.subplots(figsize=(8, 4))
            ax.plot(df_promedio_decadal['Decada_Año'], df_promedio_decadal['Prec_75%'], label='Prec_75%', color='#1f77b4', marker='o', markersize=4)
            ax.plot(df_promedio_decadal['Decada_Año'], df_promedio_decadal['Evaporacion'], label='Evaporacion', color='#ff7f0e', marker='o', markersize=4)
            ax.plot(df_promedio_decadal['Decada_Año'], df_promedio_decadal['RET'], label='RET', color='#d62728', marker='o', markersize=4)
            ax.set_title('Comportamiento Hídrico Decadal', fontsize=12)
            ax.set_xlabel('Década del Año (1 al 36)', fontsize=10)
            ax.set_ylabel('Volumen (mm/década)', fontsize=10)
            ax.legend()
            ax.grid(True, linestyle='--', alpha=0.7)
            
            img_buffer = io.BytesIO()
            fig_static.savefig(img_buffer, format='png', bbox_inches='tight', dpi=150)
            img_buffer.seek(0)
            st.session_state['imagen_clima_bytes'] = img_buffer.getvalue()
            plt.close(fig_static) 
        except Exception as e:
            st.session_state['imagen_clima_bytes'] = None

        # --- GUARDADO EN VARIABLES DE SESIÓN PARA LA PESTAÑA 2 ---
        st.session_state['latitud'] = lat_input
        st.session_state['longitud'] = lon_input
        # fecha_inicio/fin solo existen cuando la fuente es NASA POWER;
        # para WaPOR se conserva lo que ya haya en session_state (o el default).
        st.session_state['fecha_inicio_t1'] = fecha_inicio
        st.session_state['fecha_fin_t1'] = fecha_fin
        st.session_state['df_promedio'] = df_promedio_decadal
        st.session_state['df_base_diario_tab1'] = df_base_diario.copy()

        # --- MEMORIA SEPARADA POR FUENTE (NASA vs WaPOR) para Pestaña 4 ---
        if 'NASA' in fuente_datos:
            st.session_state['df_base_nasa']      = df_base_diario.copy()
            st.session_state['df_decadal_nasa']   = df_promedio_decadal.copy()
            st.session_state['fuente_nasa_lista']  = True
            st.session_state['anios_nasa']         = int(df_base_diario['Año'].nunique())
        else:
            st.session_state['df_base_wapor']     = df_base_diario.copy()
            st.session_state['df_decadal_wapor']  = df_promedio_decadal.copy()
            st.session_state['fuente_wapor_lista'] = True
            st.session_state['anios_wapor']        = int(df_base_diario['Año'].nunique())
        
        # =====================================================================
        # TABLA DE REVISIÓN DE VALORES DIARIOS EXTRAÍDOS — solo para WaPOR
        # =====================================================================
        if 'WaPOR' in fuente_datos:
            st.subheader("🔍 Revisión de valores extraídos por década (WaPOR)")
            st.caption(
                "Tabla de auditoría: muestra la **precipitación acumulada por década** "
                "(suma de todos los días de cada periodo de 10/8/11 días) junto con los "
                "valores decadales de evaporación y RET. "
                "Permite verificar que el píxel corresponde a la ubicación correcta y "
                "que los valores son coherentes con la climatología de la zona."
            )

            # Construir tabla de revisión desde df_base_diario (antes del groupby)
            df_revision = df_base_diario[['Fecha', 'Año', 'Mes', 'Día',
                                          'Precipitacion', 'Evaporacion', 'RET']].copy()
            df_revision['Fecha'] = df_revision['Fecha'].dt.strftime('%Y-%m-%d')
            df_revision = df_revision.rename(columns={
                'Precipitacion': 'P acum. (mm/décad)',
                'Evaporacion':   'E (mm/décad)',
                'RET':           'RET (mm/décad)',
            })

            # ── Filtros de exploración ──
            col_f1, col_f2, col_f3 = st.columns(3)
            with col_f1:
                anios_disp = sorted(df_revision['Año'].unique())
                anio_sel = st.selectbox(
                    "Filtrar por año", options=["Todos"] + [int(a) for a in anios_disp],
                    key="rev_anio"
                )
            with col_f2:
                meses_disp = sorted(df_revision['Mes'].unique())
                mes_nombres = {1:"Enero",2:"Febrero",3:"Marzo",4:"Abril",5:"Mayo",6:"Junio",
                               7:"Julio",8:"Agosto",9:"Septiembre",10:"Octubre",
                               11:"Noviembre",12:"Diciembre"}
                mes_opciones = ["Todos"] + [f"{m} – {mes_nombres[m]}" for m in meses_disp]
                mes_sel = st.selectbox("Filtrar por mes", options=mes_opciones, key="rev_mes")
            with col_f3:
                p_min_rev = st.number_input(
                    "Mostrar solo P ≥ (mm/décad)", value=0.0, min_value=0.0,
                    step=1.0, format="%.1f", key="rev_pmin"
                )

            # Aplicar filtros
            df_rev_filt = df_revision.copy()
            if anio_sel != "Todos":
                df_rev_filt = df_rev_filt[df_rev_filt['Año'] == int(anio_sel)]
            if mes_sel != "Todos":
                mes_num = int(mes_sel.split(" – ")[0])
                df_rev_filt = df_rev_filt[df_rev_filt['Mes'] == mes_num]
            if p_min_rev > 0.0:
                df_rev_filt = df_rev_filt[df_rev_filt['P acum. (mm/décad)'] >= p_min_rev]

            # ── Métricas de resumen del filtro activo ──
            n_dias = len(df_rev_filt)
            n_lluvia = (df_rev_filt['P acum. (mm/décad)'] > 0).sum()
            p_max_filt = df_rev_filt['P acum. (mm/décad)'].max() if n_dias > 0 else 0.0
            p_sum_filt = df_rev_filt['P acum. (mm/décad)'].sum()

            mc1, mc2, mc3, mc4 = st.columns(4)
            mc1.metric("Décadas visibles", f"{n_dias:,}")
            mc2.metric("Décadas con lluvia (P > 0)", f"{n_lluvia:,}")
            mc3.metric("P máxima (mm/décad)", f"{p_max_filt:.1f}")
            mc4.metric("P acumulada filtro (mm)", f"{p_sum_filt:.1f}")

            # ── Tabla con color condicional por intensidad de lluvia ──
            def _color_precip(val):
                """Colorea la celda de precipitación según intensidad."""
                try:
                    v = float(val)
                except (TypeError, ValueError):
                    return ''
                if v == 0.0:
                    return 'background-color: #f5f5f5; color: #aaa'
                elif v < 5.0:
                    return 'background-color: #e3f2fd'          # azul muy claro
                elif v < 15.0:
                    return 'background-color: #90caf9'          # azul claro
                elif v < 30.0:
                    return 'background-color: #42a5f5'          # azul medio
                else:
                    return 'background-color: #1565c0; color: white'  # azul oscuro

            df_display = df_rev_filt.reset_index(drop=True)

            styled = (
                df_display
                .style
                .format({
                    'P acum. (mm/décad)': '{:.1f}',
                    'E (mm/décad)':       '{:.2f}',
                    'RET (mm/décad)':     '{:.2f}',
                    'Año':                '{:d}',
                    'Mes':                '{:d}',
                    'Día':                '{:d}',
                })
                .map(_color_precip, subset=['P acum. (mm/décad)'])
                .set_properties(**{'font-size': '12px'})
            )

            st.dataframe(styled, use_container_width=True, height=420)

            # Leyenda de colores
            st.caption(
                "🎨 **Leyenda precipitación decadal:** "
                "⬜ Sin lluvia (0 mm) · "
                "🔵 Ligera (< 5 mm) · "
                "🔷 Moderada (5–15 mm) · "
                "💧 Fuerte (15–30 mm) · "
                "🌊 Muy fuerte (≥ 30 mm)"
            )

            # ── Descarga de la tabla de revisión filtrada ──
            csv_rev = df_display.to_csv(index=False, sep=";")
            st.download_button(
                "📥 Descargar tabla de revisión (CSV)",
                data=csv_rev,
                file_name=f"Revision_Diaria_WaPOR_{lat_input:.4f}_{lon_input:.4f}.csv",
                mime="text/csv",
                key="btn_down_revision"
            )

            st.divider()

        # --- DESCARGA DE DATOS ---
        st.subheader("📥 Descarga de Resultados")

        # ── Diagnóstico de valores crudos (solo para NASA) ──────────────
        if 'NASA' in fuente_datos:
            st.subheader("🔬 Verificación de valores crudos NASA POWER")
            st.caption(
                "Esta tabla muestra los datos **diarios originales** tal como los entrega la API de NASA POWER, "
                "sin ninguna transformación. Los valores de Precipitación son **mm/día**. "
                "Al acumular ~10 días por década, la precipitación decadal típica en Colombia "
                "oscila entre 5 y 150 mm/décad según la zona y la época del año."
            )

            df_raw_show = df_base_diario[['Fecha', 'Año', 'Mes', 'Día',
                                          'Precipitacion', 'Evaporacion', 'RET']].copy()
            df_raw_show.columns = ['Fecha', 'Año', 'Mes', 'Día',
                                   'P diaria (mm/día)', 'Evap diaria (mm/día)', 'RET diaria (mm/día)']

            # Métricas de diagnóstico
            p_raw_max   = df_base_diario['Precipitacion'].max()
            p_raw_media = df_base_diario['Precipitacion'].mean()
            p_raw_ceros = (df_base_diario['Precipitacion'] == 0).mean() * 100
            n_anios_raw = df_base_diario['Año'].nunique()
            n_dias_raw  = len(df_base_diario)

            mc1, mc2, mc3, mc4, mc5 = st.columns(5)
            mc1.metric("Días totales",        f"{n_dias_raw:,}")
            mc2.metric("Años",                f"{n_anios_raw}")
            mc3.metric("P máx diaria",        f"{p_raw_max:.2f} mm/día")
            mc4.metric("P media diaria",      f"{p_raw_media:.2f} mm/día")
            mc5.metric("% días sin lluvia",   f"{p_raw_ceros:.1f}%")

            # Equivalencia decadal estimada
            p_dec_estimada = p_raw_media * 10
            st.info(
                f"📐 **Interpretación:** P media diaria = {p_raw_media:.2f} mm/día → "
                f"P decadal estimada ≈ **{p_dec_estimada:.1f} mm/décad** (× 10 días). "
                f"Si los valores de P en la gráfica parecen muy bajos, verifique que la zona "
                f"y el período consultados corresponden a lo esperado. "
                f"En zonas áridas (La Guajira, Tatacoa) P media < 1 mm/día es normal en verano."
            )

            # Alerta si P media parece sospechosamente baja
            if p_raw_media < 0.5:
                st.warning(
                    f"⚠️ **P media diaria = {p_raw_media:.3f} mm/día** — Este valor es muy bajo. "
                    "Posibles causas: (1) Zona muy árida, (2) Período mayormente seco, "
                    "(3) Coordenadas incorrectas, (4) Cache de consulta anterior. "
                    "Verifique las coordenadas y limpie el caché con Ctrl+F5 o reiniciando la app."
                )
            elif p_raw_media < 1.5:
                st.warning(
                    f"⚠️ **P media diaria = {p_raw_media:.3f} mm/día** — Valor bajo. "
                    "Puede corresponder a una zona semiárida o a un período predominantemente seco. "
                    "Confirme que las coordenadas y el período son los correctos."
                )

            with st.expander("👁️ Ver datos diarios crudos de NASA POWER (primeros 60 registros)"):
                st.dataframe(
                    df_raw_show.head(60).style.format({
                        'P diaria (mm/día)':   '{:.3f}',
                        'Evap diaria (mm/día)': '{:.3f}',
                        'RET diaria (mm/día)':  '{:.3f}',
                    }).background_gradient(subset=['P diaria (mm/día)'], cmap='Blues'),
                    use_container_width=True
                )
                st.caption(
                    "🔵 Degradado azul: mayor intensidad = mayor precipitación diaria. "
                    "Valores negativos (-999) son datos faltantes reemplazados por 0. "
                    "EVPTRNS = Transpiration de vegetación (no evapotranspiración total). "
                    "RET calculada por Hargreaves-Samani sobre T_Max y T_Min de NASA."
                )

            # Mostrar también el resumen decadal para comparar
            with st.expander("📊 Ver resumen decadal generado (P75% y promedios)"):
                df_dec_show = df_decadal_anual.copy()
                df_dec_show.columns = ['Año', 'Décad/Año', 'P acum décad (mm)', 'E acum décad (mm)', 'RET acum décad (mm)']
                st.caption("Estos son los valores SUM de cada ~10 días que luego entran al cálculo P75%:")
                st.dataframe(df_dec_show.round(2).head(72), use_container_width=True)

        col_down1, col_down2 = st.columns(2)
        
        with col_down1:
            with st.expander("Ver tabla de resultados decadales (P75%, E y RET promedio)"):
                st.dataframe(df_promedio_decadal.round(2).style.format("{:.2f}"))
            csv_promedio = df_promedio_decadal.round(2).to_csv(index=False, sep=";")
            st.download_button("📥 Descargar Datos Decadales (CSV)", data=csv_promedio, file_name=f"Balance_Decadal_{lat_input}_{lon_input}.csv", mime="text/csv", key="btn_down_promedios")

        with col_down2:
            _lbl_diario = "diarios (mm/día)" if 'NASA' in fuente_datos else "decadales (mm/décad)"
            with st.expander(f"Ver datos {_lbl_diario} originales"):
                df_show_orig = df_base_diario[['Fecha', 'Año', 'Mes', 'Día', 'Precipitacion', 'Evaporacion', 'RET']].copy()
                if 'NASA' in fuente_datos:
                    df_show_orig = df_show_orig.rename(columns={
                        'Precipitacion': 'P (mm/día)',
                        'Evaporacion':   'Evap (mm/día)',
                        'RET':           'RET (mm/día)'
                    })
                    st.caption("⚠️ Estos son valores **diarios** (mm/día) — la precipitación decadal se obtiene sumando ~10 días.")
                else:
                    df_show_orig = df_show_orig.rename(columns={
                        'Precipitacion': 'P acum (mm/décad)',
                        'Evaporacion':   'E (mm/décad)',
                        'RET':           'RET (mm/décad)'
                    })
                    st.caption("Valores decadales WaPOR (ya acumulados por décad).")
                st.dataframe(df_show_orig.head(100))
                st.caption("Mostrando los primeros 100 registros.")
            csv_diario = df_base_diario.round(2).to_csv(index=False, sep=";")
            st.download_button(f"📥 Descargar Serie {_lbl_diario} (CSV)", data=csv_diario, file_name=f"Serie_{_lbl_diario.split()[0]}_{lat_input}_{lon_input}.csv", mime="text/csv", key="btn_down_diario")
# =====================================================================
# --- PESTAÑA 2: DESCARGA AUTOMÁTICA NASA POWER --- (Sin cambios)
# =====================================================================

with tab2:
        st.markdown("### Balance Hídrico y Diseño Hidráulico")
        st.markdown("Determinación paso a paso de las necesidades netas, brutas y el dimensionamiento de caudales de diseño por sector.")
        
        # --- SECCIÓN 1: DATOS CLIMÁTICOS ---
        st.subheader("1. Ubicación y Periodo")
        fuente_clima_t2 = st.radio(
            "Fuente climática para Pestaña 2:",
            ["Usar datos procesados en Pestaña 1", "NASA POWER (API Online)"],
            horizontal=True,
            key="fuente_clima_t2"
        )

        # Leer coordenadas y fechas ejecutadas en Pestaña 1 (si existen)
        _lat_def  = float(st.session_state.get('latitud',   8.848795))
        _lon_def  = float(st.session_state.get('longitud', -73.609039))
        _fi_def   = st.session_state.get('fecha_inicio_t1', pd.to_datetime("2018-01-01"))
        _ff_def   = st.session_state.get('fecha_fin_t1',   pd.to_datetime("2025-12-31"))

        if st.session_state.get('latitud') is not None:
            st.info(
                f"📍 Coordenadas y periodo heredados de Pestaña 1: "
                f"Lat **{_lat_def:.6f}** | Lon **{_lon_def:.6f}** | "
                f"{_fi_def} → {_ff_def}. "
                f"Puedes modificarlos si necesitas un análisis diferente."
            )

        col1, col2 = st.columns(2)
        with col1:
            lat_input = st.number_input("Latitud", value=_lat_def, format="%.6f", key="lat_nasa_t2")
        with col2:
            lon_input = st.number_input("Longitud", value=_lon_def, format="%.6f", key="lon_nasa_t2")

        col3, col4 = st.columns(2)
        with col3:
            fecha_inicio = st.date_input(
                "Fecha Inicio", value=_fi_def,
                min_value=pd.to_datetime("2000-01-01"), key="fi_nasa_t2"
            )
        with col4:
            fecha_fin = st.date_input("Fecha Fin", value=_ff_def, key="ff_nasa_t2")
        st.caption(
            "ℹ️ Es posible consultar NASA POWER desde el año 2000. Si no se obtiene información para "
            "fechas más antiguas en el punto consultado, la aplicación lo señalará para que se ajuste "
            "el periodo de análisis."
        )

        # --- SECCIÓN 2: PARÁMETROS AGRONÓMICOS ---
        st.subheader("2. Parámetros Agronómicos y Fenología")
        ca1, ca2, ca3 = st.columns(3)
        with ca1:
            area_total_m2_input = st.number_input("Área Total (m²)", value=5000.0, step=100.0, min_value=1.0, key="area_tot")
            area_total = area_total_m2_input / 10000.0  # Conversión interna a ha para cálculos
        with ca2:
            num_sectores = st.number_input("Número de Sectores", value=1, min_value=1, step=1, key="num_sect")
        with ca3:
            decada_inicio = st.number_input("Década de Inicio (1-36)", min_value=1, max_value=36, value=1, key="dec_ini")
            
        siembra_escalonada = st.checkbox("¿Aplicar siembra escalonada entre sectores?", value=True, key="check_esc")
        paso_escalonamiento = st.number_input("Décadas de espera entre siembra", min_value=0, value=1, key="paso_esc") if siembra_escalonada else 0

        st.markdown("**Selección de Cultivo (Seguridad Alimentaria - FAO 56)**")
        
        # Base de datos de cultivos (Duraciones en décadas de 10 días) - Sin Plátano
        base_cultivos = {
            "Maíz (Grano Seco)": {"kc_ini": 0.30, "kc_mid": 1.20, "kc_end": 0.50, "L_ini": 3, "L_dev": 4, "L_mid": 4, "L_late": 3},
            "Maíz (Dulce/Húmedo)": {"kc_ini": 0.30, "kc_mid": 1.15, "kc_end": 1.05, "L_ini": 2, "L_dev": 3, "L_mid": 3, "L_late": 1},
            "Frijol Seco": {"kc_ini": 0.40, "kc_mid": 1.15, "kc_end": 0.35, "L_ini": 2, "L_dev": 3, "L_mid": 4, "L_late": 2},
            "Yuca (Cassava)": {"kc_ini": 0.30, "kc_mid": 1.10, "kc_end": 0.50, "L_ini": 2, "L_dev": 4, "L_mid": 15, "L_late": 6},
            "Ñame (Yam)": {"kc_ini": 0.30, "kc_mid": 1.10, "kc_end": 0.60, "L_ini": 6, "L_dev": 8, "L_mid": 12, "L_late": 4},
            # Cacao (Theobroma cacao) - cultivo perenne de cobertura permanente, SIN valor
            # oficial en la Tabla 12 de FAO-56 (no está entre los ~84 cultivos tabulados).
            # Valores adaptados con la metodología FAO-56 (Allen et al., 2006) a partir de
            # literatura regional para plantación adulta en plena producción (dosel cerrado):
            # Kc_mid ≈ 1.10 (dosel completo, cultivo permanente) y Kc promedio ≈ 1.05.
            # Al ser perenne y de hoja persistente NO tiene fase de senescencia como los anuales;
            # aquí "L_ini/L_dev" representan el rebrote foliar y la floración-cuajado (menor Kc
            # por dosel parcialmente abierto tras poda) y "L_mid/L_late" representan llenado y
            # maduración de mazorca con dosel cerrado. Las 4 fases suman 36 décadas (año completo)
            # para reflejar su ciclo reproductivo continuo (sin período de descanso/barbecho).
            "Cacao (Theobroma cacao)": {"kc_ini": 0.90, "kc_mid": 1.10, "kc_end": 1.05, "L_ini": 4, "L_dev": 4, "L_mid": 20, "L_late": 8},
            "Personalizado": {"kc_ini": 0.40, "kc_mid": 1.10, "kc_end": 0.60, "L_ini": 3, "L_dev": 4, "L_mid": 4, "L_late": 3}
        }

        cultivo_seleccionado = st.selectbox("Seleccione el cultivo a establecer:", list(base_cultivos.keys()))
        datos_c = base_cultivos[cultivo_seleccionado]

        # Mostrar inputs editables por si el usuario quiere ajustar la base de datos localmente
        ck1, ck2, ck3, ck4 = st.columns(4)
        with ck1:
            kc_ini = st.number_input("Kc Inicial", value=datos_c["kc_ini"], step=0.05, key="kc_i")
            L_ini = st.number_input("Dur. Inicial (Décadas)", value=datos_c["L_ini"], min_value=1, key="l_i")
        with ck2:
            kc_mid = st.number_input("Kc Medio", value=datos_c["kc_mid"], step=0.05, key="kc_m")
            L_dev = st.number_input("Dur. Desarrollo", value=datos_c["L_dev"], min_value=1, key="l_d")
        with ck3:
            kc_end = st.number_input("Kc Final", value=datos_c["kc_end"], step=0.05, key="kc_f")
            L_mid = st.number_input("Dur. Media", value=datos_c["L_mid"], min_value=1, key="l_m")
        with ck4:
            st.write("") # Espaciador
            st.write("")
            L_late = st.number_input("Dur. Final (Maduración)", value=datos_c["L_late"], min_value=1, key="l_l")

        duracion_total = int(L_ini + L_dev + L_mid + L_late)
        
        # --- CONSTRUCCIÓN DE LA CURVA KC (Interpolación Lineal FAO) ---
        curva_kc = []
        # 1. Fase Inicial (Constante)
        curva_kc.extend([kc_ini] * int(L_ini))
        # 2. Fase de Desarrollo (Interpolación de kc_ini a kc_mid)
        if L_dev > 0:
            paso_dev = (kc_mid - kc_ini) / L_dev
            curva_kc.extend([kc_ini + paso_dev * (i + 1) for i in range(int(L_dev))])
        # 3. Fase Media (Constante)
        curva_kc.extend([kc_mid] * int(L_mid))
        # 4. Fase Final (Interpolación de kc_mid a kc_end)
        if L_late > 0:
            paso_late = (kc_end - kc_mid) / L_late
            curva_kc.extend([kc_mid + paso_late * (i + 1) for i in range(int(L_late))])

        # Visualización de la Curva
        import plotly.graph_objects as go
        fig_kc = go.Figure()
        fig_kc.add_trace(go.Scatter(x=list(range(1, duracion_total + 1)), y=curva_kc, mode='lines+markers', name='Curva Kc', line=dict(color='DarkGreen', width=3)))
        fig_kc.update_layout(title=f"Curva Fenológica del {cultivo_seleccionado} (Ciclo: {duracion_total*10} días)", xaxis_title="Décadas (10 días)", yaxis_title="Coeficiente de Cultivo (Kc)", height=350, margin=dict(l=0, r=0, t=40, b=0))
        st.plotly_chart(fig_kc, use_container_width=True)

        sembrar_multiple = st.checkbox("Habilitar múltiples ciclos de producción", value=True, key="check_mult") if duracion_total < 36 else False
        descanso = st.number_input("Décadas de descanso", min_value=0, value=1, key="descanso") if sembrar_multiple else 0
        
        # --- SECCIÓN 3: CONFIGURACIÓN DEL SISTEMA DE RIEGO ---
        st.subheader("3. Configuración del Sistema de Riego")
        
        # Nuevo selector para elegir el tipo de riego
        tipo_riego = st.radio("Tipo de Riego", options=["Riego por goteo", "Riego por aspersión"], horizontal=True, key="tipo_riego")
        
        if tipo_riego == "Riego por goteo":
            cs1, cs2, cs3, cs4 = st.columns(4)
            with cs1: dist_emisores = st.number_input("Dist. Emisores (m)", value=0.20, step=0.05, min_value=0.01, key="dist_e")
            with cs2: dist_laterales = st.number_input("Dist. Laterales (m)", value=1.20, step=0.05, min_value=0.01, key="dist_l")
            with cs3: emisores_planta = st.number_input("Emisores / Planta", value=2, min_value=1, key="em_pl")
            with cs4: caudal_emisor_lh = st.number_input("Caudal Gotero (L/h)", value=1.00, step=0.1, min_value=0.01, key="q_got")
        else:
            # Inputs exclusivos para aspersión
            cs1, cs2 = st.columns(2)
            with cs1: num_aspersores_ha = st.number_input("Número de aspersores por ha", value=100, min_value=1, step=1, key="num_asp_ha")
            with cs2: caudal_emisor_lh = st.number_input("Caudal del Aspersor (L/h)", value=500.0, step=10.0, min_value=1.0, key="q_asp")

        # Inputs de eficiencia
        ce1, ce2, ce3 = st.columns(3)
        with ce1: area_sombreada = st.number_input("% Sombreado", value=65.0, step=5.0, min_value=0.0, max_value=100.0, key="a_som")
        with ce2: pct_sustrato = st.number_input("% Sustrato", value=100.0, step=5.0, min_value=0.0, max_value=100.0, key="p_sus")

        # Eficiencia global generalizada = Ef.Conducción × Ef.Distribución × Ef.Aplicación,
        # concentradas en un solo valor representativo según el tipo de riego seleccionado,
        # en vez de tres campos numéricos independientes que el usuario debía multiplicar
        # mentalmente. Los rangos siguen los valores típicos de literatura de diseño de riego
        # (FAO 56 / ADR): aspersión 75-85%, goteo 85-95%.
        if tipo_riego == "Riego por goteo":
            opciones_eficiencia = {
                "85% — Típica (sistema convencional, cierto deterioro)": 0.85,
                "90% — Buena (sistema bien mantenido)": 0.90,
                "95% — Óptima (excelente, mejor caso de mercado)": 0.95,
            }
            default_ef_idx = 1  # 90%
        else:
            opciones_eficiencia = {
                "75% — Típica (aspersión convencional)": 0.75,
                "80% — Buena (aspersión bien diseñada)": 0.80,
                "85% — Óptima (mejor caso de mercado)": 0.85,
            }
            default_ef_idx = 2  # 85%
        with ce3:
            etiqueta_ef_sel = st.selectbox(
                "Eficiencia global del sistema (Cond. × Dist. × Aplicación)",
                options=list(opciones_eficiencia.keys()),
                index=default_ef_idx,
                key="ef_global_sel",
                help="Concentra en un solo valor la eficiencia de conducción, distribución y aplicación "
                     "de riego, cuyo producto es la eficiencia generalizada usada en el requerimiento bruto."
            )
        ef_global = opciones_eficiencia[etiqueta_ef_sel]

        # ---------------------------------------------------------------
        # BOTÓN DE EJECUCIÓN PRINCIPAL - PESTAÑA 2
        # ---------------------------------------------------------------
        st.divider()
        if st.button("🚀 Calcular Balance Hídrico y Caudales de Diseño", type="primary", key="btn_calcular_t2"):

            # 1. Obtener datos climáticos según la fuente elegida
            df_clima_t2 = None
            if fuente_clima_t2 == "Usar datos procesados en Pestaña 1":
                df_clima_t2 = st.session_state.get('df_base_diario_tab1', None)
                if df_clima_t2 is None:
                    st.error("⚠️ No hay datos de la Pestaña 1. Ejecuta primero el análisis climático allí.")
                    st.stop()
            else:
                # NASA POWER directo desde Tab 2 — usa el mismo kRS que eligió en Tab 1
                try:
                    with st.spinner("Descargando datos de NASA POWER..."):
                        krs_t2 = st.session_state.get('krs_usado_tab1', 0.0023)
                        df_clima_t2 = preparar_base_nasa(lat_input, lon_input, fecha_inicio, fecha_fin, krs=krs_t2)
                except Exception as e:
                    st.error(f"Error al conectar con NASA POWER: {e}")
                    st.stop()

            if df_clima_t2 is None or df_clima_t2.empty:
                st.error("⚠️ No se pudieron obtener datos climáticos. Verifica la fuente seleccionada.")
                st.stop()

            with st.spinner("Calculando balance hídrico y caudales de diseño... 💧"):

                # 2. Agregar décadas si no tienen esa columna
                if 'Decada_Año' not in df_clima_t2.columns:
                    df_clima_t2 = agregar_decadas(df_clima_t2)

                # 3. Precipitación efectiva (SCS-USDA) a nivel decadal por año
                df_decadal = df_clima_t2.groupby(['Año', 'Decada_Año'])[['Precipitacion', 'Evaporacion', 'RET']].sum().reset_index()

                def pe_scs(p):
                    umbral = 250.0 / 3.0
                    if p <= umbral:
                        return p * (125.0 - 0.2 * p) / 125.0
                    else:
                        return (125.0 / 3.0) + 0.1 * p

                # P75 por década — método Blom con exclusión de ceros
                # (coherente con el cálculo de Pestaña 1)
                def _p75_blom(s):
                    arr = np.array([v for v in s if v > 0.0], dtype=float)
                    N = len(arr)
                    if N == 0:
                        return 0.0
                    arr_desc = np.sort(arr)[::-1]
                    m_arr = np.arange(1, N + 1)
                    prob  = (m_arr - 0.375) / (N + 0.25) * 100.0
                    idx   = np.searchsorted(prob, 75.0)
                    return float(arr_desc[min(idx, N - 1)])

                prec_p75 = df_decadal.groupby('Decada_Año')['Precipitacion'].apply(
                    _p75_blom
                ).reset_index()
                prec_p75.columns = ['Decada_Año', 'P75_mm']
                prec_p75['Pe_mm'] = prec_p75['P75_mm'].apply(pe_scs)

                # Evaporacion y RET promedio decadal
                prom_decadal = df_decadal.groupby('Decada_Año')[['Evaporacion', 'RET']].mean().reset_index()
                df_balance = pd.merge(prec_p75, prom_decadal, on='Decada_Año')

                # 4. Construir curva Kc sobre las 36 décadas del año
                # Expandir la curva kc para cubrir toda la rotación anual (36 décadas)
                kc_anual = np.zeros(36)
                ciclo_total = int(L_ini + L_dev + L_mid + L_late)
                descanso_dec = int(descanso) if sembrar_multiple else 0
                ciclo_con_descanso = ciclo_total + descanso_dec

                inicio = int(decada_inicio) - 1  # índice 0-based

                if sembrar_multiple:
                    d = inicio
                    while d < 36:
                        for k_idx, kc_val in enumerate(curva_kc):
                            pos = (d + k_idx) % 36 if (d + k_idx) < 36 else None
                            if pos is None:
                                break
                            kc_anual[pos] = max(kc_anual[pos], kc_val)
                        d += ciclo_con_descanso
                        if d >= 36:
                            break
                else:
                    for k_idx, kc_val in enumerate(curva_kc):
                        pos = inicio + k_idx
                        if pos < 36:
                            kc_anual[pos] = kc_val

                df_balance['Kc'] = kc_anual
                df_balance['ETc_mm'] = df_balance['RET'] * df_balance['Kc']
                df_balance['Rn_mm']  = np.maximum(df_balance['ETc_mm'] - df_balance['Pe_mm'], 0.0)

                # 5. Requerimiento bruto según tipo de riego
                ef_total = max(ef_global, 0.01)
                df_balance['Rb_mm'] = df_balance['Rn_mm'] / ef_total
                st.session_state['t2_ef_total_calc'] = ef_total

                # 6. Caudal de diseño decadal (m³/s o L/s según tipo de riego)
                dias_decada = np.array([10,10,11, 10,10,8, 10,10,11, 10,10,10,
                                        10,10,11, 10,10,10, 10,10,11, 10,10,11,
                                        10,10,10, 10,10,11, 10,10,10, 10,10,11])

                area_total_m2 = area_total * 10000.0  # Ha → m²
                area_por_sector = area_total_m2 / max(num_sectores, 1)

                q_diseno = np.zeros(36)
                q_diseno_goteo = np.zeros(36)
                q_diseno_aspersion = np.zeros(36)

                for i in range(36):
                    dias_i = dias_decada[i]
                    rb_m  = df_balance.loc[df_balance['Decada_Año'] == i + 1, 'Rb_mm'].values
                    rb_m  = rb_m[0] if len(rb_m) > 0 else 0.0

                    # Goteo: q en L/s
                    vol_total_l = (rb_m / 1000.0) * area_por_sector * 1000.0
                    t_riego_s_got = float(st.session_state.get('t_max_val', 12)) * 3600.0 * dias_i
                    q_diseno_goteo[i] = vol_total_l / t_riego_s_got if t_riego_s_got > 0 else 0.0

                    # Aspersión: q en m³/s
                    vol_total_m3 = (rb_m / 1000.0) * area_por_sector
                    t_riego_s_asp = 86400.0 * dias_i
                    q_diseno_aspersion[i] = vol_total_m3 / t_riego_s_asp if t_riego_s_asp > 0 else 0.0

                if tipo_riego == "Riego por goteo":
                    # q en L/s: caudal total del sistema para regar un sector
                    q_diseno = q_diseno_goteo
                else:
                    # Aspersión: q en m³/s
                    q_diseno = q_diseno_aspersion

                # 7. Construir df_chrono (serie cronológica completa con Kc y Rb)
                df_chrono = df_decadal.copy()
                kc_map = dict(zip(df_balance['Decada_Año'], df_balance['Kc']))
                rb_map = dict(zip(df_balance['Decada_Año'], df_balance['Rb_mm']))
                df_chrono['Kc']    = df_chrono['Decada_Año'].map(kc_map).fillna(0)
                df_chrono['Rb_mm'] = df_chrono['Decada_Año'].map(rb_map).fillna(0)

                # 8. Guardar en session_state para Pestaña 3 y Anexo 3
                t_max_horas = 12  # valor por defecto; ajustar si hay input de horas de riego
                st.session_state['df_chrono']          = df_chrono
                st.session_state['q_diseno_decadal']   = q_diseno
                st.session_state['q_diseno_decadal_goteo']      = q_diseno_goteo
                st.session_state['q_diseno_decadal_aspersion']  = q_diseno_aspersion
                st.session_state['df_balance_t2']      = df_balance
                st.session_state['t_max']              = t_max_horas
                st.session_state['t_max_val']          = t_max_horas
                st.session_state['area_total_ha']      = area_total
                st.session_state['tipo_riego_calc']      = tipo_riego
                st.session_state['cultivo_calc']         = cultivo_seleccionado

            st.success("✅ Balance hídrico calculado correctamente. Puedes continuar en la Pestaña 3.")

            # --- VISUALIZACIÓN DE RESULTADOS ---
            st.subheader("📊 Resumen del Balance Hídrico Decadal")

            fig_balance = px.bar(
                df_balance, x='Decada_Año',
                y=['ETc_mm', 'Pe_mm', 'Rn_mm', 'Rb_mm'],
                barmode='group',
                labels={'Decada_Año': 'Década del Año (1–36)', 'value': 'mm/década', 'variable': 'Variable'},
                color_discrete_map={
                    'ETc_mm': '#2ecc71', 'Pe_mm': '#3498db',
                    'Rn_mm': '#e67e22', 'Rb_mm': '#e74c3c'
                },
                title="Balance Hídrico: ETc, Pe, Rn y Rb por Década"
            )
            fig_balance.update_layout(hovermode="x unified", height=400)
            st.plotly_chart(fig_balance, use_container_width=True)

            # Tabla resumen
            with st.expander("📋 Ver tabla de balance hídrico decadal"):
                df_mostrar = df_balance[['Decada_Año', 'P75_mm', 'Pe_mm', 'RET', 'Kc', 'ETc_mm', 'Rn_mm', 'Rb_mm']].copy()
                df_mostrar.columns = ['Década', 'P75 (mm)', 'Pe (mm)', 'RET (mm)', 'Kc', 'ETc (mm)', 'Rn (mm)', 'Rb (mm)']
                st.dataframe(df_mostrar.round(3).style.format("{:.3f}"))

            # Curva Kc resultante sobre las 36 décadas
            st.subheader("🌿 Distribución de Kc en el Año (36 Décadas)")
            fig_kc36 = px.area(
                x=list(range(1, 37)), y=kc_anual,
                labels={'x': 'Década del Año', 'y': 'Coeficiente de Cultivo (Kc)'},
                color_discrete_sequence=['#27ae60'],
                title=f"Curva Kc anual – {cultivo_seleccionado}"
            )
            fig_kc36.update_layout(height=300)
            st.plotly_chart(fig_kc36, use_container_width=True)

            st.info(
                f"💧 **Caudal máximo de diseño:** {q_diseno.max():.4f} {'L/s' if tipo_riego == 'Riego por goteo' else 'm³/s'}  |  "
                f"**Requerimiento bruto máximo:** {df_balance['Rb_mm'].max():.2f} mm/década  |  "
                f"**Eficiencia global del sistema:** {ef_total*100:.1f}%"
            )

# =====================================================================
# --- PESTAÑA 3: FUNCIONAMIENTO RESERVORIO
# =====================================================================

with tab3:
    st.markdown("### Simulación Cronológica del Reservorio")
    st.markdown("Tránsito del embalse frente a la serie climática histórica para evaluar el riesgo de déficit.")

    # ─────────────────────────────────────────────────────────────────────
    # SELECTOR DE FUENTE CLIMÁTICA + GESTIÓN DE CONSULTAS GUARDADAS
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("0. Fuente Climática para la Simulación")

    col_src1, col_src2 = st.columns([3, 1])
    with col_src1:
        fuente_sim_t3 = st.radio(
            "Seleccione la fuente de datos climáticos para simular el reservorio:",
            ["NASA POWER (datos cargados en Pestaña 1)", "WaPOR v3 (datos cargados en Pestaña 1)"],
            horizontal=True,
            key="fuente_sim_t3"
        )

    # Estado de disponibilidad de cada fuente
    nasa_disponible  = st.session_state.get('fuente_nasa_lista', False)
    wapor_disponible = st.session_state.get('fuente_wapor_lista', False)

    col_s1, col_s2 = st.columns(2)
    with col_s1:
        if nasa_disponible:
            st.success(f"✅ NASA POWER: {st.session_state.get('anios_nasa', '?')} años disponibles")
        else:
            st.warning("⚠️ NASA POWER: sin datos — ejecute Pestaña 1 con fuente NASA primero")
    with col_s2:
        if wapor_disponible:
            st.success(f"✅ WaPOR v3: {st.session_state.get('anios_wapor', '?')} años disponibles")
        else:
            st.warning("⚠️ WaPOR v3: sin datos — ejecute Pestaña 1 con fuente WaPOR primero")

    # Seleccionar el df_chrono y q_diseno correspondiente a la fuente elegida
    # (df_chrono viene de Pestaña 2, que usa df_base_diario_tab1; acá remapeamos si el usuario cambió de fuente)
    usar_nasa_sim  = "NASA" in fuente_sim_t3
    usar_wapor_sim = "WaPOR" in fuente_sim_t3

    # Gestión de consultas guardadas (historial de simulaciones)
    st.divider()
    st.subheader("📂 Historial de Simulaciones Guardadas")

    # Inicializar historial en session_state
    if 'historial_simulaciones' not in st.session_state:
        st.session_state['historial_simulaciones'] = {}

    col_hist1, col_hist2 = st.columns([4, 1])
    with col_hist1:
        historial = st.session_state['historial_simulaciones']
        if historial:
            nombres_guardados = list(historial.keys())
            sim_seleccionada = st.selectbox(
                "Seleccione una simulación guardada para cargarla:",
                options=["(Nueva simulación)"] + nombres_guardados,
                key="sel_historial"
            )
        else:
            st.info("No hay simulaciones guardadas. Ejecute una simulación y guárdela con el botón '💾 Guardar simulación'.")
            sim_seleccionada = "(Nueva simulación)"

    with col_hist2:
        st.write("")
        st.write("")
        if st.button("🗑️ Limpiar todo el historial", key="btn_limpiar_historial"):
            st.session_state['historial_simulaciones'] = {}
            st.rerun()

    # Cargar simulación guardada si el usuario la selecciona
    if sim_seleccionada != "(Nueva simulación)" and historial:
        datos_guardados = historial[sim_seleccionada]
        st.success(f"📥 Simulación cargada: **{sim_seleccionada}** — Fuente: {datos_guardados.get('fuente', '?')} | Años: {datos_guardados.get('n_anios', '?')} | V_max: {datos_guardados.get('v_max', 0):.1f} m³")
        col_cg1, col_cg2 = st.columns(2)
        with col_cg1:
            if st.button("📤 Usar esta simulación como activa (Tablas 7-8-9)", key="btn_usar_guardada"):
                st.session_state['df_simulacion_reservorio'] = datos_guardados['df_sim'].copy()
                st.session_state['v_rippl_optimo']           = datos_guardados.get('v_rippl', None)
                st.session_state['volumen_maximo_sistema']   = datos_guardados.get('v_max', 0)
                fuente_activa = datos_guardados.get('fuente', '')
                if 'NASA' in fuente_activa:
                    st.session_state['df_simulacion_nasa']   = datos_guardados['df_sim'].copy()
                    st.session_state['v_rippl_nasa']         = datos_guardados.get('v_rippl', None)
                else:
                    st.session_state['df_simulacion_wapor']  = datos_guardados['df_sim'].copy()
                    st.session_state['v_rippl_wapor']        = datos_guardados.get('v_rippl', None)
                st.success("✅ Simulación activa actualizada. Revise los resultados en Pestaña 4.")
        with col_cg2:
            if st.button("❌ Eliminar esta simulación del historial", key="btn_eliminar_guardada"):
                del st.session_state['historial_simulaciones'][sim_seleccionada]
                st.rerun()

        # Mostrar preview
        with st.expander("👁️ Vista previa de la simulación guardada"):
            st.dataframe(datos_guardados['df_sim'].head(20), use_container_width=True)

    st.divider()

    # Nombre para guardar la simulación
    _tipo_riego_tag = "Goteo" if st.session_state.get('tipo_riego_calc', st.session_state.get('tipo_riego', 'Riego por goteo')) == "Riego por goteo" else "Aspersion"
    nombre_guardado_input = st.text_input(
        "💾 Nombre para guardar esta simulación (se guarda al finalizar la simulación):",
        value=f"Sim_{'NASA' if usar_nasa_sim else 'WaPOR'}_{_tipo_riego_tag}_{pd.Timestamp.now().strftime('%Y%m%d_%H%M')}",
        key="nombre_guardado_sim",
        help="Incluya el tipo de riego en el nombre (ya sugerido por defecto) para poder distinguir "
             "en el historial las simulaciones de goteo y de aspersión sin que se sobrescriban entre sí."
    )

    st.divider()

    # 1. Selección de Infraestructura Principal
    tipo_almacenamiento = st.radio(
        "Seleccione el tipo de estructura de almacenamiento:",
        ["Opción 1: Tanque Australiano (Cilíndrico)", "Opción 3: Reservorio Excavado (Vaso Irregular)"],
        help="La Opción 3 utiliza curvas de nivel para un cálculo más preciso en terrenos irregulares."
    )

    # Variables de inicialización
    es_excavado = False
    vol_max_sistema = 0.0
    area_fija_espejo = 0.0

    if tipo_almacenamiento == "Opción 1: Tanque Australiano (Cilíndrico)":
        st.subheader("1. Dimensiones del Tanque Australiano")

        # ── Catálogo ITM: diámetros y capacidades por piso ───────────────
        _itm_diametros = [2.24, 2.98, 3.73, 4.47, 5.22, 5.96, 6.71, 7.45, 8.20, 8.94, 9.69, 10.43]
        _itm_cap_1p    = [4650, 8250, 12850, 18550, 25200, 32950, 41700, 51450, 62300, 74100, 87000, 100900]
        _itm_cap_2p    = [9050, 16050, 25100, 36100, 49150, 64200, 81250, 100300, 121400, 144450, 169550, 196650]
        _itm_cap_3p    = [13350, 23750, 37100, 53400, 72650, 94900, 120100, 148300, 179450, 213550, 250650, 290650]
        _itm_h         = {1: 1.20, 2: 2.30, 3: 3.40}  # altura estándar por piso

        # Modo de selección
        modo_tanque = st.radio(
            "Modo de dimensionamiento:",
            ["📋 Catálogo ITM (dimensiones estándar)", "✏️ Personalizado (radio y altura libre)"],
            horizontal=True, key="modo_tanque"
        )

        if modo_tanque == "📋 Catálogo ITM (dimensiones estándar)":
            # Selector de piso
            pisos_sel = st.radio(
                "Número de anillos (piso):",
                [1, 2, 3],
                format_func=lambda x: f"{x} anillo{'s' if x>1 else ''} — h = {_itm_h[x]:.2f} m",
                horizontal=True, key="pisos_itm"
            )
            h_std = _itm_h[pisos_sel]
            caps  = {1: _itm_cap_1p, 2: _itm_cap_2p, 3: _itm_cap_3p}[pisos_sel]

            # Tabla visual del catálogo para el piso seleccionado
            with st.expander(f"📊 Ver catálogo completo — {pisos_sel} anillo(s) | h = {h_std} m", expanded=True):
                df_cat = pd.DataFrame({
                    "Diámetro (m)": _itm_diametros,
                    "Radio (m)":    [round(d/2, 4) for d in _itm_diametros],
                    "Capacidad (L)": caps,
                    "Capacidad (m³)": [round(c/1000, 2) for c in caps],
                })
                # Destacar fila seleccionada
                st.dataframe(df_cat, use_container_width=True, hide_index=True)

            # Selector de referencia
            opciones_label = [
                f"Ø {_itm_diametros[i]:.2f} m — {caps[i]:,} L ({caps[i]/1000:.1f} m³)"
                for i in range(len(_itm_diametros))
            ]
            idx_sel = st.selectbox(
                "Seleccione la referencia del catálogo:",
                options=range(len(opciones_label)),
                format_func=lambda i: opciones_label[i],
                key="ref_itm"
            )

            diametro_itm   = _itm_diametros[idx_sel]
            radio_tanque   = round(diametro_itm / 2, 4)
            altura_tanque  = h_std
            cap_litros_itm = caps[idx_sel]

            # Mostrar métricas
            col_m1, col_m2, col_m3, col_m4 = st.columns(4)
            col_m1.metric("Diámetro", f"{diametro_itm:.2f} m")
            col_m2.metric("Radio", f"{radio_tanque:.4f} m")
            col_m3.metric("Altura estándar", f"{altura_tanque:.2f} m")
            col_m4.metric("Capacidad nominal", f"{cap_litros_itm:,} L  ({cap_litros_itm/1000:.1f} m³)")

            vol_calc_m3 = math.pi * (radio_tanque**2) * altura_tanque
            st.info(
                f"📐 Volumen geométrico calculado: **{vol_calc_m3:.2f} m³** "
                f"({vol_calc_m3*1000:.0f} L) — "
                f"Capacidad nominal catálogo: **{cap_litros_itm/1000:.2f} m³**"
            )

            caudal_concesion = st.number_input("Caudal Concesión Constante (L/s)", value=0.0, step=0.1, min_value=0.0, key="cc_itm")

        else:
            # Modo personalizado — comportamiento original
            col1, col2, col3 = st.columns(3)
            with col1:
                radio_tanque = st.number_input("Radio del Tanque (m)", value=5.0, step=0.5, min_value=0.5, key="radio_pers")
            with col2:
                altura_tanque = st.number_input("Altura Útil Máxima (m)", value=1.50, step=0.1, min_value=0.5, key="alt_pers")
            with col3:
                caudal_concesion = st.number_input("Caudal Concesión Constante (L/s)", value=0.0, step=0.1, min_value=0.0, key="cc_pers")

            vol_calc_m3 = math.pi * (radio_tanque**2) * altura_tanque
            col_ref1, col_ref2 = st.columns(2)
            col_ref1.metric("Volumen calculado", f"{vol_calc_m3:.2f} m³  ({vol_calc_m3*1000:.0f} L)")
            col_ref2.metric("Área espejo", f"{math.pi*(radio_tanque**2):.2f} m²")

            # Referencia rápida al catálogo: mostrar el tanque ITM más cercano
            vol_l = vol_calc_m3 * 1000
            mejor_piso, mejor_idx, mejor_diff = 1, 0, float('inf')
            for p_i, cap_list in [(1, _itm_cap_1p), (2, _itm_cap_2p), (3, _itm_cap_3p)]:
                for k, c in enumerate(cap_list):
                    diff = abs(c - vol_l)
                    if diff < mejor_diff:
                        mejor_diff, mejor_piso, mejor_idx = diff, p_i, k
            caps_ref   = {1: _itm_cap_1p,   2: _itm_cap_2p,   3: _itm_cap_3p  }[mejor_piso][mejor_idx]
            st.caption(
                f"🔍 Referencia ITM más cercana: Ø {_itm_diametros[mejor_idx]:.2f} m, "
                f"{mejor_piso} anillo(s), h={_itm_h[mejor_piso]:.2f} m — "
                f"{caps_ref:,} L"
            )

        area_fija_espejo = math.pi * (radio_tanque**2)
        vol_max_sistema  = area_fija_espejo * altura_tanque

        st.subheader("2. Cosecha de Aguas Lluvias (Opción 2)")
        habilitar_cosecha = st.checkbox("¿Implementar cosecha de aguas lluvias mediante cubierta?", value=False, key="check_cosecha")
        
        if habilitar_cosecha:
            col4, col5, col6 = st.columns(3)
            with col4:
                largo_tejado = st.number_input("Largo Tejado (m)", value=10.0, step=1.0, min_value=0.0)
            with col5:
                ancho_tejado = st.number_input("Ancho Tejado (m)", value=10.0, step=1.0, min_value=0.0)
            with col6:
                coef_escorrentia = st.number_input("Coeficiente de Escorrentía", value=0.90, step=0.05, min_value=0.0, max_value=1.0)
        else:
            largo_tejado, ancho_tejado, coef_escorrentia = 0.0, 0.0, 0.0

    else:
        st.subheader("1. Diseño de Reservorio Excavado")
        st.info("Se deshabilitan dimensiones de tanques comerciales para usar batimetría de campo.")
        es_excavado = True
        
        col1, col2 = st.columns(2)
        with col1:
            prof_max = st.number_input("Profundidad máxima (m)", value=2.0, step=0.25, min_value=0.5)
        with col2:
            caudal_concesion = st.number_input("Caudal Concesión (L/s)", value=0.0, step=0.1, min_value=0.0)

        # Generación de tabla de batimetría
        intervalos = np.arange(0, prof_max + 0.25, 0.25)
        df_bat_init = pd.DataFrame({
            "Altura (m)": intervalos,
            "Área Espejo (m2)": [0.0] * len(intervalos),
            "Volumen Acumulado (m3)": [0.0] * len(intervalos)
        })
        
        st.write("Ingrese los datos de la batimetría proyectada:")
        df_bat_usuario = st.data_editor(df_bat_init, num_rows="fixed", use_container_width=True)
        
        if df_bat_usuario["Volumen Acumulado (m3)"].max() > 0:
            vol_max_sistema = df_bat_usuario["Volumen Acumulado (m3)"].max()
            # Lógica interna para interpolación
            h_vals = df_bat_usuario["Altura (m)"].values
            a_vals = df_bat_usuario["Área Espejo (m2)"].values
            v_vals = df_bat_usuario["Volumen Acumulado (m3)"].values
            
            # Funciones de apoyo para la simulación
            func_area = lambda v: np.interp(v, v_vals, a_vals)
        else:
            st.warning("⚠️ Complete la tabla de batimetría para habilitar la simulación.")

    if st.button("Simular Tránsito del Reservorio", type="primary"):
        # Validar que la fuente seleccionada esté disponible
        if usar_nasa_sim and not nasa_disponible:
            st.error("⚠️ No hay datos NASA POWER cargados. Ejecute Pestaña 1 con fuente NASA primero.")
            st.stop()
        if usar_wapor_sim and not wapor_disponible:
            st.error("⚠️ No hay datos WaPOR v3 cargados. Ejecute Pestaña 1 con fuente WaPOR primero.")
            st.stop()

        if 'df_chrono' not in st.session_state or 'q_diseno_decadal' not in st.session_state:
            st.warning("⚠️ Por favor, ejecuta primero el cálculo en la 'Pestaña 2' para generar las matrices de demanda.")
        else:
            with st.spinner("Simulando balance volumétrico y calculando optimización agronómica... 🌊"):
                # Reconstruir df_chrono desde la fuente elegida si es posible
                # Pestaña 2 genera df_chrono desde df_base_diario_tab1 (última fuente cargada en P1)
                # Si el usuario quiere simular con la otra fuente, usamos el df_base guardado
                if usar_nasa_sim and 'df_base_nasa' in st.session_state:
                    df_base_fuente = st.session_state['df_base_nasa'].copy()
                    etiqueta_fuente = "NASA POWER"
                elif usar_wapor_sim and 'df_base_wapor' in st.session_state:
                    df_base_fuente = st.session_state['df_base_wapor'].copy()
                    etiqueta_fuente = "WaPOR v3"
                else:
                    df_base_fuente = None
                    etiqueta_fuente = "última fuente cargada"

                df_chrono = st.session_state['df_chrono'].copy()

                # Si tenemos la base de la fuente seleccionada, reconstruir df_chrono con ella
                if df_base_fuente is not None:
                    if 'Decada_Año' not in df_base_fuente.columns:
                        df_base_fuente = agregar_decadas(df_base_fuente)
                    df_chrono_fuente = df_base_fuente.groupby(['Año', 'Decada_Año'])[['Precipitacion', 'Evaporacion', 'RET']].sum().reset_index()
                    # Agregar Kc y Rb_mm desde el df_balance de Pestaña 2
                    df_balance_t2 = st.session_state.get('df_balance_t2', None)
                    if df_balance_t2 is not None:
                        kc_map_f = dict(zip(df_balance_t2['Decada_Año'], df_balance_t2['Kc']))
                        rb_map_f = dict(zip(df_balance_t2['Decada_Año'], df_balance_t2['Rb_mm']))
                        df_chrono_fuente['Kc']    = df_chrono_fuente['Decada_Año'].map(kc_map_f).fillna(0)
                        df_chrono_fuente['Rb_mm'] = df_chrono_fuente['Decada_Año'].map(rb_map_f).fillna(0)
                    df_chrono = df_chrono_fuente
                    st.info(f"🔄 Simulando con datos de **{etiqueta_fuente}** ({df_chrono['Año'].nunique()} años).")
                q_diseno_decadal = st.session_state['q_diseno_decadal']
                t_max = st.session_state.get('t_max', 12)
                area_cultivo_ha = st.session_state.get('area_total_ha', 0.5)
                tipo_riego = st.session_state.get('tipo_riego_calc', st.session_state.get('tipo_riego', 'Riego por goteo'))
                
                # ---------------------------------------------------------
                # 1. GEOMETRÍA INICIAL (AQUÍ ENTRA EL IF/ELSE DEL TIPO DE RESERVORIO)
                # ---------------------------------------------------------
                if es_excavado:
                    v_max = vol_max_sistema # Viene de la tabla batimétrica (arriba)
                    area_tejado_efectiva = 0.0 # Se asume que no hay cosecha de techos para el reservorio
                else:
                    area_tanque = math.pi * (radio_tanque ** 2)
                    v_max = area_tanque * altura_tanque
                    area_tejado_efectiva = (largo_tejado * ancho_tejado) * coef_escorrentia
                
                dias_d = np.array([10,10,11, 10,10,8, 10,10,11, 10,10,10, 10,10,11, 10,10,10, 10,10,11, 10,10,11, 10,10,10, 10,10,11, 10,10,10, 10,10,11])
                
                resultados_simulacion = []
                v_actual = v_max  # Inicia lleno
                deficit_maximo_registrado = 0.0 
                
                for index, row in df_chrono.iterrows():
                    año, decada_año = int(row['Año']), int(row['Decada_Año'])
                    decada_idx = decada_año - 1 
                    dias, p_dec_mm, e_dec_mm = dias_d[decada_idx], row['Precipitacion'], row['Evaporacion']
                    
                    # ---------------------------------------------------------
                    # 2. ÁREA DINÁMICA DE EVAPORACIÓN/LLUVIA
                    # ---------------------------------------------------------
                    if es_excavado:
                        # Evalúa el polinomio para sacar el área según el volumen que nos queda
                        area_espejo_actual = func_area(v_actual) if v_actual > 0 else func_area(0)
                    else:
                        area_espejo_actual = area_tanque

                    # ENTRADAS
                    e_cp = (caudal_concesion * 86400 * dias) / 1000.0
                    e_ll = area_espejo_actual * (p_dec_mm / 1000.0) # Lluvia directa usa el área dinámica
                    e_es = area_tejado_efectiva * (p_dec_mm / 1000.0) if not es_excavado else 0.0
                    
                    # SALIDAS
                    if tipo_riego == "Riego por goteo":
                        s_d = (q_diseno_decadal[decada_idx] * t_max * 3600 * dias) / 1000.0
                    else:
                        s_d = (q_diseno_decadal[decada_idx] * 86400 * dias) 
                    
                    s_e = area_espejo_actual * (e_dec_mm / 1000.0) # Evaporación usa el área dinámica
                    s_i = s_e * 0.10 # Asumes infiltración como 10% de evaporación
                    
                    # BALANCE
                    v_temp = v_actual + e_cp + e_ll + e_es - s_d - s_e - s_i
                    derramado, deficit_decada = 0.0, 0.0
                    
                    if v_temp > v_max:
                        v_final, derramado, estado = v_max, v_temp - v_max, "Lleno (Derrama)"
                    elif v_temp < 0:
                        v_final, deficit_decada, estado = 0.0, abs(v_temp), "Déficit Crítico ⚠️"
                        v_actual_matematico = v_actual + e_cp + e_ll + e_es - s_d - s_e - s_i
                        if abs(v_actual_matematico) > deficit_maximo_registrado: 
                            deficit_maximo_registrado = abs(v_actual_matematico)
                    else:
                        v_final, estado = v_temp, "Operación Normal"
                        
                    # ---------------------------------------------------------
                    # 3. ALTURA DE LÁMINA FINAL DE LA DÉCADA
                    # ---------------------------------------------------------
                    if es_excavado:
                        # Si la UI lo definió, interpolamos. (h_vals y v_vals se definieron al llenar la tabla)
                        altura_vaso = np.interp(v_final, v_vals, h_vals) if v_final > 0 else 0
                    else:
                        altura_vaso = v_final / area_tanque if area_tanque > 0 else 0
                    
                    resultados_simulacion.append({
                        'Año': año, 'Decada': decada_año, 'Altura Vaso (m)': round(altura_vaso, 2), 'Volumen Inicial (m3)': round(v_actual, 2),
                        'Entrada Concesion (m3)': round(e_cp, 2), 'Entrada Lluvia (m3)': round(e_ll, 2), 'Entrada Escorrentia (m3)': round(e_es, 2),
                        'Salida Riego (m3)': round(s_d, 2), 'Salida Evaporación (m3)': round(s_e, 2), 'Salida Infiltración (m3)': round(s_i, 2),
                        'Volumen Final (m3)': round(v_final, 2), 'Déficit Hídrico (m3)': round(deficit_decada, 2), 'Volumen Derramado (m3)': round(derramado, 2), 'Estado': estado
                    })
                    v_actual = v_final
                
                df_simulacion = pd.DataFrame(resultados_simulacion)
                st.session_state['df_simulacion_reservorio'] = df_simulacion # Lo guardamos para exportarlo al Word después

                # ── Guardar separado por fuente para Pestaña 4 ──
                if usar_nasa_sim:
                    st.session_state['df_simulacion_nasa']  = df_simulacion.copy()
                    st.session_state['fuente_sim_activa']   = 'NASA POWER'
                else:
                    st.session_state['df_simulacion_wapor'] = df_simulacion.copy()
                    st.session_state['fuente_sim_activa']   = 'WaPOR v3'

                # ── Guardar en historial ──
                nombre_guardado_key = st.session_state.get('nombre_guardado_sim', f'Sim_{etiqueta_fuente}')
                st.session_state['historial_simulaciones'][nombre_guardado_key] = {
                    'df_sim':  df_simulacion.copy(),
                    'fuente':  etiqueta_fuente,
                    'n_anios': df_chrono['Año'].nunique(),
                    'v_max':   v_max,
                    'v_rippl': None,  # se actualizará tras el cálculo Rippl abajo
                    'tipo_riego': st.session_state.get('tipo_riego_calc', tipo_riego),
                    'eficiencia_global': st.session_state.get('t2_ef_total_calc', None),
                    'timestamp': pd.Timestamp.now().strftime('%Y-%m-%d %H:%M'),
                }

                # --- Guardar parámetros del reservorio para el Anexo 3 ---
                st.session_state['tipo_almacenamiento_elegido'] = tipo_almacenamiento
                st.session_state['volumen_maximo_sistema']      = vol_max_sistema
                st.session_state['es_excavado_flag']            = es_excavado
                st.session_state['radio_tanque_val']            = locals().get('radio_tanque', 0.0)
                st.session_state['altura_tanque_val']           = locals().get('altura_tanque', 0.0)
                st.session_state['habilitar_cosecha_val']       = locals().get('habilitar_cosecha', False)
                st.session_state['largo_tejado_val']            = locals().get('largo_tejado', 0.0)
                st.session_state['ancho_tejado_val']            = locals().get('ancho_tejado', 0.0)
                st.session_state['area_tejado_fisica_val']      = locals().get('area_tejado_fisica', 0.0)
                st.session_state['area_tejado_efectiva_val']    = locals().get('area_tejado_efectiva', 0.0)
                st.session_state['caudal_concesion_val']        = locals().get('caudal_concesion', 0.0)
                st.session_state['t_max_riego_val']             = locals().get('t_max', 12)
                if es_excavado and 'df_bat_usuario' in locals():
                    st.session_state['df_batimetria_val']       = df_bat_usuario.copy()
                else:
                    st.session_state['df_batimetria_val']       = None

                st.success("✅ Simulación de tránsito del reservorio finalizada.")

                # =============================================================
                # GRÁFICA 1: COMPORTAMIENTO DEL VOLUMEN DEL RESERVORIO
                # =============================================================
                st.subheader("📈 Comportamiento del Volumen del Reservorio")

                etiquetas = [f"{int(r['Año'])}-D{int(r['Decada'])}" for _, r in df_simulacion.iterrows()]
                vols_final = df_simulacion['Volumen Final (m3)'].values

                fig_sim, ax_sim = plt.subplots(figsize=(14, 5))
                ax_sim.fill_between(range(len(vols_final)), vols_final, alpha=0.35, color='#2196F3', label='Volumen almacenado (m³)')
                ax_sim.plot(range(len(vols_final)), vols_final, color='#1565C0', linewidth=1.2)
                ax_sim.axhline(v_max, color='#E53935', linestyle='--', linewidth=1.2, label=f'Capacidad máxima = {v_max:.1f} m³')
                ax_sim.axhline(0, color='#FF6F00', linestyle=':', linewidth=1.0, label='Nivel cero (déficit)')

                deficit_idx = [i for i, vf in enumerate(vols_final) if vf == 0]
                if deficit_idx:
                    ax_sim.scatter(deficit_idx, [0]*len(deficit_idx), color='red', zorder=5, s=20, label='Déficit hídrico')

                tick_pos  = list(range(0, len(etiquetas), 36))
                tick_labs = [etiquetas[i].split('-')[0] for i in tick_pos]
                ax_sim.set_xticks(tick_pos)
                ax_sim.set_xticklabels(tick_labs, rotation=45, fontsize=8)
                ax_sim.set_xlabel("Año de simulación", fontsize=10)
                ax_sim.set_ylabel("Volumen (m³)", fontsize=10)
                ax_sim.set_title("Tránsito del Reservorio — Volumen almacenado por Década", fontsize=11, fontweight='bold')
                ax_sim.legend(fontsize=8)
                ax_sim.grid(axis='y', linestyle='--', alpha=0.4)
                plt.tight_layout()
                st.pyplot(fig_sim)

                buf_sim = io.BytesIO()
                fig_sim.savefig(buf_sim, format='png', dpi=150, bbox_inches='tight')
                buf_sim.seek(0)
                st.session_state['imagen_simulacion_bytes'] = buf_sim.getvalue()
                plt.close(fig_sim)

                # =============================================================
                # GRÁFICA 2: ÁREA Y VOLUMEN EN FUNCIÓN DE LA ELEVACIÓN
                # =============================================================
                st.subheader("📐 Área y Volumen en función de la Elevación")

                if es_excavado and 'h_vals' in locals() and h_vals is not None and len(h_vals) > 1:
                    coef_a = np.polyfit(h_vals, a_vals, 2)
                    coef_v = np.polyfit(h_vals, v_vals, 2)
                    h_fit  = np.linspace(h_vals.min(), h_vals.max(), 100)
                    a_fit  = np.polyval(coef_a, h_fit)
                    v_fit  = np.polyval(coef_v, h_fit)

                    fig_av, (ax_a, ax_v) = plt.subplots(1, 2, figsize=(12, 4))
                    ax_a.scatter(h_vals, a_vals, color='#1565C0', zorder=5, s=30)
                    ax_a.plot(h_fit, a_fit, color='#E53935', linewidth=2,
                              label=f'A={coef_a[0]:.3f}h²+{coef_a[1]:.3f}h+{coef_a[2]:.2f}')
                    ax_a.set_xlabel("Elevación (m)"); ax_a.set_ylabel("Área espejo (m²)")
                    ax_a.set_title("Curva Cota–Área"); ax_a.legend(fontsize=7); ax_a.grid(alpha=0.3)

                    ax_v.scatter(h_vals, v_vals, color='#2E7D32', zorder=5, s=30)
                    ax_v.plot(h_fit, v_fit, color='#F57F17', linewidth=2,
                              label=f'V={coef_v[0]:.3f}h²+{coef_v[1]:.3f}h+{coef_v[2]:.2f}')
                    ax_v.set_xlabel("Elevación (m)"); ax_v.set_ylabel("Volumen acumulado (m³)")
                    ax_v.set_title("Curva Cota–Volumen"); ax_v.legend(fontsize=7); ax_v.grid(alpha=0.3)

                else:
                    _r_aust  = locals().get('radio_tanque', 0.0)
                    _h_aust  = locals().get('altura_tanque', 1.0)
                    area_aust = math.pi * (_r_aust ** 2) if _r_aust > 0 else 1.0
                    h_lin = np.linspace(0, _h_aust, 50)
                    a_lin = np.full_like(h_lin, area_aust)
                    v_lin = area_aust * h_lin

                    fig_av, (ax_a, ax_v) = plt.subplots(1, 2, figsize=(12, 4))
                    ax_a.plot(h_lin, a_lin, color='#1565C0', linewidth=2,
                              label=f'Área constante = {area_aust:.2f} m²')
                    ax_a.set_xlabel("Elevación (m)"); ax_a.set_ylabel("Área espejo (m²)")
                    ax_a.set_title("Curva Cota–Área (Cilíndrico)"); ax_a.legend(fontsize=8); ax_a.grid(alpha=0.3)

                    ax_v.plot(h_lin, v_lin, color='#2E7D32', linewidth=2,
                              label=f'V = {area_aust:.2f} × h')
                    ax_v.set_xlabel("Elevación (m)"); ax_v.set_ylabel("Volumen (m³)")
                    ax_v.set_title("Curva Cota–Volumen (Cilíndrico)"); ax_v.legend(fontsize=8); ax_v.grid(alpha=0.3)

                plt.suptitle("Curvas Cota–Área y Cota–Volumen del Reservorio", fontsize=11, fontweight='bold')
                plt.tight_layout()
                st.pyplot(fig_av)

                buf_av = io.BytesIO()
                fig_av.savefig(buf_av, format='png', dpi=150, bbox_inches='tight')
                buf_av.seek(0)
                st.session_state['imagen_area_volumen_bytes'] = buf_av.getvalue()
                plt.close(fig_av)

                # Mostrar tabla de resultados
                st.subheader("📋 Tabla de Simulación Decadal")
                st.dataframe(df_simulacion)

               # --- PLANITO ESQUEMÁTICO A ESCALA REAL ---
                st.divider()
                st.subheader("🗺️ Esquema Espacial del Proyecto (Vista en Planta)")
                
                import plotly.graph_objects as go
                fig_esq = go.Figure()

                # --- 1. DEFINIR VARIABLES BASE PRIMERO ---
                # Traemos el área del cultivo (Si no existe, asumimos 0.5 hectáreas por defecto)
                area_cultivo_ha = st.session_state.get('area_total_ha', 0.5)
                area_cultivo_m2 = area_cultivo_ha * 10000
                lado_cultivo = math.sqrt(area_cultivo_m2)
                margen = 5.0 # Margen de separación entre obras en metros
                
                # --- 2. DEFINIR DIMENSIONES DEL ALMACENAMIENTO ---
                if es_excavado:
                    # Calcular el área máxima evaluando el polinomio en el volumen máximo
                    area_maxima = func_area(vol_max_sistema) if vol_max_sistema > 0 else 100.0
                    
                    # Asumimos una forma cuadrada para el espejo de agua máximo
                    lado_reservorio = math.sqrt(area_maxima)
                    distancia_centro = lado_reservorio / 2
                    
                    xc_tanque = lado_cultivo + margen + distancia_centro
                    yc_tanque = margen + distancia_centro
                    forma_dibujo = "rectangulo"
                    
                else:
                    # Usar el radio del tanque australiano (con valor seguro por defecto)
                    distancia_centro = locals().get('radio_tanque', 5.0)
                    xc_tanque = lado_cultivo + margen + distancia_centro
                    yc_tanque = margen + distancia_centro
                    forma_dibujo = "circulo"
                
                # --- 3. DIBUJAR EL CULTIVO ---
                fig_esq.add_shape(
                    type="rect",
                    x0=0, y0=0, x1=lado_cultivo, y1=lado_cultivo,
                    line_color="DarkGreen", fillcolor="LightGreen", opacity=0.3
                )
                fig_esq.add_annotation(x=lado_cultivo/2, y=lado_cultivo/2, text=f"Área de Cultivo<br>({area_cultivo_ha * 10000:,.0f} m²)", showarrow=False)

                # --- 4. DIBUJAR EL ALMACENAMIENTO ---
                if forma_dibujo == "circulo":
                    fig_esq.add_shape(
                        type="circle",
                        x0=xc_tanque - distancia_centro, y0=yc_tanque - distancia_centro,
                        x1=xc_tanque + distancia_centro, y1=yc_tanque + distancia_centro,
                        line_color="DarkBlue", fillcolor="LightSkyBlue"
                    )
                    fig_esq.add_annotation(x=xc_tanque, y=yc_tanque, text="Tanque<br>Australiano", showarrow=False)
                
                elif forma_dibujo == "rectangulo":
                    fig_esq.add_shape(
                        type="rect",
                        x0=xc_tanque - distancia_centro, y0=yc_tanque - distancia_centro,
                        x1=xc_tanque + distancia_centro, y1=yc_tanque + distancia_centro,
                        line_color="SaddleBrown", fillcolor="MediumTurquoise",
                        opacity=0.8
                    )
                    fig_esq.add_annotation(x=xc_tanque, y=yc_tanque, text="Reservorio<br>Excavado", showarrow=False)

                # --- 5. DIBUJAR COSECHA DE AGUAS LLUVIAS (Si aplica) ---
                area_tejado_fisica = locals().get('area_tejado_fisica', 0)
                habilitar_cosecha = locals().get('habilitar_cosecha', False)

                if habilitar_cosecha and area_tejado_fisica > 0:
                    largo_tejado = locals().get('largo_tejado', 10)
                    ancho_tejado = locals().get('ancho_tejado', 10)
                    y_tej_base = yc_tanque + distancia_centro + margen
                    
                    fig_esq.add_shape(
                        type="rect", 
                        x0=lado_cultivo + margen, y0=y_tej_base, 
                        x1=lado_cultivo + margen + largo_tejado, y1=y_tej_base + ancho_tejado, 
                        line=dict(color="DimGray", width=2), fillcolor="rgba(169,169,169,0.6)"
                    )
                    fig_esq.add_annotation(
                        x=lado_cultivo + margen + (largo_tejado/2), 
                        y=y_tej_base + (ancho_tejado/2), 
                        text=f"Cubierta<br>({area_tejado_fisica:,.0f} m²)", showarrow=False
                    )

                # --- 6. CONFIGURACIÓN Y RENDERIZADO DEL GRÁFICO ---
                fig_esq.update_layout(
                    xaxis=dict(scaleanchor="y", scaleratio=1, showgrid=False, zeroline=False, visible=False),
                    yaxis=dict(showgrid=False, zeroline=False, visible=False),
                    plot_bgcolor="white", margin=dict(l=0, r=0, t=30, b=0),
                    title_text="Distribución Espacial a Escala Real", title_x=0.5
                )
                
                st.plotly_chart(fig_esq, use_container_width=True)
                st.caption("🔍 *Nota: Este plano geométrico está renderizado a escala real (1:1). Ayuda a dimensionar la magnitud de las obras civiles frente a la extensión agrícola.*")

                # Guardar imagen del esquema usando matplotlib (no requiere kaleido)
                try:
                    _area_cult_ha  = st.session_state.get('area_total_ha', 0.5)
                    _area_cult_m2  = _area_cult_ha * 10000
                    _lado_cult     = math.sqrt(_area_cult_m2)
                    _margen        = 5.0

                    fig_esq_mpl, ax_esq = plt.subplots(figsize=(10, 6))
                    ax_esq.set_aspect('equal')

                    # Cultivo
                    cult_rect = mpatches.FancyBboxPatch(
                        (0, 0), _lado_cult, _lado_cult,
                        boxstyle="round,pad=1", linewidth=1.5,
                        edgecolor='darkgreen', facecolor='#90EE90', alpha=0.5
                    )
                    ax_esq.add_patch(cult_rect)
                    ax_esq.text(_lado_cult/2, _lado_cult/2,
                                f"Área de Cultivo\n{_area_cult_ha * 10000:,.0f} m²",
                                ha='center', va='center', fontsize=9, color='darkgreen', fontweight='bold')

                    # Reservorio
                    if es_excavado:
                        _area_max_esq = func_area(vol_max_sistema) if vol_max_sistema > 0 else 100.0
                        _lado_res     = math.sqrt(_area_max_esq)
                        _xr = _lado_cult + _margen
                        res_rect = mpatches.FancyBboxPatch(
                            (_xr, 0), _lado_res, _lado_res,
                            boxstyle="round,pad=0.5", linewidth=1.5,
                            edgecolor='saddlebrown', facecolor='#40E0D0', alpha=0.75
                        )
                        ax_esq.add_patch(res_rect)
                        ax_esq.text(_xr + _lado_res/2, _lado_res/2,
                                    f"Reservorio\nExcavado\n{vol_max_sistema:.1f} m³",
                                    ha='center', va='center', fontsize=9, color='saddlebrown', fontweight='bold')
                        _x_max = _xr + _lado_res + _margen
                        _y_max = max(_lado_cult, _lado_res) + _margen
                    else:
                        _r_esq = locals().get('radio_tanque', 5.0)
                        _xc    = _lado_cult + _margen + _r_esq
                        _yc    = _r_esq + _margen
                        circ   = plt.Circle((_xc, _yc), _r_esq,
                                             linewidth=1.5, edgecolor='darkblue', facecolor='#87CEEB', alpha=0.75)
                        ax_esq.add_patch(circ)
                        ax_esq.text(_xc, _yc,
                                    f"Tanque\nAustraliano\n{v_max:.1f} m³",
                                    ha='center', va='center', fontsize=9, color='darkblue', fontweight='bold')
                        _x_max = _xc + _r_esq + _margen
                        _y_max = max(_lado_cult, 2*_r_esq) + 2*_margen

                    # Tejado / cosecha (si aplica)
                    _hab_cos  = locals().get('habilitar_cosecha', False)
                    _largo_t  = locals().get('largo_tejado', 0.0)
                    _ancho_t  = locals().get('ancho_tejado', 0.0)
                    _area_tej = locals().get('area_tejado_fisica', 0.0)
                    if _hab_cos and _area_tej > 0:
                        _xt = _lado_cult + _margen
                        _yt = _y_max - _ancho_t - _margen
                        tej_rect = mpatches.FancyBboxPatch(
                            (_xt, _yt), _largo_t, _ancho_t,
                            boxstyle="round,pad=0.3", linewidth=1.5,
                            edgecolor='dimgray', facecolor='#D3D3D3', alpha=0.8
                        )
                        ax_esq.add_patch(tej_rect)
                        ax_esq.text(_xt + _largo_t/2, _yt + _ancho_t/2,
                                    f"Cubierta\n{_area_tej:.0f} m²",
                                    ha='center', va='center', fontsize=8, color='dimgray', fontweight='bold')
                        _x_max = max(_x_max, _xt + _largo_t + _margen)

                    ax_esq.set_xlim(-_margen, _x_max)
                    ax_esq.set_ylim(-_margen, _y_max)
                    ax_esq.set_xlabel("Distancia Este (m)", fontsize=9)
                    ax_esq.set_ylabel("Distancia Norte (m)", fontsize=9)
                    ax_esq.set_title("Esquema Representativo de Áreas (escala 1:1)", fontsize=11, fontweight='bold')
                    ax_esq.grid(True, linestyle='--', alpha=0.3)
                    plt.tight_layout()

                    buf_esq = io.BytesIO()
                    fig_esq_mpl.savefig(buf_esq, format='png', dpi=150, bbox_inches='tight')
                    buf_esq.seek(0)
                    st.session_state['imagen_esquema_bytes'] = buf_esq.getvalue()
                    plt.close(fig_esq_mpl)
                except Exception as _e_esq:
                    st.session_state['imagen_esquema_bytes'] = None
                    st.warning(f"No se pudo exportar el esquema de áreas: {_e_esq}")

                # --- ANÁLISIS DE RESILIENCIA Y AUTO-DIMENSIONAMIENTO ---
                st.divider()
                st.subheader("🛠️ Diagnóstico de Resiliencia y Auto-Dimensionamiento")
                
                col_diag1, col_diag2 = st.columns(2)
                
                with col_diag1:
                    # 1. Resumen seguro (evita el NameError de radio_tanque)
                    if es_excavado:
                        texto_almacenamiento = f"* Tipo: Reservorio Excavado\n* Volumen Máximo: {v_max:.2f} m³"
                    else:
                        radio_seguro = locals().get('radio_tanque', 0.0)
                        texto_almacenamiento = f"* Tipo: Tanque Australiano\n* Radio: {radio_seguro} m\n* Volumen: {v_max:.2f} m³"

                    # Validar si existe cosecha de techos
                    habilitar_cosecha = locals().get('habilitar_cosecha', False)
                    area_tejado_fisica = locals().get('area_tejado_fisica', 0.0)
                    area_tejado_texto = f"{area_tejado_fisica:.2f} m²" if habilitar_cosecha and area_tejado_fisica > 0 else "No implementada"
                    area_cultivo_ha = st.session_state.get('area_total_ha', 0.5)

                    st.info(f"**Diseño Actual Analizado:**\n{texto_almacenamiento}\n* Área Cultivo: {area_cultivo_ha * 10000:,.0f} m²\n* Área Ramada: {area_tejado_texto}")
                
                with col_diag2:
                    if deficit_maximo_registrado > 0:
                        st.error(f"🚨 **ALERTA DE QUIEBRE:** El sistema colapsó en época seca. Faltaron hasta **{deficit_maximo_registrado:.2f} m³** de agua en el peor momento.")
                        
                        # 1. Ajuste Estructural
                        volumen_ideal = v_max + deficit_maximo_registrado
                        
                        # 2. Ajuste Agronómico (Búsqueda Binaria del Área Óptima)
                        low, high = 0.001, area_cultivo_ha
                        area_optima = 0.0
                        
                        for _ in range(20): # 20 iteraciones (precisión de 0.001 Ha)
                            mid = (low + high) / 2
                            factor_area = mid / area_cultivo_ha
                            v_act_sim = v_max
                            fallo_sim = False
                            
                            for idx_s, row_s in df_chrono.iterrows():
                                d_idx = int(row_s['Decada_Año']) - 1
                                p_mm, e_mm = row_s['Precipitacion'], row_s['Evaporacion']
                                
                                # --- Lógica dual: Excavado vs Tanque ---
                                if es_excavado:
                                    area_sim = func_area(v_act_sim) if v_act_sim > 0 else func_area(0)
                                    e_es_s = 0.0 # No se asume cosecha de techos para reservorio excavado
                                else:
                                    area_sim = locals().get('area_tanque', 0.0)
                                    area_tej_ef = locals().get('area_tejado_efectiva', 0.0)
                                    e_es_s = area_tej_ef * (p_mm / 1000.0)
                                
                                e_cp_s = (caudal_concesion * 86400 * dias_d[d_idx]) / 1000.0
                                e_ll_s = area_sim * (p_mm / 1000.0)
                                
                                # Salida de riego ajustada por el factor de iteración
                                if tipo_riego == "Riego por goteo":
                                    s_d_s = (q_diseno_decadal[d_idx] * factor_area * t_max * 3600 * dias_d[d_idx]) / 1000.0
                                else:
                                    s_d_s = (q_diseno_decadal[d_idx] * factor_area * 86400 * dias_d[d_idx])
                                
                                s_e_s = area_sim * (e_mm / 1000.0)
                                s_i_s = s_e_s * 0.10
                                
                                v_temp_s = v_act_sim + e_cp_s + e_ll_s + e_es_s - s_d_s - s_e_s - s_i_s
                                
                                if v_temp_s < 0:
                                    fallo_sim = True
                                    break
                                v_act_sim = v_max if v_temp_s > v_max else v_temp_s
                                
                            if fallo_sim:
                                high = mid # El área sigue siendo muy grande
                            else:
                                low = mid  # El área soporta bien, intentemos un poco más
                                area_optima = mid
                        
                        # Mostrar Soluciones (Diferenciadas por infraestructura)
                        if es_excavado:
                            st.info(f"🏗️ **Opción 1 (Estructural):** Se requiere rediseñar la topografía/batimetría del reservorio para alcanzar un volumen útil de al menos **{volumen_ideal:.2f} m³**.")
                        else:
                            altura_segura = locals().get('altura_tanque', 1.5)
                            radio_ideal = math.sqrt(volumen_ideal / (math.pi * altura_segura))
                            
                            if radio_ideal > 20.0:
                                st.warning(f"🏗️ **Opción 1 (Estructural):** Subir el radio a **{radio_ideal:.2f} m** es inviable (>20m). Te recomendamos aumentar el área de la ramada.")
                            else:
                                st.info(f"🏗️ **Opción 1 (Estructural):** Incrementa el radio del tanque a **{radio_ideal:.2f} metros** para mantener los {area_cultivo_ha * 10000:,.0f} m² actuales de cultivo.")
                            
                        if area_optima > 0.01:
                            st.success(f"🌱 **Opción 2 (Agronómica):** Si no puedes agrandar el reservorio, debes reducir el área de cultivo a máximo **{area_optima * 10000:,.0f} m²** para garantizar agua todo el año.")
                        else:
                            st.error(f"🌱 **Opción 2 (Agronómica):** El reservorio propuesto es tan pequeño que no puede sostener ni 0.01 Ha. ¡Necesitas rediseñar urgentemente o buscar una concesión de agua!")
                    else:
                        st.success("🏆 **DISEÑO ÓPTIMO:** El reservorio propuesto es resiliente y no se vació durante toda la serie climática analizada.")

                # ─────────────────────────────────────────────────────────
                # RECOMENDACIÓN RIPPL: VOLUMEN ÚTIL MÍNIMO NECESARIO
                # ─────────────────────────────────────────────────────────
                st.divider()
                st.subheader("💡 Recomendación por Método de Rippl — Volumen Útil Necesario")

                try:
                    # Búsqueda iterativa Rippl moderno: menor V que satisface no-vaciado en toda la serie
                    v_candidatos = np.arange(10, 5001, 5)  # de 10 a 5000 m³ en pasos de 5
                    v_rippl = None

                    for v_cand in v_candidatos:
                        v_sim = v_cand
                        fallo = False
                        for _, row_r in df_chrono.iterrows():
                            d_idx_r = int(row_r['Decada_Año']) - 1
                            p_mm_r, e_mm_r = row_r['Precipitacion'], row_r['Evaporacion']
                            if es_excavado:
                                area_r = func_area(min(v_sim, vol_max_sistema)) if v_sim > 0 else func_area(0)
                                e_es_r = 0.0
                            else:
                                area_r = math.pi * (radio_tanque ** 2)
                                area_tej_ef_r = area_tejado_efectiva
                                e_es_r = area_tej_ef_r * (p_mm_r / 1000.0)
                            e_cp_r = (caudal_concesion * 86400 * dias_d[d_idx_r]) / 1000.0
                            e_ll_r = area_r * (p_mm_r / 1000.0)
                            if tipo_riego == "Riego por goteo":
                                s_d_r = (q_diseno_decadal[d_idx_r] * t_max * 3600 * dias_d[d_idx_r]) / 1000.0
                            else:
                                s_d_r = (q_diseno_decadal[d_idx_r] * 86400 * dias_d[d_idx_r])
                            s_e_r = area_r * (e_mm_r / 1000.0)
                            s_i_r = s_e_r * 0.10
                            v_temp_r = v_sim + e_cp_r + e_ll_r + e_es_r - s_d_r - s_e_r - s_i_r
                            if v_temp_r < 0:
                                fallo = True
                                break
                            v_sim = min(v_cand, v_temp_r)
                        if not fallo:
                            v_rippl = v_cand
                            break

                    if v_rippl is not None:
                        diferencia_pct = ((v_rippl - v_max) / v_max * 100) if v_max > 0 else 0
                        if v_rippl <= v_max:
                            st.success(
                                f"✅ **Volumen Útil Rippl (V*):** {v_rippl:.0f} m³ — "                                f"El reservorio actual ({v_max:.0f} m³) **supera** el mínimo requerido. "                                f"Margen de seguridad: {abs(diferencia_pct):.1f}%."                            )
                        else:
                            st.warning(
                                f"⚠️ **Volumen Útil Rippl (V*):** {v_rippl:.0f} m³ — "                                f"El reservorio actual ({v_max:.0f} m³) es insuficiente por "                                f"{diferencia_pct:.1f}% para garantizar no-vaciado en toda la serie histórica."                            )
                            if not es_excavado:
                                radio_rippl = math.sqrt(v_rippl / (math.pi * altura_tanque)) if altura_tanque > 0 else 0
                                st.info(f"🔧 Para alcanzar V*={v_rippl:.0f} m³ manteniendo h={altura_tanque:.2f} m, se requiere un radio de **{radio_rippl:.2f} m** (diámetro {radio_rippl*2:.2f} m).")

                        # Guardar V* en sesión para Pestaña 4 (general y por fuente)
                        st.session_state['v_rippl_optimo'] = v_rippl
                        st.session_state['df_simulacion_reservorio_rippl'] = df_simulacion.copy()
                        if usar_nasa_sim:
                            st.session_state['v_rippl_nasa']  = v_rippl
                        else:
                            st.session_state['v_rippl_wapor'] = v_rippl
                        # Actualizar historial con V*
                        if nombre_guardado_key in st.session_state['historial_simulaciones']:
                            st.session_state['historial_simulaciones'][nombre_guardado_key]['v_rippl'] = v_rippl

                        st.caption(
                            f"📐 *Método: Búsqueda iterativa Rippl moderno (criterio no-vaciado sobre toda la serie histórica). "                            f"V* = menor volumen V que satisface V(t) ≥ 0 para todos los t de la serie analizada. "                            f"Referencia: Tabla 7 — Rippl (1883), modernizado por McMahon & Adeloye (2005).*"                        )
                    else:
                        st.error("🚫 No se encontró un volumen Rippl ≤ 5,000 m³ que satisfaga la demanda. Revise la concesión o reduzca el área de cultivo.")
                except Exception as e_rippl:
                    st.warning(f"No se pudo calcular el volumen Rippl: {e_rippl}")

                # ─────────────────────────────────────────────────────────
                # V* MÍNIMO — REFERENCIA CONVENCIONAL (CICLO ÚNICO, ÚLTIMO AÑO)
                # Replica la búsqueda iterativa de Rippl restringida al último
                # año disponible de la serie cargada (NASA o WaPOR), emulando
                # el criterio de dimensionamiento convencional sin continuidad
                # multianual. Por construcción, V* aquí debe ser ≤ que el V*
                # obtenido sobre la serie histórica completa.
                # ─────────────────────────────────────────────────────────
                try:
                    anio_max_serie = int(df_chrono['Año'].max())
                    df_ultimo_anio = df_chrono[df_chrono['Año'] == anio_max_serie]

                    v_candidatos_u = np.arange(10, 5001, 5)
                    v_rippl_ultimo = None
                    for v_cand_u in v_candidatos_u:
                        v_sim_u = v_cand_u
                        fallo_u = False
                        for _, row_u in df_ultimo_anio.iterrows():
                            d_idx_u = int(row_u['Decada_Año']) - 1
                            p_mm_u, e_mm_u = row_u['Precipitacion'], row_u['Evaporacion']
                            if es_excavado:
                                area_u = func_area(min(v_sim_u, vol_max_sistema)) if v_sim_u > 0 else func_area(0)
                                e_es_u = 0.0
                            else:
                                area_u = math.pi * (radio_tanque ** 2)
                                e_es_u = area_tejado_efectiva * (p_mm_u / 1000.0)
                            e_cp_u = (caudal_concesion * 86400 * dias_d[d_idx_u]) / 1000.0
                            e_ll_u = area_u * (p_mm_u / 1000.0)
                            if tipo_riego == "Riego por goteo":
                                s_d_u = (q_diseno_decadal[d_idx_u] * t_max * 3600 * dias_d[d_idx_u]) / 1000.0
                            else:
                                s_d_u = (q_diseno_decadal[d_idx_u] * 86400 * dias_d[d_idx_u])
                            s_e_u = area_u * (e_mm_u / 1000.0)
                            s_i_u = s_e_u * 0.10
                            v_temp_u = v_sim_u + e_cp_u + e_ll_u + e_es_u - s_d_u - s_e_u - s_i_u
                            if v_temp_u < 0:
                                fallo_u = True
                                break
                            v_sim_u = min(v_cand_u, v_temp_u)
                        if not fallo_u:
                            v_rippl_ultimo = v_cand_u
                            break

                    if v_rippl_ultimo is not None:
                        st.info(
                            f"📌 **Referencia convencional (ciclo único, año {anio_max_serie}):** "                            f"V* = {v_rippl_ultimo:.0f} m³ — volumen mínimo necesario para sostener el cultivo "                            f"usando únicamente el último año disponible de la serie, sin continuidad multianual. "                            f"Este valor sustituye la referencia P75% de década crítica en la Tabla 7."                        )
                        if usar_nasa_sim:
                            st.session_state['v_rippl_ultimo_anio_nasa'] = v_rippl_ultimo
                            st.session_state['anio_ultimo_nasa'] = anio_max_serie
                        else:
                            st.session_state['v_rippl_ultimo_anio_wapor'] = v_rippl_ultimo
                            st.session_state['anio_ultimo_wapor'] = anio_max_serie
                    else:
                        st.warning("🚫 No se encontró un V* de referencia convencional (último año) ≤ 5,000 m³.")
                except Exception as e_ultimo:
                    st.warning(f"No se pudo calcular el V* de referencia convencional (último año): {e_ultimo}")

                # ─────────────────────────────────────────────────────────
                # CURVA PARAMÉTRICA V* vs VENTANA DE AÑOS SIMULADOS (NASA POWER)
                # Sección 7.1 — sensibilidad del volumen óptimo frente a la
                # longitud de la ventana histórica analizada (Rippl moderno
                # con ventanas crecientes sobre la serie NASA POWER).
                # ─────────────────────────────────────────────────────────
                if usar_nasa_sim:
                    st.divider()
                    st.subheader("📈 Sensibilidad de V* frente a la Ventana Histórica Simulada (NASA POWER)")
                    st.caption(
                        "Curva paramétrica V* vs número de años simulados (ventanas crecientes: 3, 5, 8, 10, 15 "
                        "y 18 años), construida ejecutando la búsqueda iterativa de Rippl moderno sobre tramos "
                        "consecutivos de la serie NASA POWER. Permite evidenciar empíricamente la sensibilidad "
                        "del V* frente a la longitud del horizonte histórico y reforzar la recomendación "
                        "metodológica sobre el periodo mínimo de análisis."
                    )
                    try:
                        import plotly.graph_objects as go
                        VOLUMEN_COMERCIAL_REF = 290.0  # m³ — reservorio comercial de referencia

                        def buscar_v_rippl_subserie(df_sub):
                            v_candidatos_s = np.arange(10, 5001, 5)
                            for v_cand_s in v_candidatos_s:
                                v_sim_s = v_cand_s
                                fallo_s = False
                                for _, row_s2 in df_sub.iterrows():
                                    d_idx_s = int(row_s2['Decada_Año']) - 1
                                    p_mm_s, e_mm_s = row_s2['Precipitacion'], row_s2['Evaporacion']
                                    if es_excavado:
                                        area_s = func_area(min(v_sim_s, vol_max_sistema)) if v_sim_s > 0 else func_area(0)
                                        e_es_s2 = 0.0
                                    else:
                                        area_s = math.pi * (radio_tanque ** 2)
                                        e_es_s2 = area_tejado_efectiva * (p_mm_s / 1000.0)
                                    e_cp_s2 = (caudal_concesion * 86400 * dias_d[d_idx_s]) / 1000.0
                                    e_ll_s2 = area_s * (p_mm_s / 1000.0)
                                    if tipo_riego == "Riego por goteo":
                                        s_d_s2 = (q_diseno_decadal[d_idx_s] * t_max * 3600 * dias_d[d_idx_s]) / 1000.0
                                    else:
                                        s_d_s2 = (q_diseno_decadal[d_idx_s] * 86400 * dias_d[d_idx_s])
                                    s_e_s2 = area_s * (e_mm_s / 1000.0)
                                    s_i_s2 = s_e_s2 * 0.10
                                    v_temp_s2 = v_sim_s + e_cp_s2 + e_ll_s2 + e_es_s2 - s_d_s2 - s_e_s2 - s_i_s2
                                    if v_temp_s2 < 0:
                                        fallo_s = True
                                        break
                                    v_sim_s = min(v_cand_s, v_temp_s2)
                                if not fallo_s:
                                    return v_cand_s
                            return None

                        anios_disponibles_chrono = sorted(df_chrono['Año'].unique())
                        n_anios_totales = len(anios_disponibles_chrono)
                        ventanas_objetivo = [3, 5, 8, 10, 15, 18]

                        filas_curva = []
                        for vent in ventanas_objetivo:
                            if vent > n_anios_totales:
                                continue
                            anios_ventana = anios_disponibles_chrono[:vent]
                            df_ventana = df_chrono[df_chrono['Año'].isin(anios_ventana)]
                            v_star_vent = buscar_v_rippl_subserie(df_ventana)
                            if v_star_vent is not None:
                                aprovechamiento_pct = (v_star_vent / VOLUMEN_COMERCIAL_REF) * 100
                                filas_curva.append({
                                    "Ventana (años)": vent,
                                    "V* (m³)": v_star_vent,
                                    "Margen operacional V*/290 (%)": round(aprovechamiento_pct, 1)
                                })

                        if len(filas_curva) >= 2:
                            df_curva_vstar = pd.DataFrame(filas_curva)
                            fig_curva_vstar = go.Figure()
                            fig_curva_vstar.add_trace(go.Scatter(
                                x=df_curva_vstar["Ventana (años)"], y=df_curva_vstar["V* (m³)"],
                                mode="lines+markers", name="V* (m³)", line=dict(color="#1f77b4", width=2),
                                marker=dict(size=8)
                            ))
                            fig_curva_vstar.add_hline(
                                y=VOLUMEN_COMERCIAL_REF, line_dash="dash", line_color="#d62728",
                                annotation_text=f"Reservorio comercial de referencia ({VOLUMEN_COMERCIAL_REF:.0f} m³)"
                            )
                            fig_curva_vstar.update_layout(
                                title="V* vs Número de Años Simulados (Rippl moderno, ventanas crecientes — NASA POWER)",
                                xaxis_title="Ventana de simulación (años)",
                                yaxis_title="Volumen Útil Rippl V* (m³)",
                                height=400
                            )
                            st.plotly_chart(fig_curva_vstar, use_container_width=True)

                            st.dataframe(
                                df_curva_vstar.style.format({
                                    "V* (m³)": "{:.0f}", "Margen operacional V*/290 (%)": "{:.1f}%"
                                }),
                                use_container_width=True, hide_index=True
                            )
                            st.caption(
                                "El margen operacional disponible (V*/290 × 100) expresa el porcentaje del "
                                "reservorio comercial de referencia (290 m³) que sería requerido por el V* "
                                "identificado en cada ventana histórica. Valores crecientes con la longitud de "
                                "la ventana evidencian la necesidad de series históricas suficientemente largas "
                                "para no subestimar el volumen de diseño."
                            )
                        else:
                            st.info(
                                "No hay suficientes años disponibles en la serie NASA POWER para construir la "
                                "curva con las ventanas solicitadas (3, 5, 8, 10, 15, 18 años). Amplíe el periodo "
                                "de descarga en la Pestaña 1/2."
                            )
                    except Exception as e_curva_vstar:
                        st.warning(f"No se pudo construir la curva V* vs ventana histórica: {e_curva_vstar}")


# --- PESTAÑA 4: REPORTE COMPARATIVO NASA vs WaPOR + ANÁLISIS ANUAL RIPPL ---

with tab4:
    st.header("📊 Reporte Técnico Comparativo: NASA POWER vs WaPOR v3")
    st.markdown(
        "Este módulo consolida los resultados de ambas fuentes climáticas y genera el reporte estructurado "
        "con las **Tablas 7, 8 y 9** del análisis de tesis, análisis anual por episodios ENSO y "
        "generación del documento Word de memorias de cálculo."
    )

    # ─────────────────────────────────────────────────────────────────────
    # ESTADO DE MEMORIAS
    # ─────────────────────────────────────────────────────────────────────
    col_status1, col_status2, col_status3 = st.columns(3)
    with col_status1:
        nasa_ok = st.session_state.get('fuente_nasa_lista', False)
        if nasa_ok:
            st.success(f"✅ NASA POWER: {st.session_state.get('anios_nasa', '?')} años cargados")
        else:
            st.warning("⚠️ NASA POWER: sin datos — ejecute Pestaña 1 con fuente NASA")
    with col_status2:
        wapor_ok = st.session_state.get('fuente_wapor_lista', False)
        if wapor_ok:
            st.success(f"✅ WaPOR v3: {st.session_state.get('anios_wapor', '?')} años cargados")
        else:
            st.warning("⚠️ WaPOR v3: sin datos — ejecute Pestaña 1 con fuente WaPOR")
    with col_status3:
        sim_ok = st.session_state.get('df_simulacion_reservorio', None) is not None
        v_rippl = st.session_state.get('v_rippl_optimo', None)
        if sim_ok:
            txt_vrippl = f" | V* Rippl = {v_rippl:.0f} m³" if v_rippl else ""
            st.success(f"✅ Simulación Pestaña 3 lista{txt_vrippl}")
        else:
            st.warning("⚠️ Simulación: ejecute Pestaña 3 primero")

    st.divider()

    # ─────────────────────────────────────────────────────────────────────
    # SECCIÓN A: TABLA 7 — VOLUMEN ÓPTIMO SEGÚN FUENTE Y VENTANA
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("📋 Tabla 7 — Volumen Óptimo del Reservorio por Fuente y Ventana Histórica")
    st.caption("Equivalente metodológico a la Tabla 7 del documento de tesis. V* = menor volumen que satisface no-vaciado en toda la ventana analizada.")

    df_nasa_base = st.session_state.get('df_base_nasa', None)
    df_wapor_base = st.session_state.get('df_base_wapor', None)
    v_max_sim = st.session_state.get('volumen_maximo_sistema', 0.0)
    v_rippl_opt = st.session_state.get('v_rippl_optimo', None)
    v_rippl_nasa  = st.session_state.get('v_rippl_nasa', None)
    v_rippl_wapor = st.session_state.get('v_rippl_wapor', None)
    anios_nasa_n = st.session_state.get('anios_nasa', 0)
    anios_wapor_n = st.session_state.get('anios_wapor', 0)

    # Simulaciones separadas por fuente
    df_sim_nasa   = st.session_state.get('df_simulacion_nasa', None)
    df_sim_wapor  = st.session_state.get('df_simulacion_wapor', None)
    df_sim_activa = st.session_state.get('df_simulacion_reservorio', None)  # la última simulada

    # Indicadores de disponibilidad de simulaciones
    col_sim_ind1, col_sim_ind2 = st.columns(2)
    with col_sim_ind1:
        if df_sim_nasa is not None:
            st.success(f"✅ Simulación NASA: {df_sim_nasa['Año'].nunique()} años | V* = {v_rippl_nasa:.0f} m³" if v_rippl_nasa else f"✅ Simulación NASA disponible ({df_sim_nasa['Año'].nunique()} años)")
        else:
            st.warning("⚠️ Simulación NASA: ejecute Pestaña 3 con fuente NASA POWER")
    with col_sim_ind2:
        if df_sim_wapor is not None:
            st.success(f"✅ Simulación WaPOR: {df_sim_wapor['Año'].nunique()} años | V* = {v_rippl_wapor:.0f} m³" if v_rippl_wapor else f"✅ Simulación WaPOR disponible ({df_sim_wapor['Año'].nunique()} años)")
        else:
            st.warning("⚠️ Simulación WaPOR: ejecute Pestaña 3 con fuente WaPOR v3")

    st.divider()

    # Filas ciclo único (referencia convencional) — V* Rippl calculado
    # únicamente sobre el último año disponible de cada fuente, en lugar
    # del volumen físico del sistema. Esto produce un V* de referencia
    # comparable metodológicamente con las simulaciones multianuales, y
    # se espera que sea inferior al V* obtenido con series históricas
    # largas (NASA/WaPOR), evidenciando el riesgo de subdimensionar con
    # el criterio convencional de ciclo único.
    v_rippl_ultimo_nasa = st.session_state.get('v_rippl_ultimo_anio_nasa', None)
    v_rippl_ultimo_wapor = st.session_state.get('v_rippl_ultimo_anio_wapor', None)
    anio_ultimo_nasa = st.session_state.get('anio_ultimo_nasa', None)
    anio_ultimo_wapor = st.session_state.get('anio_ultimo_wapor', None)

    filas_t7 = []
    if v_rippl_ultimo_wapor is not None:
        filas_t7.append({
            "Ventana de análisis": f"Ciclo único — último año WaPOR ({anio_ultimo_wapor})",
            "Fuente": "Referencia convencional",
            "Años": "1",
            "V* estimado (m³)": f"{v_rippl_ultimo_wapor:.0f}",
            "Eventos ENSO cubiertos": "0 (no aplica)",
            "Observación": "Dimensionamiento convencional sin continuidad multianual (Rippl sobre 1 año, fuente WaPOR)"
        })
    if v_rippl_ultimo_nasa is not None:
        filas_t7.append({
            "Ventana de análisis": f"Ciclo único — último año NASA ({anio_ultimo_nasa})",
            "Fuente": "Referencia convencional",
            "Años": "1",
            "V* estimado (m³)": f"{v_rippl_ultimo_nasa:.0f}",
            "Eventos ENSO cubiertos": "0 (no aplica)",
            "Observación": "Dimensionamiento convencional sin continuidad multianual (Rippl sobre 1 año, fuente NASA POWER)"
        })
    if not filas_t7:
        filas_t7.append({
            "Ventana de análisis": "Ciclo único (referencia convencional)",
            "Fuente": "Referencia convencional",
            "Años": "1",
            "V* estimado (m³)": "N/D (ejecute Pestaña 3)",
            "Eventos ENSO cubiertos": "0 (no aplica)",
            "Observación": "Ejecute la simulación en Pestaña 3 con NASA y/o WaPOR para calcular el V* del último año"
        })

    # Fila WaPOR — usa su propio V* Rippl
    if wapor_ok:
        v_wapor_str = (f"{v_rippl_wapor:.0f}" if v_rippl_wapor
                       else ("Ejecute Pestaña 3 con WaPOR" if df_sim_wapor is None
                             else f"{v_max_sim:.0f} (sin Rippl)"))
        filas_t7.append({
            "Ventana de análisis": f"WaPOR v3 ({anios_wapor_n} años)",
            "Fuente": "WaPOR v3",
            "Años": str(anios_wapor_n),
            "V* estimado (m³)": v_wapor_str,
            "Eventos ENSO cubiertos": "2023-2024 (único evento ENSO con cobertura completa)" if anios_wapor_n >= 6 else f"~{max(0, anios_wapor_n-2)} años disponibles",
            "Observación": "Alta resolución espacial (30-250 m). Aunque el producto WaPOR v3 existe desde ~2009, los valores con cobertura confiable y continua para esta zona inician en 2018, por lo que el evento El Niño 2015-2016 no queda representado en la serie; solo el evento ENSO 2023-2024 es analizable."
        })

    # Filas NASA según ventana — usa su propio V* Rippl
    if nasa_ok and df_nasa_base is not None:
        n_filas_antes_nasa = len(filas_t7)
        for ventana, label, enso_txt in [
            (10, "10 años", "2015-2016"),
            (20, "20 años", "2005-2016"),
            (30, "30 años (OMM)", "1997-98, 2002-03, 2015-16, 2023-24"),
        ]:
            if anios_nasa_n >= ventana:
                v_nasa_str = (f"{v_rippl_nasa:.0f}" if v_rippl_nasa
                              else ("Ejecute Pestaña 3 con NASA" if df_sim_nasa is None
                                    else f"{v_max_sim:.0f} (sin Rippl)"))
                filas_t7.append({
                    "Ventana de análisis": f"NASA POWER ({label})",
                    "Fuente": "NASA POWER",
                    "Años": str(ventana),
                    "V* estimado (m³)": v_nasa_str,
                    "Eventos ENSO cubiertos": enso_txt,
                    "Observación": "Serie larga, estándar OMM" if ventana == 30 else "Serie media"
                })
        if len(filas_t7) == n_filas_antes_nasa:  # no se agregó ninguna ventana NASA
            v_nasa_str = (f"{v_rippl_nasa:.0f}" if v_rippl_nasa else "Ejecute Pestaña 3 con NASA")
            filas_t7.append({
                "Ventana de análisis": f"NASA POWER ({anios_nasa_n} años disponibles)",
                "Fuente": "NASA POWER",
                "Años": str(anios_nasa_n),
                "V* estimado (m³)": v_nasa_str,
                "Eventos ENSO cubiertos": "Según serie",
                "Observación": "Amplíe el período de descarga para mayor cobertura ENSO"
            })

    df_tabla7 = pd.DataFrame(filas_t7)
    st.dataframe(df_tabla7, use_container_width=True, hide_index=True)
    st.caption(f"Fuente: Elaboración propia ADR {ANO_ACTUAL}. Referencia metodológica: Rippl (1883), modernizado por McMahon & Adeloye (2005).")

    st.divider()

    # ─────────────────────────────────────────────────────────────────────
    # SECCIÓN A.2 — COMPARACIÓN V* POR TIPO DE RIEGO: GOTEO vs ASPERSIÓN
    # A partir del historial acumulado de simulaciones guardadas (Pestaña 3),
    # NO de un recálculo en vivo con los parámetros compartidos actuales.
    # Esto evita que la tabla sólo refleje la última ejecución de Pestaña 2/3:
    # cada tipo de riego conserva su propia simulación guardada (con su propia
    # eficiencia y, potencialmente, su propio dimensionamiento de reservorio),
    # de manera que el V* de goteo y el V* de aspersión queden claramente
    # diferenciados aunque se hayan calculado en momentos distintos.
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("💧 Comparación del Volumen Útil (V*) — Riego por Goteo vs Riego por Aspersión")
    st.caption(
        "Compara el volumen útil mínimo (V*, método de Rippl moderno) entre la simulación guardada "
        "más reciente para riego por goteo y la simulación guardada más reciente para riego por "
        "aspersión, por cada fuente climática. Cada simulación conserva su propia eficiencia global "
        "y, si fue modificado entre ejecuciones, su propio dimensionamiento de reservorio — a "
        "diferencia de un recálculo en vivo, que solo reflejaría los parámetros de la última corrida "
        "de la Pestaña 2/3."
    )

    historial_t4 = st.session_state.get('historial_simulaciones', {})

    def _ultima_sim_por_tipo(historial, fuente_label, tipo_riego_label):
        """Devuelve (nombre, dict) de la simulación guardada más reciente que coincide
        con la fuente y el tipo de riego indicados, o (None, None) si no existe."""
        candidatas = [
            (nombre, datos) for nombre, datos in historial.items()
            if datos.get('fuente') == fuente_label
            and datos.get('tipo_riego') == tipo_riego_label
            and datos.get('v_rippl') is not None
        ]
        if not candidatas:
            return None, None
        candidatas.sort(key=lambda x: x[1].get('timestamp', ''), reverse=True)
        return candidatas[0]

    if not historial_t4:
        st.info(
            "Aún no hay simulaciones guardadas en el historial. Ejecute la Pestaña 3 (que guarda "
            "automáticamente cada corrida en el historial) primero con riego por goteo y luego con "
            "riego por aspersión — para la misma fuente — para habilitar esta comparación."
        )
    else:
        filas_riego_cmp = []
        for fuente_label_cmp in ["NASA POWER", "WaPOR v3"]:
            nombre_got, datos_got = _ultima_sim_por_tipo(historial_t4, fuente_label_cmp, "Riego por goteo")
            nombre_asp, datos_asp = _ultima_sim_por_tipo(historial_t4, fuente_label_cmp, "Riego por aspersión")
            if datos_got is None and datos_asp is None:
                continue

            v_got = datos_got['v_rippl'] if datos_got else None
            v_asp = datos_asp['v_rippl'] if datos_asp else None
            ef_got = datos_got.get('eficiencia_global') if datos_got else None
            ef_asp = datos_asp.get('eficiencia_global') if datos_asp else None

            filas_riego_cmp.append({
                "Fuente": fuente_label_cmp,
                "V* Riego por Goteo (m³)": f"{v_got:.0f}" if v_got is not None else "Sin simulación guardada",
                "Eficiencia (Goteo)": f"{ef_got*100:.0f}%" if ef_got else "—",
                "Simulación (Goteo)": f"{nombre_got} · {datos_got['timestamp']}" if datos_got else "—",
                "V* Riego por Aspersión (m³)": f"{v_asp:.0f}" if v_asp is not None else "Sin simulación guardada",
                "Eficiencia (Aspersión)": f"{ef_asp*100:.0f}%" if ef_asp else "—",
                "Simulación (Aspersión)": f"{nombre_asp} · {datos_asp['timestamp']}" if datos_asp else "—",
                "Diferencia (m³)": f"{(v_asp - v_got):+.0f}" if (v_got is not None and v_asp is not None) else "—",
            })

        if filas_riego_cmp:
            df_riego_cmp = pd.DataFrame(filas_riego_cmp)
            st.dataframe(df_riego_cmp, use_container_width=True, hide_index=True)
            faltantes = [f["Fuente"] for f in filas_riego_cmp
                         if "Sin simulación" in f["V* Riego por Goteo (m³)"] or "Sin simulación" in f["V* Riego por Aspersión (m³)"]]
            if faltantes:
                st.warning(
                    f"⚠️ Falta al menos un tipo de riego guardado para: {', '.join(faltantes)}. "
                    f"Vuelva a la Pestaña 2, seleccione el tipo de riego faltante (ajustando su "
                    f"eficiencia si corresponde), ejecute el balance y corra la simulación en la "
                    f"Pestaña 3 nuevamente — quedará guardada en el historial sin sobrescribir la anterior, "
                    f"siempre que use un nombre de simulación distinto en el campo '💾 Nombre para guardar "
                    f"esta simulación'."
                )
            st.caption(
                "El riego por aspersión, al distribuir el caudal en un periodo más prolongado (jornada "
                "completa), suele requerir reservorios de menor volumen útil que el riego por goteo bajo "
                "el mismo régimen de horas de riego, aunque ambos parten del mismo requerimiento bruto "
                "decadal (Rb_mm). La diferencia exacta depende también de la eficiencia global elegida "
                "para cada simulación."
            )
        else:
            st.info(
                "No se encontraron simulaciones guardadas para NASA POWER ni WaPOR v3 con un tipo de "
                "riego identificado. Ejecute y guarde al menos una simulación en la Pestaña 3."
            )

    st.divider()

    # ─────────────────────────────────────────────────────────────────────
    # SECCIÓN B: TABLA 8 — EPISODIOS SECOS CRÍTICOS POR AÑO
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("📋 Tabla 8 — Episodios Secos Críticos Identificados (Análisis por Año)")
    st.caption("Análisis de déficit anual y tiempo de recuperación para NASA POWER y WaPOR v3 por separado.")

    df_sim_reservorio = st.session_state.get('df_simulacion_reservorio', None)

    def analizar_episodios_anuales(df_sim, v_max_val, fuente_label):
        """Analiza déficit y recuperación por año a partir de la simulación."""
        if df_sim is None or df_sim.empty:
            return pd.DataFrame()

        filas = []
        anios_sim = sorted(df_sim['Año'].unique())
        v_max_ref = v_max_val if v_max_val > 0 else df_sim['Volumen Final (m3)'].max()

        for anio in anios_sim:
            df_anio = df_sim[df_sim['Año'] == anio].copy()
            vols = df_anio['Volumen Final (m3)'].values
            deficit_decadas = df_anio['Déficit Hídrico (m3)'].values if 'Déficit Hídrico (m3)' in df_anio.columns else np.zeros(len(vols))

            v_min = vols.min()
            v_min_pct = (v_min / v_max_ref * 100) if v_max_ref > 0 else 0
            decadas_deficit = int((deficit_decadas > 0).sum())
            deficit_total = float(deficit_decadas.sum())

            # Tiempo de recuperación: décadas desde el mínimo hasta que vuelve a >80% de V_max
            idx_min = int(np.argmin(vols))
            t_recuperacion = 0
            for k in range(idx_min, len(vols)):
                if vols[k] >= v_max_ref * 0.80:
                    t_recuperacion = k - idx_min
                    break
            else:
                t_recuperacion = len(vols) - idx_min  # no recuperó en el año

            # Detectar si hay un evento ENSO conocido
            enso_conocidos = {1997: "El Niño MF", 1998: "El Niño MF (pico)", 2002: "El Niño M",
                              2003: "El Niño M", 2015: "El Niño MF", 2016: "El Niño MF (pico)",
                              2023: "El Niño F", 2024: "El Niño F (pico)"}
            enso_label = enso_conocidos.get(anio, "—")

            estado_anio = "Déficit crítico ⚠️" if decadas_deficit > 0 else ("Operación normal ✅" if v_min_pct > 20 else "Nivel bajo ⚠️")

            filas.append({
                "Fuente": fuente_label,
                "Año": anio,
                "V mínimo (m³)": round(v_min, 1),
                "V mínimo (% V*)": round(v_min_pct, 1),
                "Décadas con déficit": decadas_deficit,
                "Déficit total (m³)": round(deficit_total, 1),
                "Décadas p/recuperar (desde mín.)": t_recuperacion,
                "Evento ENSO": enso_label,
                "Estado": estado_anio
            })

        return pd.DataFrame(filas)

    col_t8a, col_t8b = st.columns(2)

    with col_t8a:
        st.markdown("#### 🛰️ NASA POWER")
        if nasa_ok and df_sim_nasa is not None:
            df_t8_nasa = analizar_episodios_anuales(df_sim_nasa, v_max_sim, "NASA POWER")
            if not df_t8_nasa.empty:
                # Highlight ENSO years
                def highlight_enso(row):
                    if row['Evento ENSO'] != '—':
                        return ['background-color: #fff3cd'] * len(row)
                    elif row['Décadas con déficit'] > 0:
                        return ['background-color: #f8d7da'] * len(row)
                    return [''] * len(row)
                st.dataframe(
                    df_t8_nasa.drop(columns=['Fuente']).style.apply(highlight_enso, axis=1).format({
                        "V mínimo (m³)": "{:.1f}", "V mínimo (% V*)": "{:.1f}%",
                        "Déficit total (m³)": "{:.1f}"
                    }),
                    use_container_width=True, hide_index=True
                )
                st.session_state['df_t8_nasa'] = df_t8_nasa
            else:
                st.info("Sin datos de simulación NASA.")
        elif nasa_ok and df_sim_nasa is None:
            st.info("✅ Datos NASA disponibles. Ejecute la simulación en Pestaña 3 seleccionando **NASA POWER** como fuente.")
        else:
            st.info("Cargue datos NASA en Pestaña 1 y ejecute simulación en Pestaña 3.")

    with col_t8b:
        st.markdown("#### 🌍 WaPOR v3")
        if wapor_ok and df_sim_wapor is not None:
            df_t8_wapor = analizar_episodios_anuales(df_sim_wapor, v_max_sim, "WaPOR v3")
            if not df_t8_wapor.empty:
                def highlight_enso_w(row):
                    if row['Evento ENSO'] != '—':
                        return ['background-color: #d4edda'] * len(row)
                    elif row['Décadas con déficit'] > 0:
                        return ['background-color: #f8d7da'] * len(row)
                    return [''] * len(row)
                st.dataframe(
                    df_t8_wapor.drop(columns=['Fuente']).style.apply(highlight_enso_w, axis=1).format({
                        "V mínimo (m³)": "{:.1f}", "V mínimo (% V*)": "{:.1f}%",
                        "Déficit total (m³)": "{:.1f}"
                    }),
                    use_container_width=True, hide_index=True
                )
                st.session_state['df_t8_wapor'] = df_t8_wapor
            else:
                st.info("Sin datos de simulación WaPOR.")
        elif wapor_ok and df_sim_wapor is None:
            st.info("✅ Datos WaPOR disponibles. Ejecute la simulación en Pestaña 3 seleccionando **WaPOR v3** como fuente.")
        else:
            st.info("Cargue datos WaPOR en Pestaña 1 y ejecute simulación en Pestaña 3.")

    st.caption(f"🟡 Años con fondo amarillo/verde = evento ENSO identificado. 🔴 Fondo rojo = décadas con déficit hídrico. Fuente: Elaboración propia ADR {ANO_ACTUAL}.")

    st.divider()

    # ─────────────────────────────────────────────────────────────────────
    # SECCIÓN B.2 — MATRIZ AÑO × DÉCADA (HEATMAP) DEL % DE V*
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("🗺️ Matriz Año × Década — % del Volumen Útil (V*) Alcanzado por Década")
    st.caption(
        "Amplía la Tabla 8 a resolución decadal (36 décadas/año). El color representa el "
        "porcentaje del volumen mínimo decadal respecto al volumen útil de referencia (V*), "
        "evidenciando el patrón estacional del estrés hídrico dentro de cada año, especialmente "
        "en años ENSO."
    )

    def construir_heatmap_decadal(df_sim, v_ref, titulo_fuente, v_ref_label):
        """
        Normaliza el volumen mínimo decadal respecto al V* propio de cada fuente.
        Cada panel usa su propia referencia, lo que hace comparables los niveles
        de estrés hídrico de manera independiente entre NASA y WaPOR.
        """
        if df_sim is None or df_sim.empty or not v_ref or v_ref <= 0:
            return None
        piv = df_sim.pivot_table(
            index='Año', columns='Decada', values='Volumen Final (m3)', aggfunc='min'
        )
        piv = piv.reindex(columns=range(1, 37))
        # Clip a 100%: décadas con volumen > V* (lleno) se muestran como 100%
        piv_pct = (piv / v_ref * 100).clip(upper=100).round(1)
        fig_hm = px.imshow(
            piv_pct,
            labels=dict(x="Década del Año (1–36)", y="Año", color="% de V*"),
            color_continuous_scale="RdYlGn", zmin=0, zmax=100, aspect="auto",
            title=f"{titulo_fuente}<br><sup>Referencia: V* = {v_ref_label} m³ (propio de esta fuente)</sup>"
        )
        fig_hm.update_layout(height=max(280, 35 * piv_pct.shape[0] + 120))
        return fig_hm

    # ── Cada fuente usa su propio V* Rippl como referencia de normalización ──
    # Esto evita que un V* mayor (NASA) enmascare el estrés hídrico en WaPOR.
    col_hm1, col_hm2 = st.columns(2)

    with col_hm1:
        if df_sim_nasa is not None:
            if v_rippl_nasa:
                ref_nasa   = v_rippl_nasa
                label_nasa = f"{v_rippl_nasa:.0f} Rippl NASA"
            else:
                ref_nasa   = v_max_sim if v_max_sim > 0 else None
                label_nasa = f"{v_max_sim:.0f} vol. diseño (Rippl NASA no calculado)"
            fig_hm_nasa = construir_heatmap_decadal(df_sim_nasa, ref_nasa, "🛰️ NASA POWER", label_nasa)
            if fig_hm_nasa is not None:
                st.plotly_chart(fig_hm_nasa, use_container_width=True)
                if not v_rippl_nasa:
                    st.warning("⚠️ V* Rippl NASA no disponible — se normalizó con el volumen de diseño. Ejecute la simulación NASA en Pestaña 3.")
            else:
                st.info("Sin simulación NASA disponible para construir la matriz.")
        else:
            st.info("Ejecute la simulación con NASA POWER en Pestaña 3.")

    with col_hm2:
        if df_sim_wapor is not None:
            if v_rippl_wapor:
                ref_wapor   = v_rippl_wapor
                label_wapor = f"{v_rippl_wapor:.0f} Rippl WaPOR"
            else:
                ref_wapor   = v_max_sim if v_max_sim > 0 else None
                label_wapor = f"{v_max_sim:.0f} vol. diseño (Rippl WaPOR no calculado)"
            fig_hm_wapor = construir_heatmap_decadal(df_sim_wapor, ref_wapor, "🌍 WaPOR v3", label_wapor)
            if fig_hm_wapor is not None:
                st.plotly_chart(fig_hm_wapor, use_container_width=True)
                if not v_rippl_wapor:
                    st.warning("⚠️ V* Rippl WaPOR no disponible — se normalizó con el volumen de diseño. Ejecute la simulación WaPOR en Pestaña 3.")
            else:
                st.info("Sin simulación WaPOR disponible para construir la matriz.")
        else:
            st.info("Ejecute la simulación con WaPOR v3 en Pestaña 3.")

    st.caption(
        "Cada panel normaliza respecto al **V* propio de su fuente**: Rippl NASA (panel izquierdo) "
        "y Rippl WaPOR (panel derecho). Esto permite leer el nivel de estrés hídrico de cada fuente "
        "en sus propios términos, sin distorsión por diferencia de magnitud entre los dos V*. "
        "🟥 Rojo = décadas en o cerca del 0% de V* propio (déficit). "
        "🟩 Verde = décadas con volumen próximo al 100% del V* propio. "
        "El tiempo de recuperación desde el mínimo hasta el 80% de V* se reporta en la Tabla 8."
    )

    st.divider()

    # ─────────────────────────────────────────────────────────────────────
    # SECCIÓN C: TABLA 9 — COMPARACIÓN MÉTODOS
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("📋 Tabla 9 — Comparación: Dimensionamiento Ciclo Único vs Simulación Continua")

    if v_rippl_ultimo_wapor is not None and v_rippl_ultimo_nasa is not None:
        v_ciclo_unico_t9 = f"WaPOR: {v_rippl_ultimo_wapor:.0f} / NASA: {v_rippl_ultimo_nasa:.0f}"
    elif v_rippl_ultimo_wapor is not None:
        v_ciclo_unico_t9 = f"{v_rippl_ultimo_wapor:.0f} (WaPOR, último año)"
    elif v_rippl_ultimo_nasa is not None:
        v_ciclo_unico_t9 = f"{v_rippl_ultimo_nasa:.0f} (NASA, último año)"
    else:
        v_ciclo_unico_t9 = "N/D"

    df_t9_filas = [
        {"Dimensión evaluada": "Volumen óptimo identificado (m³)",
         "Ciclo único (últ. año)": v_ciclo_unico_t9,
         "Rippl NASA POWER": f"{v_rippl_nasa:.0f}" if v_rippl_nasa else ("Simule con NASA" if nasa_ok else "Sin datos"),
         "Rippl WaPOR v3": f"{v_rippl_wapor:.0f}" if v_rippl_wapor else ("Simule con WaPOR" if wapor_ok else "Sin datos"),
         "Observación": "El Rippl captura episodios ENSO severos; el ciclo único solo cubre el último año"},
        {"Dimensión evaluada": "Eventos ENSO verificados",
         "Ciclo único (últ. año)": "0 (no aplica)",
         "Rippl NASA POWER": f"Sí — {anios_nasa_n} años" if nasa_ok else "Sin datos",
         "Rippl WaPOR v3": f"Sí — {anios_wapor_n} años" if wapor_ok else "Sin datos",
         "Observación": "El ciclo único no verifica eventos extremos"},
        {"Dimensión evaluada": "Aval de no-vaciado",
         "Ciclo único (últ. año)": "Explícito, pero solo sobre 1 año (Rippl)",
         "Rippl NASA POWER": "Explícito (operacional)" if df_sim_nasa is not None else "Pendiente simulación",
         "Rippl WaPOR v3": "Explícito (operacional)" if df_sim_wapor is not None else "Pendiente simulación",
         "Observación": "Ver Tabla 8 — episodios críticos identificados"},
        {"Dimensión evaluada": "Identificación décadas críticas",
         "Ciclo único (últ. año)": "No disponible",
         "Rippl NASA POWER": "Sí (Tabla 8)" if df_sim_nasa is not None else "Pendiente",
         "Rippl WaPOR v3": "Sí (Tabla 8)" if df_sim_wapor is not None else "Pendiente",
         "Observación": "Permite anticipar manejo agronómico"},
        {"Dimensión evaluada": "Resolución espacial",
         "Ciclo único (últ. año)": "Puntual (estación/pixel)",
         "Rippl NASA POWER": "~50 km (global)",
         "Rippl WaPOR v3": "30–250 m (África/MENA/Colombia)",
         "Observación": "NASA: extensión temporal. WaPOR: resolución espacial."},
        {"Dimensión evaluada": "Extensión temporal disponible",
         "Ciclo único (últ. año)": "1 año (último disponible)",
         "Rippl NASA POWER": f"{anios_nasa_n} años (desde 1981)" if nasa_ok else "Sin datos",
         "Rippl WaPOR v3": f"{anios_wapor_n} años (cobertura útil desde 2018)" if wapor_ok else "Sin datos",
         "Observación": "NASA recomendada para análisis OMM (≥30 años)"},
        {"Dimensión evaluada": "Trazabilidad documental",
         "Ciclo único (últ. año)": "Limitada",
         "Rippl NASA POWER": "Alta (código abierto)",
         "Rippl WaPOR v3": "Alta (código abierto)",
         "Observación": "Reproducible y auditable"},
    ]

    df_tabla9 = pd.DataFrame(df_t9_filas)
    st.dataframe(df_tabla9, use_container_width=True, hide_index=True)
    st.caption(f"Fuente: Elaboración propia ADR {ANO_ACTUAL}. Referencia: Tabla 9 — metodología de la Herramienta Computacional 2026.")

    st.divider()

    # ─────────────────────────────────────────────────────────────────────
    # SECCIÓN D: GRÁFICA COMPARATIVA VOLUMEN POR AÑO (NASA vs WaPOR)
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("📈 Comportamiento Anual del Reservorio — NASA vs WaPOR")

    hay_grafica = (df_sim_nasa is not None) or (df_sim_wapor is not None)
    if hay_grafica:
        df_t8_n = st.session_state.get('df_t8_nasa', pd.DataFrame())
        df_t8_w = st.session_state.get('df_t8_wapor', pd.DataFrame())

        # Recalcular si hay sim disponible pero no df_t8 aún
        if df_sim_nasa is not None and df_t8_n.empty:
            df_t8_n = analizar_episodios_anuales(df_sim_nasa, v_max_sim, "NASA POWER")
            st.session_state['df_t8_nasa'] = df_t8_n
        if df_sim_wapor is not None and df_t8_w.empty:
            df_t8_w = analizar_episodios_anuales(df_sim_wapor, v_max_sim, "WaPOR v3")
            st.session_state['df_t8_wapor'] = df_t8_w

        import plotly.graph_objects as go

        fig_comp = go.Figure()

        if not df_t8_n.empty:
            fig_comp.add_trace(go.Bar(
                x=df_t8_n['Año'], y=df_t8_n['V mínimo (% V*)'],
                name='V mínimo % — NASA', marker_color='#3498db', opacity=0.8
            ))
            fig_comp.add_trace(go.Scatter(
                x=df_t8_n['Año'], y=df_t8_n['Décadas con déficit'],
                name='Décadas déficit — NASA', mode='lines+markers',
                line=dict(color='#e74c3c', width=2), yaxis='y2'
            ))

        if not df_t8_w.empty and df_t8_n.empty or (not df_t8_w.empty and not df_t8_n.equals(df_t8_w)):
            fig_comp.add_trace(go.Bar(
                x=df_t8_w['Año'], y=df_t8_w['V mínimo (% V*)'],
                name='V mínimo % — WaPOR', marker_color='#2ecc71', opacity=0.7
            ))

        # Mark ENSO years
        enso_anios_conocidos = [1997, 1998, 2002, 2003, 2015, 2016, 2023, 2024]
        for ea in enso_anios_conocidos:
            fig_comp.add_vrect(x0=ea-0.4, x1=ea+0.4, fillcolor="orange", opacity=0.15, line_width=0)

        fig_comp.add_hline(y=20, line_dash="dash", line_color="red", annotation_text="Nivel crítico 20%")
        fig_comp.add_hline(y=80, line_dash="dot", line_color="green", annotation_text="Recuperación 80%")

        fig_comp.update_layout(
            title="Volumen mínimo anual (% V*) con eventos ENSO marcados — NASA vs WaPOR",
            xaxis_title="Año",
            yaxis=dict(title="V mínimo (% de V*)", range=[0, 110]),
            yaxis2=dict(title="Décadas con déficit", overlaying='y', side='right', range=[0, 36]),
            barmode='group', height=420, hovermode='x unified',
            legend=dict(orientation='h', yanchor='bottom', y=1.02),
            plot_bgcolor='white'
        )
        st.plotly_chart(fig_comp, use_container_width=True)
        st.caption("🟠 Bandas naranjas = años con eventos ENSO documentados (NOAA). Línea roja = nivel crítico 20%. Línea verde = umbral de recuperación 80%.")
    else:
        st.info("Ejecute la simulación en Pestaña 3 (primero con NASA, luego con WaPOR) para visualizar la comparación.")

    st.divider()

    # ─────────────────────────────────────────────────────────────────────
    # SECCIÓN D.1 — TENDENCIA ANUAL DE PRECIPITACIÓN Y RET (NASA vs WaPOR)
    # Diagnóstico visual para identificar si el RET (evapotranspiración de
    # referencia) muestra una tendencia creciente que explique caídas del
    # volumen no atribuibles a déficit de precipitación (p. ej. 2022).
    # Se usan subgráficas (small multiples) por fuente para evitar saturar
    # una sola gráfica con series de longitud muy distinta (NASA: serie
    # larga; WaPOR: serie corta de alta resolución).
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("🌡️ Tendencia Anual de Precipitación y RET — NASA vs WaPOR")
    st.caption(
        "Cada fuente se agrega a totales **anuales** (en vez de decadales) y se muestra en su propio panel, "
        "con línea de tendencia (regresión lineal) para evidenciar si la evapotranspiración de referencia "
        "(RET) viene en aumento frente a la precipitación."
    )

    df_nasa_anual_tend, df_wapor_anual_tend = None, None
    if df_nasa_base is not None and {'Precipitacion', 'RET', 'Año'}.issubset(df_nasa_base.columns):
        df_nasa_anual_tend = df_nasa_base.groupby('Año')[['Precipitacion', 'RET']].sum().reset_index()
    if df_wapor_base is not None and {'Precipitacion', 'RET', 'Año'}.issubset(df_wapor_base.columns):
        df_wapor_anual_tend = df_wapor_base.groupby('Año')[['Precipitacion', 'RET']].sum().reset_index()

    fuentes_tend = [(d, n) for d, n in
                     [(df_nasa_anual_tend, "NASA POWER"), (df_wapor_anual_tend, "WaPOR v3")]
                     if d is not None and len(d) >= 2]

    if fuentes_tend:
        from plotly.subplots import make_subplots

        n_filas_tend = len(fuentes_tend)
        fig_tend = make_subplots(
            rows=n_filas_tend, cols=1,
            subplot_titles=[f"{nombre} — Precipitación y RET anual" for _, nombre in fuentes_tend],
            specs=[[{"secondary_y": True}] for _ in range(n_filas_tend)],
            vertical_spacing=0.16
        )

        for i_tend, (df_anual_f, nombre_f) in enumerate(fuentes_tend, start=1):
            anios_f  = df_anual_f['Año'].values
            p_f      = df_anual_f['Precipitacion'].values
            ret_f    = df_anual_f['RET'].values

            fig_tend.add_trace(
                go.Bar(x=anios_f, y=p_f, name='Precipitación anual (mm)',
                       marker_color='#0b3d91', opacity=0.75, showlegend=(i_tend == 1)),
                row=i_tend, col=1, secondary_y=False
            )
            fig_tend.add_trace(
                go.Scatter(x=anios_f, y=ret_f, name='RET anual (mm)', mode='lines+markers',
                           line=dict(color='#c0392b', width=2), showlegend=(i_tend == 1)),
                row=i_tend, col=1, secondary_y=True
            )

            if len(anios_f) >= 3:
                coef_p_tend = np.polyfit(anios_f, p_f, 1)
                coef_ret_tend = np.polyfit(anios_f, ret_f, 1)
                fig_tend.add_trace(
                    go.Scatter(x=anios_f, y=np.polyval(coef_p_tend, anios_f), name='Tendencia P',
                               mode='lines', line=dict(color='#0b3d91', width=2, dash='dash'),
                               showlegend=(i_tend == 1)),
                    row=i_tend, col=1, secondary_y=False
                )
                fig_tend.add_trace(
                    go.Scatter(x=anios_f, y=np.polyval(coef_ret_tend, anios_f), name='Tendencia RET',
                               mode='lines', line=dict(color='#c0392b', width=2, dash='dash'),
                               showlegend=(i_tend == 1)),
                    row=i_tend, col=1, secondary_y=True
                )
                st.caption(
                    f"📐 **{nombre_f}** — pendiente Precipitación: {coef_p_tend[0]:+.1f} mm/año · "
                    f"pendiente RET: {coef_ret_tend[0]:+.1f} mm/año"
                )

            fig_tend.update_yaxes(title_text="Precipitación (mm/año)", row=i_tend, col=1, secondary_y=False)
            fig_tend.update_yaxes(title_text="RET (mm/año)", row=i_tend, col=1, secondary_y=True)
            fig_tend.update_xaxes(title_text="Año", row=i_tend, col=1)

        fig_tend.update_layout(
            height=380 * n_filas_tend,
            legend=dict(orientation='h', yanchor='bottom', y=1.08),
            plot_bgcolor='white',
            barmode='group'
        )
        st.plotly_chart(fig_tend, use_container_width=True)
        st.caption(
            "🔵 Barras = precipitación anual acumulada. 🔴 Línea = RET anual acumulada. Líneas punteadas = "
            "tendencia lineal de cada variable (regresión). Una pendiente RET positiva y creciente frente a P "
            "estable o positiva sugiere que la caída de volumen en un año dado responde a mayor demanda "
            "evaporativa, no a déficit de lluvia."
        )
    else:
        st.info("Ejecute la Pestaña 1/2 con NASA y/o WaPOR (mínimo 2 años) para visualizar la tendencia de precipitación y RET.")

    st.divider()

    # ─────────────────────────────────────────────────────────────────────
    # SECCIÓN E: DATOS DEL PROYECTO + GENERACIÓN DE MEMORIAS WORD
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("📄 Generación de Memorias de Cálculo (Anexo 3, 6 y 7)")
    st.markdown("Complete los datos del proyecto y genere los documentos Word técnicos.")

    col_d1, col_d2 = st.columns(2)
    with col_d1:
        nombre_proyecto_inp = st.text_input("Nombre del Proyecto", value="", placeholder="Ej: SRIC Vereda El Paraíso", key="inp_nombre_proy")
        departamento_inp    = st.text_input("Departamento", value="", placeholder="Se determina por coordenadas DIVIPOLA", key="inp_depto")
        municipio_inp       = st.text_input("Municipio",    value="", placeholder="Se determina por coordenadas DIVIPOLA", key="inp_mpio")
    with col_d2:
        nombre_benef_inp    = st.text_input("Nombre del Potencial Beneficiario", value="", placeholder="Nombre completo", key="inp_benef")
        id_predio_inp       = st.text_input("ID del Predio", value="", placeholder="Ej: 15001000100100", key="inp_predio")

    lat_anx = st.session_state.get('latitud',  st.session_state.get('lat_nasa_t2', 0.0))
    lon_anx = st.session_state.get('longitud', st.session_state.get('lon_nasa_t2', 0.0))
    cultivo_anx     = st.session_state.get('cultivo_calc', 'No definido')
    tipo_riego_anx  = st.session_state.get('tipo_riego_calc', 'No definido')

    col_inf1, col_inf2 = st.columns(2)
    with col_inf1:
        st.info(f"📍 **Coordenadas:** Lat {lat_anx:.6f}  |  Lon {lon_anx:.6f}")
    with col_inf2:
        st.info(f"🌱 **Cultivo:** {cultivo_anx}  |  💧 **Riego:** {tipo_riego_anx}")

    st.markdown("---")
    col_anx3, col_anx6, col_anx7 = st.columns(3)

    # ANEXO 3
    with col_anx3:
        st.markdown("### 📘 Anexo 3")
        st.markdown("**Hidrología Básica e Información de Reservorios**")
        st.caption("Climatología, precipitación decadal y efectiva, dimensionamiento del reservorio, área/volumen vs elevación, simulación del tránsito y esquema de áreas.")

        if st.button("⚙️ Generar Anexo 3", type="primary", key="btn_anx3"):
            df_sim_anx3       = st.session_state.get('df_simulacion_reservorio', None)
            vol_max_anx3      = st.session_state.get('volumen_maximo_sistema', 0.0)
            imagen_clima_anx3 = st.session_state.get('imagen_clima_bytes', None)
            imagen_sim_anx3   = st.session_state.get('imagen_simulacion_bytes', None)
            imagen_av_anx3    = st.session_state.get('imagen_area_volumen_bytes', None)
            imagen_esq_anx3   = st.session_state.get('imagen_esquema_bytes', None)
            area_cult_anx3    = st.session_state.get('area_total_ha', 0.0)
            es_excavado_anx3  = st.session_state.get('es_excavado_flag', False)
            radio_anx3        = st.session_state.get('radio_tanque_val', 0.0)
            altura_anx3       = st.session_state.get('altura_tanque_val', 0.0)
            habilita_cos_anx3 = st.session_state.get('habilitar_cosecha_val', False)
            largo_tej_anx3    = st.session_state.get('largo_tejado_val', 0.0)
            ancho_tej_anx3    = st.session_state.get('ancho_tejado_val', 0.0)
            area_tej_anx3     = st.session_state.get('area_tejado_fisica_val', 0.0)
            df_bat_anx3       = st.session_state.get('df_batimetria_val', None)
            tipo_alm_anx3     = st.session_state.get('tipo_almacenamiento_elegido', "Tanque Australiano")

            # Datos para Resumen Predio
            fuente_nasa  = st.session_state.get('fuente_nasa_lista',  False)
            fuente_wapor = st.session_state.get('fuente_wapor_lista', False)
            if fuente_nasa and fuente_wapor:
                fuente_anx3 = "NASA POWER y WaPOR v3"
            elif fuente_wapor:
                fuente_anx3 = "WaPOR v3"
            elif fuente_nasa:
                fuente_anx3 = "NASA POWER"
            else:
                fuente_anx3 = "N/D"

            anios_nasa_anx3  = st.session_state.get('anios_nasa',  0)
            anios_wapor_anx3 = st.session_state.get('anios_wapor', 0)
            num_anios_anx3   = max(anios_nasa_anx3, anios_wapor_anx3)
            num_sect_anx3    = int(st.session_state.get('num_sect', 1))

            if df_sim_anx3 is None:
                st.error("⚠️ Ejecuta primero la simulación en la Pestaña 3 antes de generar este documento.")
            else:
                with st.spinner("Generando Anexo 3..."):
                    doc_anx3 = crear_memoria_hidrologia(
                        datos_clima=st.session_state.get('df_chrono', None),
                        coordenadas=None,
                        df_simulacion=df_sim_anx3,
                        tipo_almacenamiento=tipo_alm_anx3,
                        vol_max=vol_max_anx3,
                        nombre_proyecto=nombre_proyecto_inp,
                        departamento=departamento_inp,
                        municipio=municipio_inp,
                        lat_coord=lat_anx,
                        lon_coord=lon_anx,
                        nombre_beneficiario=nombre_benef_inp,
                        id_predio=id_predio_inp,
                        nombre_cultivo=cultivo_anx,
                        sistema_riego=tipo_riego_anx,
                        radio_tanque=radio_anx3,
                        altura_tanque=altura_anx3,
                        diametro_tanque=radio_anx3 * 2,
                        habilitar_cosecha=habilita_cos_anx3,
                        largo_tejado=largo_tej_anx3,
                        ancho_tejado=ancho_tej_anx3,
                        area_tejado_fisica=area_tej_anx3,
                        es_excavado=es_excavado_anx3,
                        df_batimetria=df_bat_anx3,
                        imagen_clima_bytes=imagen_clima_anx3,
                        imagen_simulacion_bytes=imagen_sim_anx3,
                        imagen_area_volumen_bytes=imagen_av_anx3,
                        imagen_esquema_bytes=imagen_esq_anx3,
                        area_cultivo_ha=area_cult_anx3,
                        fuente_datos=fuente_anx3,
                        num_anios_serie=num_anios_anx3,
                        num_sectores=num_sect_anx3,
                    )
                buffer_anx3 = io.BytesIO()
                doc_anx3.save(buffer_anx3)
                st.success("✅ Anexo 3 generado.")
                st.download_button(
                    label="📥 Descargar Anexo 3 (.docx)",
                    data=buffer_anx3.getvalue(),
                    file_name="Anexo_3_Hidrologia_Basica_ADR.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_anx3"
                )

    # ANEXO 6
    with col_anx6:
        st.markdown("### 📗 Anexo 6")
        st.markdown("**Disponibilidad y Demandas de Agua**")
        st.caption("Balance hídrico, ETc decadal, requerimiento neto y bruto, análisis oferta-demanda.")

        if st.button("⚙️ Generar Anexo 6", type="primary", key="btn_anx6"):
            df_bal_anx6 = st.session_state.get('df_balance_t2', None)
            kc_mid_anx6 = st.session_state.get('kc_m', None)
            ef_anx6     = st.session_state.get('t2_ef_total_calc', None)
            area_anx6   = st.session_state.get('area_total_ha', None)
            etc_max_anx6 = None
            if df_bal_anx6 is not None and 'ETc_mm' in df_bal_anx6.columns:
                etc_max_anx6 = float(df_bal_anx6['ETc_mm'].max())

            if df_bal_anx6 is None:
                st.warning("⚠️ Ejecuta la Pestaña 2 (Balance Hídrico) para poblar el anexo con resultados. Se generará la versión metodológica.")

            with st.spinner("Generando Anexo 6..."):
                doc_anx6 = gax.crear_memoria_demandas(
                    df_balance=df_bal_anx6,
                    cultivo=cultivo_anx,
                    tipo_riego=tipo_riego_anx,
                    kc_mid=kc_mid_anx6,
                    ef_global=ef_anx6,
                    area_ha=area_anx6,
                    etc_max=etc_max_anx6,
                    nombre_proyecto=nombre_proyecto_inp,
                    municipio=municipio_inp,
                    departamento=departamento_inp,
                    id_predios=id_predio_inp,
                )
            buffer_anx6 = io.BytesIO()
            doc_anx6.save(buffer_anx6)
            st.success("✅ Anexo 6 generado.")
            st.download_button(
                label="📥 Descargar Anexo 6 (.docx)",
                data=buffer_anx6.getvalue(),
                file_name="Anexo_6_Disponibilidad_Demandas_ADR.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_anx6"
            )

    # ANEXO 7 (Word)
    with col_anx7:
        st.markdown("### 📙 Anexo 7")
        st.markdown("**Memoria de Cálculo de Riego**")
        st.caption("Diseño agronómico e hidráulico (goteo o aspersión), caudales de diseño y eficiencias.")

        if st.button("⚙️ Generar Anexo 7 (Word)", type="primary", key="btn_anx7"):
            q_dec_anx7 = st.session_state.get('q_diseno_decadal', None)
            q_max_anx7 = None
            if q_dec_anx7 is not None:
                try:
                    q_max_anx7 = float(np.max(q_dec_anx7))
                except Exception:
                    q_max_anx7 = None
            with st.spinner("Generando Anexo 7 (Word)..."):
                doc_anx7 = gax.crear_anexo_7(
                    tipo_riego=tipo_riego_anx,
                    cultivo=cultivo_anx,
                    area_ha=st.session_state.get('area_total_ha', None),
                    q_diseno_lps=q_max_anx7,
                    ef_global=st.session_state.get('t2_ef_total_calc', None),
                    nombre_proyecto=nombre_proyecto_inp,
                    municipio=municipio_inp,
                    departamento=departamento_inp,
                    id_predios=id_predio_inp,
                    num_sectores=int(st.session_state.get('num_sect', 1)),
                )
            buffer_anx7 = io.BytesIO()
            doc_anx7.save(buffer_anx7)
            st.success("✅ Anexo 7 (Word) generado.")
            st.download_button(
                label="📥 Descargar Anexo 7 (.docx)",
                data=buffer_anx7.getvalue(),
                file_name="Anexo_7_Memoria_Riego_ADR.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_anx7"
            )

    # ─────────────────────────────────────────────────────────────────────
    # ANEXO 3a — CONSOLIDADO DE DATOS CRUDOS
    # ─────────────────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📘 Anexo 3a — Consolidación de Datos Climáticos Crudos")
    st.caption(
        "Reúne las series diarias crudas (precipitación, evaporación y RET) cargadas en "
        "la Pestaña 1 y procesadas en la Pestaña 2, con una gráfica decadal suavizada de "
        "las tres variables por fuente. Facilita la revisión trazable de los datos fuente."
    )
    col_3a_l, col_3a_r = st.columns([1, 2])
    with col_3a_l:
        max_filas_3a = st.number_input(
            "Filas de la serie diaria a incluir (por fuente)",
            min_value=50, max_value=5000, value=400, step=50, key="max_filas_3a",
            help="La serie diaria completa puede tener miles de filas; se trunca en el Word "
                 "para mantenerlo legible. El CSV completo sigue disponible en la Pestaña 1."
        )
    with col_3a_r:
        if st.button("⚙️ Generar Anexo 3a", type="primary", key="btn_anx3a"):
            df_base_nasa_3a  = st.session_state.get('df_base_nasa', None)
            df_dec_nasa_3a   = st.session_state.get('df_decadal_nasa', None)
            df_base_wapor_3a = st.session_state.get('df_base_wapor', None)
            df_dec_wapor_3a  = st.session_state.get('df_decadal_wapor', None)

            if df_base_nasa_3a is None and df_base_wapor_3a is None:
                st.error("⚠️ No hay series climáticas cargadas. Ejecuta la Pestaña 1 (NASA POWER o WaPOR v3) primero.")
            else:
                with st.spinner("Generando Anexo 3a..."):
                    doc_3a = gax.crear_anexo_3a(
                        df_base_nasa=df_base_nasa_3a,
                        df_decadal_nasa=df_dec_nasa_3a,
                        df_base_wapor=df_base_wapor_3a,
                        df_decadal_wapor=df_dec_wapor_3a,
                        lat=lat_anx, lon=lon_anx,
                        cultivo=cultivo_anx,
                        nombre_proyecto=nombre_proyecto_inp,
                        municipio=municipio_inp,
                        departamento=departamento_inp,
                        max_filas_diarias=int(max_filas_3a),
                    )
                buffer_3a = io.BytesIO()
                doc_3a.save(buffer_3a)
                st.success("✅ Anexo 3a generado.")
                st.download_button(
                    label="📥 Descargar Anexo 3a (.docx)",
                    data=buffer_3a.getvalue(),
                    file_name="Anexo_3a_Datos_Climaticos_Crudos_ADR.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_anx3a"
                )

    # ─────────────────────────────────────────────────────────────────────
    # ANEXO 7a — HOJA DE CÁLCULO HIDRÁULICO (Excel) con inputs previos
    # ─────────────────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📗 Anexo 7a — Hoja de Cálculo Hidráulico (Excel)")
    st.caption(
        "Genera el libro Excel con el cálculo hidráulico funcional por etapas "
        "(Hazen-Williams en conducción, Darcy-Weisbach en múltiple y lateral críticos). "
        "Complete las entradas de diseño; los valores marcados en amarillo son editables "
        "en el propio Excel tras la descarga."
    )

    _es_goteo_ui = "gote" in str(tipo_riego_anx).lower()
    with st.expander("📝 Entradas de diseño del Anexo 7a", expanded=True):
        c1, c2, c3 = st.columns(3)
        with c1:
            in_area_sector = st.number_input(
                "Área por sector (ha)", min_value=0.001,
                value=float(st.session_state.get('area_total_ha', 0.05) or 0.05),
                step=0.01, format="%.3f", key="7a_area_sector")
            in_nsect = st.number_input("N.º de sectores", min_value=1,
                value=int(st.session_state.get('num_sect', 1)), step=1, key="7a_nsect")
            in_valv = st.number_input("Válvulas funcionando a la vez", min_value=1,
                value=1, step=1, key="7a_valv")
            in_ef = st.number_input(
                "Eficiencia de riego", min_value=0.1, max_value=1.0,
                value=float(st.session_state.get('t2_ef_total_calc', 0.90 if _es_goteo_ui else 0.85) or 0.90),
                step=0.01, format="%.2f", key="7a_ef")
            in_kc = st.number_input("Coeficiente Kc", min_value=0.1, max_value=1.5,
                value=float(st.session_state.get('kc_m', 1.10) or 1.10), step=0.05,
                format="%.2f", key="7a_kc")
        with c2:
            _eto_def = 5.0
            _dfb = st.session_state.get('df_balance_t2', None)
            if _dfb is not None and 'RET' in _dfb.columns:
                try: _eto_def = float(_dfb['RET'].max())
                except Exception: _eto_def = 5.0
            in_eto = st.number_input("ETo máxima (mm/día)", min_value=0.5,
                value=round(_eto_def, 2), step=0.1, format="%.2f", key="7a_eto")
            in_sep_e = st.number_input("Separación emisores (m)", min_value=0.05,
                value=0.5 if _es_goteo_ui else 6.6, step=0.1, format="%.2f", key="7a_sepe")
            in_sep_l = st.number_input("Separación laterales (m)", min_value=0.05,
                value=3.0 if _es_goteo_ui else 6.6, step=0.1, format="%.2f", key="7a_sepl")
            in_q_emisor = st.number_input("Caudal del emisor (l/h)", min_value=0.1,
                value=4.0 if _es_goteo_ui else 170.1, step=0.1, format="%.1f", key="7a_qemi")
            in_p_emisor = st.number_input("Presión trabajo emisor (m.c.a.)", min_value=1.0,
                value=10.0 if _es_goteo_ui else 14.0, step=1.0, format="%.1f", key="7a_pemi")
        with c3:
            in_diam_hum = st.number_input("Diámetro humedecido (m)", min_value=0.1,
                value=0.8 if _es_goteo_ui else 13.2, step=0.1, format="%.1f", key="7a_dhum")
            in_jornada = st.number_input("Jornada de operación (h/día)", min_value=1,
                value=int(st.session_state.get('t_max_val', 8) or 8), step=1, key="7a_jorn")
            in_p_toma = st.number_input("Presión disponible en toma (m.c.a.)", min_value=1.0,
                value=40.0, step=1.0, format="%.1f", key="7a_ptoma")

        st.markdown("**Tramos de conducción** (nodo → nodo, longitud y cotas)")
        cc1, cc2, cc3, cc4, cc5 = st.columns(5)
        with cc1:
            t1_long = st.number_input("T1 Long (m) Caseta→Filtro", value=1.0, step=1.0, key="7a_t1l")
            t2_long = st.number_input("T2 Long (m) Filtro→Válvula", value=40.0, step=1.0, key="7a_t2l")
        with cc2:
            t1_ci = st.number_input("T1 Cota ini", value=2698.8, step=0.1, format="%.1f", key="7a_t1ci")
            t2_ci = st.number_input("T2 Cota ini", value=2698.8, step=0.1, format="%.1f", key="7a_t2ci")
        with cc3:
            t1_cf = st.number_input("T1 Cota fin", value=2698.8, step=0.1, format="%.1f", key="7a_t1cf")
            t2_cf = st.number_input("T2 Cota fin", value=2702.5, step=0.1, format="%.1f", key="7a_t2cf")
        with cc4:
            m_long = st.number_input("Múltiple Long (m)", value=13.2, step=0.1, format="%.1f", key="7a_ml")
            m_cf   = st.number_input("Múltiple Cota fin", value=2704.2, step=0.1, format="%.1f", key="7a_mcf")
        with cc5:
            l_long = st.number_input("Lateral Long (m)", value=6.1, step=0.1, format="%.1f", key="7a_ll")
            l_cf   = st.number_input("Lateral Cota fin", value=2705.0, step=0.1, format="%.1f", key="7a_lcf")

    if st.button("⚙️ Generar Anexo 7a (Excel)", type="primary", key="btn_anx7a"):
        inputs_7a = {
            "departamento": departamento_inp, "municipio": municipio_inp,
            "cultivo": cultivo_anx, "tipo_riego": tipo_riego_anx,
            "area_sector_ha": float(in_area_sector), "num_sectores": int(in_nsect),
            "valvulas_simultaneas": int(in_valv), "eficiencia": float(in_ef),
            "eto_max": float(in_eto), "kc": float(in_kc),
            "sep_emisores": float(in_sep_e), "sep_laterales": float(in_sep_l),
            "caudal_emisor_lph": float(in_q_emisor),
            "presion_trabajo_emisor": float(in_p_emisor),
            "diam_humedecido": float(in_diam_hum), "jornada_h": int(in_jornada),
            "presion_disponible_toma": float(in_p_toma),
            "tramos_conduccion": [
                {"nodo_ini": "Caseta", "nodo_fin": "Filtro", "longitud": float(t1_long),
                 "cota_ini": float(t1_ci), "cota_fin": float(t1_cf)},
                {"nodo_ini": "Filtro", "nodo_fin": "Válvula", "longitud": float(t2_long),
                 "cota_ini": float(t2_ci), "cota_fin": float(t2_cf)},
            ],
            "long_multiple": float(m_long), "cota_ini_mult": float(t2_cf),
            "cota_fin_mult": float(m_cf),
            "long_lateral": float(l_long), "cota_ini_lat": float(m_cf),
            "cota_fin_lat": float(l_cf),
        }
        with st.spinner("Generando Anexo 7a (Excel)..."):
            buf_7a = gax.crear_anexo_7a_excel(inputs_7a)
        st.success("✅ Anexo 7a (Excel) generado. Ábralo y recalcule (F9) si su visor no lo hace automáticamente.")
        st.download_button(
            label="📥 Descargar Anexo 7a (.xlsx)",
            data=buf_7a.getvalue(),
            file_name="Anexo_7a_Calculo_Hidraulico_ADR.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_anx7a"
        )

    # ─────────────────────────────────────────────────────────────────────
    # SECCIÓN F: INSTRUCCIONES CONFIG STREAMLIT para ZIP grandes
    # ─────────────────────────────────────────────────────────────────────
    st.divider()
    with st.expander("⚙️ Configuración recomendada para archivos ZIP grandes (>200 MB)"):
        st.markdown("""
Para habilitar la carga de archivos WaPOR de **hasta 2 GB**, crea o edita el archivo `.streamlit/config.toml`
en la raíz de tu proyecto con el siguiente contenido:

```toml
[server]
maxUploadSize = 2048
```

Esto incrementa el límite de Streamlit de 200 MB (predeterminado) a **2048 MB (2 GB)**.
Reinicia la aplicación tras guardar el archivo.
        """)

    with st.expander("ℹ️ ¿Por qué no aparecen mis datos? — Instrucciones de guardado"):
        st.markdown("""
Para que los reportes se generen correctamente, asegúrese de ejecutar las pestañas en orden:

1. **Pestaña 1** → Cargar NASA POWER (para serie larga 30 años) → Cargar WaPOR v3 (para serie corta 8 años)
2. **Pestaña 2** → Calcular Balance Hídrico y Caudales de Diseño
3. **Pestaña 3** → Simular Tránsito del Reservorio (incluye cálculo Rippl automático)
4. **Pestaña 4** → Los reportes y tablas comparativas estarán disponibles aquí

La memoria de sesión almacena **por separado** los datos de NASA y WaPOR, de modo que puede
cargar ambas fuentes en Pestaña 1 de forma secuencial sin perder los datos de la primera carga.
        """)
