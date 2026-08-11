# -*- coding: utf-8 -*-
"""
generadores_anexos.py
=====================================================================
Generadores de anexos para la Herramienta Computacional ADR (HidroApp).

Provee cuatro funciones que consumen st.session_state y devuelven
documentos listos para descargar:

  · crear_anexo_3a(...)        -> Document  (Word: datos crudos + gráfica 3 var.)
  · crear_memoria_demandas(...)-> Document  (Word: Anexo 6, versión ampliada)
  · crear_anexo_7(...)         -> Document  (Word: Anexo 7, memoria de riego)
  · crear_anexo_7a_excel(...)  -> BytesIO   (Excel: hoja de cálculo hidráulico)

Diseñados para respetar el tipo de riego activo (goteo/aspersión).
=====================================================================
"""
import io
import math
import datetime as _dt

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

ANO_ACTUAL = _dt.datetime.now().year

# ════════════════════════════════════════════════════════════════════
# HELPERS DE FORMATO WORD
# ════════════════════════════════════════════════════════════════════
def _p_just(doc, texto, negrita=False, italica=False, tam=None, centrado=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centrado else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texto)
    run.bold = negrita
    run.italic = italica
    if tam:
        run.font.size = Pt(tam)
    return p

def _caption(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.italic = True
    r.font.size = Pt(9)
    return p

def _fuente(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Fuente: Elaboración propia ADR, {ANO_ACTUAL}.")
    r.italic = True
    r.font.size = Pt(8)
    return p

def _ecuacion(doc, texto, numero=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.italic = True
    r.font.size = Pt(11)
    if numero:
        r2 = p.add_run(f"    (Ecuación {numero})")
        r2.font.size = Pt(9)
    return p

def _tabla_desde_df(doc, df, estilo="Table Grid", decimales=2, max_filas=None,
                    encabezados=None):
    """Inserta un DataFrame como tabla Word. Trunca a max_filas si se indica."""
    cols = list(df.columns) if encabezados is None else encabezados
    n = len(df) if max_filas is None else min(len(df), max_filas)
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = estilo
    for i, c in enumerate(cols):
        cell = t.rows[0].cells[i]
        cell.text = str(c)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for _, fila in df.head(n).iterrows():
        celdas = t.add_row().cells
        for i, c in enumerate(df.columns):
            v = fila[c]
            if isinstance(v, (float, np.floating)):
                celdas[i].text = f"{v:.{decimales}f}"
            else:
                celdas[i].text = str(v)
            for run in celdas[i].paragraphs[0].runs:
                run.font.size = Pt(7.5)
    return t


# ════════════════════════════════════════════════════════════════════
# GRÁFICAS
# ════════════════════════════════════════════════════════════════════
def _grafica_tres_variables(df_decadal, fuente_txt="NASA POWER"):
    """
    Gráfica decadal suavizada de las tres variables climáticas.
    Espera columnas: Decada_Año, Prec_75%, Evaporacion, RET
    Devuelve bytes PNG.
    """
    fig, ax = plt.subplots(figsize=(10, 4.5))
    x = df_decadal["Decada_Año"]
    if "Prec_75%" in df_decadal.columns:
        ax.plot(x, df_decadal["Prec_75%"], color="#1f77b4", marker="o",
                markersize=3, linewidth=1.6, label="Precipitación P75% (mm)")
    elif "Precipitacion" in df_decadal.columns:
        ax.plot(x, df_decadal["Precipitacion"], color="#1f77b4", marker="o",
                markersize=3, linewidth=1.6, label="Precipitación (mm)")
    ax.plot(x, df_decadal["Evaporacion"], color="#ff7f0e", marker="s",
            markersize=3, linewidth=1.6, label="Evaporación (mm)")
    ax.plot(x, df_decadal["RET"], color="#d62728", marker="^",
            markersize=3, linewidth=1.6, label="RET / ET0 (mm)")
    ax.set_xlabel("Década del año (1–36)")
    ax.set_ylabel("Lámina de agua (mm)")
    ax.set_title(f"Serie decadal de variables climáticas — {fuente_txt}")
    ax.grid(True, alpha=0.3)
    ax.legend(loc="upper right", fontsize=8)
    ax.set_xlim(1, 36)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130)
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════════
# ANEXO 3a — CONSOLIDADO DE DATOS CRUDOS
# ════════════════════════════════════════════════════════════════════
def crear_anexo_3a(df_base_nasa=None, df_decadal_nasa=None,
                   df_base_wapor=None, df_decadal_wapor=None,
                   lat=None, lon=None, cultivo="N/D",
                   nombre_proyecto="", municipio="", departamento="",
                   max_filas_diarias=400):
    """
    Anexo 3a — Consolidación de datos crudos de precipitación, evaporación
    y evapotranspiración de referencia (RET), por fuente climática.

    Para cada fuente disponible (NASA POWER / WaPOR v3):
      · Tabla de la serie diaria cruda (truncada por legibilidad).
      · Gráfica decadal suavizada de las tres variables.
    """
    doc = Document()

    titulo = doc.add_heading("Anexo 3a. Consolidación de Datos Climáticos Crudos", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _p_just(doc,
        "El presente anexo consolida los archivos fuente de precipitación, "
        "evaporación y evapotranspiración de referencia (RET) empleados en el "
        "análisis hidrológico del Anexo 3. Su propósito es facilitar la revisión "
        "trazable de la información cargada en la Pestaña 1 y procesada en la "
        "Pestaña 2 de la herramienta computacional, permitiendo a cualquier "
        "evaluador auditar los datos crudos que sustentan el dimensionamiento "
        "del reservorio.")

    # --- Datos del análisis
    doc.add_heading("1. Identificación del Análisis", level=1)
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    filas_id = [
        ("Nombre del Proyecto", nombre_proyecto or "N/D"),
        ("Departamento", departamento or "N/D"),
        ("Municipio", municipio or "N/D"),
        ("Coordenadas (Lat, Lon)",
         f"{lat:.6f}°, {lon:.6f}°" if lat is not None and lon is not None else "N/D"),
        ("Cultivo Analizado", cultivo or "N/D"),
        ("Variables Consolidadas", "Precipitación · Evaporación · RET (ET0 Hargreaves-Samani)"),
    ]
    for campo, valor in filas_id:
        row = tabla.add_row().cells
        row[0].text = campo
        row[1].text = str(valor)
        for run in row[0].paragraphs[0].runs:
            run.bold = True

    doc.add_heading("2. Metodología de Consolidación", level=1)
    _p_just(doc,
        "Las series diarias se obtienen de NASA POWER (variables PRECTOTCORR, "
        "T2M_MAX, T2M_MIN y EVPTRNS) o se procesan a partir de rásteres WaPOR v3 "
        "(productos de precipitación y evapotranspiración). La evapotranspiración "
        "de referencia (RET) se calcula por el método de Hargreaves-Samani (1985):")
    _ecuacion(doc, "ET0 = kRS · Ra · (Tmean + 17.8) · √(Tmax − Tmin)", numero=1)
    _p_just(doc,
        "Posteriormente, las series diarias se agregan a escala decadal (períodos "
        "de 8, 10 u 11 días). La precipitación decadal se procesa con el método de "
        "Blom (Critchley & Siegert, 1996) para estimar el valor confiable al 75% de "
        "probabilidad de excedencia (P75%), mientras que la evaporación y la RET se "
        "promedian por década.")

    fuentes = [
        ("NASA POWER", df_base_nasa, df_decadal_nasa),
        ("WaPOR v3",   df_base_wapor, df_decadal_wapor),
    ]
    seccion = 3
    hay_datos = False
    for nombre_fuente, df_base, df_dec in fuentes:
        if df_base is None or (hasattr(df_base, "empty") and df_base.empty):
            continue
        hay_datos = True
        doc.add_heading(f"{seccion}. Fuente: {nombre_fuente}", level=1)

        n_reg = len(df_base)
        anios = int(df_base["Fecha"].dt.year.nunique()) if "Fecha" in df_base.columns else "N/D"
        f_ini = df_base["Fecha"].min().date() if "Fecha" in df_base.columns else "N/D"
        f_fin = df_base["Fecha"].max().date() if "Fecha" in df_base.columns else "N/D"
        _p_just(doc,
            f"Serie diaria de {nombre_fuente}: {n_reg} registros diarios que abarcan "
            f"{anios} años ({f_ini} a {f_fin}).")

        # --- Gráfica decadal de tres variables
        if df_dec is not None and not df_dec.empty:
            doc.add_heading(f"{seccion}.1 Comportamiento Decadal de las Tres Variables",
                            level=2)
            img = _grafica_tres_variables(df_dec, fuente_txt=nombre_fuente)
            doc.add_picture(io.BytesIO(img), width=Inches(6.3))
            _caption(doc,
                f"Ilustración {seccion}-1. Serie decadal de precipitación P75%, "
                f"evaporación y RET — {nombre_fuente}.")
            _fuente(doc)

        # --- Tabla de la serie diaria cruda (truncada)
        doc.add_heading(f"{seccion}.2 Extracto de la Serie Diaria Cruda", level=2)
        cols_disp = [c for c in ["Fecha", "Precipitacion", "T_Max", "T_Min",
                                 "Evaporacion", "RET"] if c in df_base.columns]
        df_show = df_base[cols_disp].copy()
        if "Fecha" in df_show.columns:
            df_show["Fecha"] = df_show["Fecha"].dt.strftime("%Y-%m-%d")
        _caption(doc,
            f"Tabla {seccion}-1. Serie diaria cruda de {nombre_fuente} "
            f"(se muestran las primeras {min(n_reg, max_filas_diarias)} de {n_reg} filas).")
        _tabla_desde_df(doc, df_show, decimales=2, max_filas=max_filas_diarias)
        _fuente(doc)

        # --- Tabla decadal consolidada
        if df_dec is not None and not df_dec.empty:
            doc.add_heading(f"{seccion}.3 Serie Decadal Consolidada", level=2)
            cols_dec = [c for c in ["Decada_Año", "Prec_75%", "Precipitacion",
                                    "Evaporacion", "RET"] if c in df_dec.columns]
            _caption(doc, f"Tabla {seccion}-2. Valores decadales consolidados — {nombre_fuente}.")
            _tabla_desde_df(doc, df_dec[cols_dec], decimales=2)
            _fuente(doc)

        seccion += 1

    if not hay_datos:
        _p_just(doc,
            "⚠️ No se encontraron series climáticas cargadas. Ejecute la Pestaña 1 "
            "(carga NASA POWER o WaPOR v3) antes de generar este anexo.",
            negrita=True)

    return doc


# ════════════════════════════════════════════════════════════════════
# ANEXO 6 — DISPONIBILIDAD Y DEMANDAS DE AGUA (versión ampliada)
# ════════════════════════════════════════════════════════════════════
def crear_memoria_demandas(df_balance=None, cultivo="N/D", tipo_riego="N/D",
                           kc_mid=None, ef_global=None, area_ha=None,
                           etc_max=None, nombre_proyecto="", municipio="",
                           departamento="", id_predios=""):
    """
    Anexo 6 — Disponibilidad y Demandas de Agua.
    Replica la estructura del documento ADR de referencia, poblado con los
    resultados de la ejecución (df_balance de la Pestaña 2).
    """
    doc = Document()

    enc = doc.add_paragraph()
    enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = enc.add_run("PROYECTOS DE SISTEMAS DE RIEGO INDIVIDUALES O COMUNITARIOS")
    r.bold = True
    tit = doc.add_heading("Anexo 6. Disponibilidad y Demandas de Agua", 0)
    tit.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if nombre_proyecto:
        _p_just(doc, f"PROYECTO: {nombre_proyecto}", negrita=True)
    if municipio or departamento:
        _p_just(doc, f"LOCALIZACIÓN: {municipio}, {departamento}.", negrita=True)
    if id_predios:
        _p_just(doc, f"PREDIOS: {id_predios}")

    _p_just(doc,
        "Esta memoria describe el procedimiento para evaluar la disponibilidad "
        "hídrica en un sistema productivo y las demandas de agua, con el objetivo "
        "de seleccionar un sistema de almacenamiento destinado a satisfacer las "
        "necesidades de un cultivo.", negrita=False)

    # 1. Demanda hídrica
    doc.add_heading("1. Determinación de la Demanda Hídrica del Cultivo", level=1)
    _p_just(doc,
        "El primer paso consiste en cuantificar los requerimientos de agua del "
        "cultivo a partir de la evapotranspiración de referencia (ETo), el "
        "coeficiente de cultivo (Kc) según FAO-56 y la evapotranspiración del "
        "cultivo (ETc):")
    _ecuacion(doc, "ETc = ETo · Kc", numero=1)

    if cultivo and cultivo != "N/D":
        txt_kc = f"Para el cultivo analizado ({cultivo})"
        if kc_mid is not None:
            txt_kc += f", se adopta un Kc medio de {kc_mid:.2f}"
        txt_kc += ". "
        if etc_max is not None:
            txt_kc += (f"El valor máximo de ETc calculado, correspondiente a la "
                       f"década más crítica, es de {etc_max:.2f} mm/día, empleado "
                       f"para asegurar la cobertura de los picos de demanda.")
        _p_just(doc, txt_kc)

    # 2. Balance agroclimatológico
    doc.add_heading("2. Balance Agroclimatológico Preliminar", level=1)
    _p_just(doc,
        "Se compara la oferta natural de agua (precipitación efectiva) con la "
        "demanda bruta del cultivo a lo largo del año, identificando los períodos "
        "de déficit (la demanda supera la lluvia) y de superávit. Este balance "
        "justifica la necesidad de un sistema de almacenamiento.")

    # Gráfica de balance si hay datos
    if df_balance is not None and not df_balance.empty:
        cols_ok = all(c in df_balance.columns for c in ["Decada_Año", "Pe_mm", "Rb_mm"])
        if cols_ok:
            fig, ax = plt.subplots(figsize=(10, 4.2))
            ax.bar(df_balance["Decada_Año"], df_balance["Pe_mm"],
                   color="#2b8cbe", alpha=0.7, label="Precipitación efectiva (mm)")
            ax.plot(df_balance["Decada_Año"], df_balance["Rb_mm"],
                    color="#d62728", marker="o", markersize=3, linewidth=1.8,
                    label="Demanda bruta Rb (mm)")
            ax.set_xlabel("Década del año (1–36)")
            ax.set_ylabel("Lámina (mm)")
            ax.set_title("Balance decadal: demanda bruta vs. precipitación efectiva")
            ax.grid(True, alpha=0.3)
            ax.legend(fontsize=8)
            ax.set_xlim(0.5, 36.5)
            fig.tight_layout()
            b = io.BytesIO(); fig.savefig(b, format="png", dpi=130); plt.close(fig)
            b.seek(0)
            doc.add_picture(io.BytesIO(b.getvalue()), width=Inches(6.3))
            _caption(doc, "Ilustración 1. Comportamiento de la demanda bruta frente "
                          "a la precipitación efectiva.")
            _fuente(doc)

    # 3. Requerimiento neto y bruto
    doc.add_heading("3. Requerimiento Neto y Bruto de Riego", level=1)
    _p_just(doc,
        "Considerando la precipitación efectiva, el requerimiento neto (Rn) y el "
        "requerimiento bruto (Rb) se definen como:")
    _ecuacion(doc, "Rn = ETc − Pe", numero=2)
    _ecuacion(doc, "Rb = Rn / Eficiencia", numero=3)
    if ef_global is not None:
        tipo_txt = ("riego por goteo" if "gote" in str(tipo_riego).lower()
                    else "riego por aspersión")
        _p_just(doc,
            f"Para el {tipo_txt} adoptado, se emplea una eficiencia global de "
            f"{ef_global:.2f} ({ef_global*100:.0f}%), que integra las eficiencias "
            f"de conducción, distribución y aplicación.")

    # Tabla resumen del balance
    if df_balance is not None and not df_balance.empty:
        doc.add_heading("4. Resultados del Balance Hídrico Decadal", level=1)
        cols_bal = [c for c in ["Decada_Año", "RET", "Kc", "ETc_mm", "Pe_mm",
                                "Rn_mm", "Rb_mm"] if c in df_balance.columns]
        _caption(doc, "Tabla 1. Balance hídrico decadal (36 décadas del año).")
        _tabla_desde_df(doc, df_balance[cols_bal], decimales=2)
        _fuente(doc)

    # 5. Diseño del reservorio (metodológico)
    doc.add_heading("5. Diseño y Dimensionamiento del Reservorio", level=1)
    _p_just(doc,
        "Para asegurar el suministro durante los periodos de déficit se diseña un "
        "reservorio de almacenamiento. El modelo geométrico corresponde a un "
        "cilindro (armotanque australiano), de fácil construcción, portabilidad y "
        "mantenimiento. El dimensionamiento y la simulación de tránsito se detallan "
        "en el Anexo 3.")

    return doc


# ════════════════════════════════════════════════════════════════════
# ANEXO 7 — MEMORIA DE CÁLCULO DE RIEGO (Word)
# ════════════════════════════════════════════════════════════════════
def crear_anexo_7(tipo_riego="Riego por goteo", cultivo="N/D",
                  area_ha=None, q_diseno_lps=None, ef_global=None,
                  nombre_proyecto="", municipio="", departamento="",
                  id_predios="", num_sectores=1, params_hidraulicos=None):
    """
    Anexo 7 — Metodología aplicada en el sistema de riego.
    Respeta el tipo de riego activo: emplea Hazen-Williams (conducción) y
    Darcy-Weisbach (múltiple/lateral crítico), replicando la estructura del
    documento de referencia ADR.
    """
    es_goteo = "gote" in str(tipo_riego).lower()
    metodo_emisor = "goteros autocompensados" if es_goteo else "microaspersión"

    doc = Document()
    enc = doc.add_paragraph()
    enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = enc.add_run("PROYECTOS DE SISTEMAS DE RIEGO INDIVIDUALES O COMUNITARIOS")
    r.bold = True

    sub = "GOTEO" if es_goteo else "ASPERSIÓN"
    tit = doc.add_heading(
        f"Anexo 7. Metodología Aplicada en el Sistema de Riego por {sub}", 0)
    tit.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if nombre_proyecto:
        _p_just(doc, f"PROYECTO: {nombre_proyecto}", negrita=True)
    if municipio or departamento:
        _p_just(doc, f"LOCALIZACIÓN: {municipio}, {departamento}.", negrita=True)
    if id_predios:
        _p_just(doc, f"PREDIOS: {id_predios}")

    # 1. Objetivo
    doc.add_heading("1. OBJETIVO", level=1)
    _p_just(doc,
        "El objetivo del diseño hidráulico es satisfacer la demanda de agua del "
        f"cultivo de {cultivo}, asegurando que el caudal y la presión sean "
        "adecuados para el uso previsto de riego.")

    # 2. Descripción
    doc.add_heading("2. BREVE DESCRIPCIÓN", level=1)
    _p_just(doc,
        "Un buen diseño optimiza el uso del recurso hídrico, minimizando pérdidas "
        "por fugas, evaporación y deriva; contempla la eficiencia energética en el "
        "uso de bombas y pérdidas de presión, evitando sobrepresiones, cavitación y "
        "golpes de ariete; considera la topografía para el arreglo hidráulico y el "
        "balance hidrológico frente a la variabilidad climática de la zona. "
        "Adicionalmente, el diseño debe ser accesible, simple, duradero y alineado "
        "con las regulaciones locales, regionales e internacionales.")

    # 3. Metodología
    doc.add_heading("3. METODOLOGÍA APLICADA", level=1)
    _p_just(doc, "Para el cálculo hidráulico se consideran los siguientes pasos:")
    pasos = [
        "Curvas de nivel: contorno con intervalo de 0.25 m interpolado del ráster "
        "base, para reflejar el relieve que afecta el cálculo hidráulico.",
        "Área de riego: define la configuración del arreglo hidráulico en el terreno.",
        "Pendiente: se evalúa con los deltas de curvas de nivel y la distancia de "
        "recorrido del agua, para establecer pérdidas admisibles.",
        "Planificación de distribución de tuberías: los diámetros se ordenan de mayor "
        "a menor entre conducción, múltiple y lateral para uniformidad del caudal.",
        "Longitud de conducción: se determina la longitud equivalente (Pitágoras "
        "entre delta de longitud y pendiente).",
        "Selección de accesorios: se reserva una pérdida de 4 m.c.a. por accesorios.",
        f"Selección del sistema de riego: por {metodo_emisor}, acorde a la necesidad "
        "del cultivo y al concepto agronómico.",
        f"Selección del emisor: se escoge un emisor de {metodo_emisor} que se ajuste "
        "al caudal disponible y a la demanda del sector o módulo de riego.",
        "Capacidad del sistema: define si el diseño aplica a módulo, sector o jornada "
        "de riego para cuantificar el caudal de operación.",
        "Filtros: filtro de discos en la etapa inicial de conducción; se reserva una "
        "pérdida de 10 m.c.a.",
        "Diámetro de la tubería: determina velocidad y presión; se recomienda "
        "velocidad entre 0.6 m/s y 2.5 m/s.",
        "Alturas inicial y final: las cotas se chequean para prever conflictos de "
        "presión y velocidad.",
        "Presión de operación: se admiten pérdidas < 20% en múltiple y < 10% en "
        "lateral de riego.",
        "Suma de pérdidas de presión: filtros + válvulas + conducción + múltiple + "
        "lateral + presión del emisor = presión que debe suplir la bomba.",
        "Bomba: máquina hidráulica para impulsar el agua; se instala al inicio de la "
        "conducción y se dimensiona con la solución de bombeo solar del proyecto.",
    ]
    for i, paso in enumerate(pasos, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(paso)

    # Tabla 1. Consideraciones de diseño
    _caption(doc, "Tabla 1. Consideraciones de diseño.")
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    hdr = ["ETAPA", "PRESIÓN", "CAUDAL", "VELOCIDAD", "LONGITUD"]
    for i, h in enumerate(hdr):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(9)
    filas_consid = [
        ("Conducción", "< 20%", "— —", "1.5 ≤ V ≤ 2.0 m/s", "— —"),
        ("Múltiple de riego", "< 15%", "— —", "0.6 ≤ V ≤ 2.5 m/s", "— —"),
        ("Lateral de riego (crítico)", "< 10%", "< 10%", "0.6 ≤ V ≤ 2.5 m/s", "< 70 m"),
    ]
    for fila in filas_consid:
        c = t.add_row().cells
        for i, v in enumerate(fila):
            c[i].text = v
            for run in c[i].paragraphs[0].runs:
                run.font.size = Pt(8)
    _fuente(doc)

    # 4. Ecuaciones generales
    doc.add_heading("4. ECUACIONES GENERALES DEL CÁLCULO HIDRÁULICO", level=1)

    doc.add_heading("4.1 Etapa de Conducción — Hazen-Williams", level=2)
    _ecuacion(doc, "Q = V · A     →     V = Q / A", numero="01")
    _ecuacion(doc, "J = (10.668 / D^4.87) · (Q / C)^1.852     [C = 150 para PVC]", numero="02")
    _ecuacion(doc, "hf = L · J", numero="03")
    _p_just(doc,
        "La ecuación 01 despeja la velocidad para el primer chequeo de diámetros. "
        "La ecuación 02 calcula la fricción unitaria J. La ecuación 03 obtiene las "
        "pérdidas por fricción en todo el tramo.")

    doc.add_heading("4.2 Etapa de Múltiple de Riego (crítico) — Darcy-Weisbach", level=2)
    _ecuacion(doc, "Q = V · A     →     V = Q / A", numero="01")
    _ecuacion(doc, "f = 1.325 / [ln(e/(3.7·D) + 5.74/R^0.9)]²", numero="02")
    _ecuacion(doc, "J = f · V² / (2 · g · D)     [g = 9.81 m/s²]", numero="03")
    _ecuacion(doc, "hf = J · L (total)", numero="04")

    doc.add_heading("4.3 Etapa Lateral de Riego (crítico) — Darcy-Weisbach", level=2)
    if es_goteo:
        _ecuacion(doc, "Q = V · A ≈ qd · (N.º de goteros)", numero="01")
    else:
        _ecuacion(doc, "Q = V · A ≈ qd · (N.º de aspersores)", numero="01")
    _ecuacion(doc, "f = 1.325 / [ln(e/(3.7·D) + 5.74/Re^0.9)]²", numero="02")
    _ecuacion(doc, "Re = V · A / f", numero="03")
    _ecuacion(doc, "J = f · V² / (2 · g · D)", numero="04")
    _ecuacion(doc, "hf = J · L (total)", numero="05")
    _p_just(doc,
        "En el lateral crítico se busca régimen laminar; el número de Reynolds se "
        "verifica contra el diagrama de Moody para confirmar el factor de fricción.")

    # 5. Parámetros de la ejecución
    doc.add_heading("5. PARÁMETROS DE DISEÑO DE LA EJECUCIÓN ACTUAL", level=1)
    tabla = doc.add_table(rows=0, cols=3)
    tabla.style = "Table Grid"
    filas_param = [
        ("Modalidad de riego", tipo_riego, ""),
        ("Cultivo", cultivo, ""),
        ("Área total", f"{area_ha:.3f}" if area_ha else "N/D", "ha"),
        ("N.º de sectores de riego", str(num_sectores), ""),
        ("Eficiencia global", f"{ef_global:.2f}" if ef_global else "N/D", ""),
        ("Caudal de diseño (máx. decadal)",
         f"{q_diseno_lps:.3f}" if q_diseno_lps else "N/D", "L/s"),
    ]
    if params_hidraulicos:
        for k, v in params_hidraulicos.items():
            filas_param.append((k, str(v), ""))
    hdr = ["ITEM", "VALOR", "UNIDAD"]
    r = tabla.add_row().cells
    for i, h in enumerate(hdr):
        r[i].text = h
        for run in r[i].paragraphs[0].runs:
            run.bold = True
    for campo, valor, unidad in filas_param:
        c = tabla.add_row().cells
        c[0].text = campo; c[1].text = str(valor); c[2].text = unidad

    _p_just(doc,
        "El detalle tramo a tramo del cálculo hidráulico (velocidades, fricción, "
        "pérdidas de energía, cotas y presión) se presenta en el Anexo 7a, "
        "hoja de cálculo hidráulico.", italica=True)

    return doc


# ════════════════════════════════════════════════════════════════════
# ANEXO 7a — HOJA DE CÁLCULO HIDRÁULICO (Excel)
# ════════════════════════════════════════════════════════════════════
def crear_anexo_7a_excel(inputs):
    """
    Genera el Excel del Anexo 7a con el cálculo hidráulico funcional
    (fórmulas nativas de Excel por etapas), replicando la lógica del
    archivo de referencia ADR.

    'inputs' es un dict con las entradas previas (recogidas en la UI):
      departamento, municipio, cultivo, tipo_riego,
      area_sector_ha, num_sectores, valvulas_simultaneas,
      eficiencia, eto_max, sep_emisores, sep_laterales, caudal_emisor_lph,
      kc, presion_trabajo_emisor, caudal_descarga_emisor, diam_humedecido,
      jornada_h, presion_disponible_toma,
      # tramos de conducción (lista de dicts): nodo_ini, nodo_fin, longitud, cota_ini, cota_fin
      tramos_conduccion,
      long_multiple, cota_ini_mult, cota_fin_mult,
      long_lateral, cota_ini_lat, cota_fin_lat
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()

    # Estilos
    F_TIT = Font(bold=True, size=14, color="1F4E79")
    F_SEC = Font(bold=True, size=11, color="FFFFFF")
    F_HDR = Font(bold=True, size=9)
    F_LBL = Font(bold=True, size=9)
    FILL_SEC = PatternFill("solid", fgColor="2E75B6")
    FILL_HDR = PatternFill("solid", fgColor="DDEBF7")
    FILL_IN  = PatternFill("solid", fgColor="FFF2CC")   # celdas de entrada
    FILL_OK  = PatternFill("solid", fgColor="E2EFDA")
    thin = Side(style="thin", color="BFBFBF")
    BORDER = Border(left=thin, right=thin, top=thin, bottom=thin)
    CEN = Alignment(horizontal="center", vertical="center", wrap_text=True)

    es_goteo = "gote" in str(inputs.get("tipo_riego", "")).lower()

    # ───────────────────── HOJA 1: INGRESO DE DATOS ─────────────────────
    ws = wb.active
    ws.title = "INGRESO DE DATOS"
    ws.column_dimensions["A"].width = 34
    ws.column_dimensions["B"].width = 4
    ws.column_dimensions["C"].width = 16
    for col in "DEFG":
        ws.column_dimensions[col].width = 13

    ws["A1"] = "ANEXO 7a — INGRESO DE DATOS DE DISEÑO"
    ws["A1"].font = F_TIT
    ws.merge_cells("A1:G1")

    entradas = [
        ("Departamento", inputs.get("departamento", "")),
        ("Municipio", inputs.get("municipio", "")),
        ("Cultivo", inputs.get("cultivo", "")),
        ("Modalidad de riego", inputs.get("tipo_riego", "")),
        ("Área por sector (ha)", inputs.get("area_sector_ha", 0.05)),
        ("N.º de sectores", inputs.get("num_sectores", 1)),
        ("N.º de válvulas funcionando", inputs.get("valvulas_simultaneas", 1)),
        ("Eficiencia de riego", inputs.get("eficiencia", 0.90 if es_goteo else 0.85)),
        ("Evapotranspiración máxima ETo (mm/día)", inputs.get("eto_max", 5.0)),
        ("Coeficiente de cultivo Kc", inputs.get("kc", 1.10)),
        ("Separación entre emisores (m)", inputs.get("sep_emisores", 0.5 if es_goteo else 6.6)),
        ("Separación entre laterales (m)", inputs.get("sep_laterales", 3.0 if es_goteo else 6.6)),
        ("Caudal del emisor (l/h)", inputs.get("caudal_emisor_lph", 4.0 if es_goteo else 170.1)),
        ("Presión de trabajo del emisor (m.c.a.)", inputs.get("presion_trabajo_emisor", 10 if es_goteo else 14)),
        ("Diámetro humedecido / mojado (m)", inputs.get("diam_humedecido", 0.8 if es_goteo else 13.2)),
        ("Jornada de operación (h/día)", inputs.get("jornada_h", 8)),
        ("Presión disponible en la toma (m.c.a.)", inputs.get("presion_disponible_toma", 40)),
    ]
    fila = 3
    ws[f"A{fila}"] = "1. PARÁMETROS GENERALES"
    ws[f"A{fila}"].font = F_SEC; ws[f"A{fila}"].fill = FILL_SEC
    ws.merge_cells(f"A{fila}:C{fila}")
    fila += 1
    ref = {}   # nombre -> celda para referencias
    for etiqueta, valor in entradas:
        ws[f"A{fila}"] = etiqueta; ws[f"A{fila}"].font = F_LBL
        ws[f"C{fila}"] = valor
        ws[f"C{fila}"].fill = FILL_IN
        ws[f"C{fila}"].border = BORDER
        ref[etiqueta] = f"'INGRESO DE DATOS'!C{fila}"
        fila += 1

    # Área de riego (fórmula)
    ws[f"A{fila}"] = "Área de riego total (ha)"; ws[f"A{fila}"].font = F_LBL
    row_area_sect = 3 + 1 + 4   # fila de "Área por sector"
    row_nsect     = 3 + 1 + 5   # fila de "N.º de sectores"
    ws[f"C{fila}"] = f"=C{row_area_sect}*C{row_nsect}"
    ref["Area de Riego (Ha)"] = f"'INGRESO DE DATOS'!C{fila}"
    fila += 2

    # Tramos de conducción
    ws[f"A{fila}"] = "2. TRAMOS DE CONDUCCIÓN"
    ws[f"A{fila}"].font = F_SEC; ws[f"A{fila}"].fill = FILL_SEC
    ws.merge_cells(f"A{fila}:G{fila}")
    fila += 1
    hdr_tramos = ["Nodo inicial", "Nodo final", "Longitud (m)",
                  "Cota inicial (msnm)", "Cota final (msnm)", "Vel. (m/s)", "Chequeo"]
    for i, h in enumerate(hdr_tramos):
        cell = ws.cell(row=fila, column=1 + i, value=h)
        cell.font = F_HDR; cell.fill = FILL_HDR; cell.border = BORDER; cell.alignment = CEN
    fila += 1
    fila_tramo_ini = fila
    tramos = inputs.get("tramos_conduccion", [
        {"nodo_ini": "Caseta", "nodo_fin": "Filtro", "longitud": 1,
         "cota_ini": 2698.8, "cota_fin": 2698.8},
        {"nodo_ini": "Filtro", "nodo_fin": "Válvula", "longitud": 40,
         "cota_ini": 2698.8, "cota_fin": 2702.5},
    ])
    for k, tr in enumerate(tramos):
        r = fila
        ws.cell(row=r, column=1, value=tr.get("nodo_ini", ""))
        ws.cell(row=r, column=2, value=tr.get("nodo_fin", ""))
        ws.cell(row=r, column=3, value=tr.get("longitud", 0)).fill = FILL_IN
        ws.cell(row=r, column=4, value=tr.get("cota_ini", 0)).fill = FILL_IN
        ws.cell(row=r, column=5, value=tr.get("cota_fin", 0)).fill = FILL_IN
        # Velocidad y chequeo vienen de la hoja HIDRÁULICO
        ws.cell(row=r, column=6, value=f"=HIDRÁULICO!T{26 + k}")
        ws.cell(row=r, column=7, value=f'=IF(F{r}<2,"OK","Aumente Diám")')
        for c in range(1, 8):
            ws.cell(row=r, column=c).border = BORDER
        fila += 1
    fila_tramo_fin = fila - 1
    fila += 1

    # Múltiple y lateral
    ws[f"A{fila}"] = "3. MÚLTIPLE Y LATERAL CRÍTICOS"
    ws[f"A{fila}"].font = F_SEC; ws[f"A{fila}"].fill = FILL_SEC
    ws.merge_cells(f"A{fila}:G{fila}")
    fila += 1
    ws[f"A{fila}"] = "Longitud múltiple (m)"; ws[f"C{fila}"] = inputs.get("long_multiple", 13.2)
    ws[f"C{fila}"].fill = FILL_IN
    row_lmult = fila; fila += 1
    ws[f"A{fila}"] = "Cota inicial múltiple (msnm)"; ws[f"C{fila}"] = inputs.get("cota_ini_mult", 2702.5)
    ws[f"C{fila}"].fill = FILL_IN
    row_cimult = fila; fila += 1
    ws[f"A{fila}"] = "Cota final múltiple (msnm)"; ws[f"C{fila}"] = inputs.get("cota_fin_mult", 2704.2)
    ws[f"C{fila}"].fill = FILL_IN
    row_cfmult = fila; fila += 1
    ws[f"A{fila}"] = "Longitud lateral (m)"; ws[f"C{fila}"] = inputs.get("long_lateral", 6.1)
    ws[f"C{fila}"].fill = FILL_IN
    row_llat = fila; fila += 1
    ws[f"A{fila}"] = "Cota inicial lateral (msnm)"; ws[f"C{fila}"] = inputs.get("cota_ini_lat", 2704.2)
    ws[f"C{fila}"].fill = FILL_IN
    row_cilat = fila; fila += 1
    ws[f"A{fila}"] = "Cota final lateral (msnm)"; ws[f"C{fila}"] = inputs.get("cota_fin_lat", 2705.0)
    ws[f"C{fila}"].fill = FILL_IN
    row_cflat = fila; fila += 1

    # ───────────────────── HOJA 2: DATOS PREDIO ─────────────────────
    wsp = wb.create_sheet("DATOS PREDIO")
    for col, w in {"A": 32, "B": 6, "C": 16, "D": 26, "E": 22, "F": 14,
                   "H": 30, "I": 16, "J": 8}.items():
        wsp.column_dimensions[col].width = w
    wsp["A1"] = f"RIEGO INDIVIDUAL — {inputs.get('municipio','')}"
    wsp["A1"].font = F_TIT; wsp.merge_cells("A1:F1")

    wsp["A3"] = "DISEÑO AGRONÓMICO DEL SISTEMA"; wsp["A3"].font = F_SEC
    wsp["A3"].fill = FILL_SEC; wsp.merge_cells("A3:F3")
    R = lambda k: ref[k]
    filas_pred = [
        ("Modalidad de riego", f"={R('Modalidad de riego')}", ""),
        ("Cultivo", f"={R('Cultivo')}", ""),
        ("Eficiencia de aplicación", f"={R('Eficiencia de riego')}", ""),
        ("Jornada de operación", f"={R('Jornada de operación (h/día)')}", "h/día"),
        ("Coeficiente de cultivo Kc", f"={R('Coeficiente de cultivo Kc')}", ""),
        ("ETo (mm/día)", f"={R('Evapotranspiración máxima ETo (mm/día)')}", "mm/día"),
        ("ETc = Kc·ETo (mm/día)",
         f"={R('Coeficiente de cultivo Kc')}*{R('Evapotranspiración máxima ETo (mm/día)')}", "mm/día"),
        ("Lámina de riego a aplicar (mm/día)",
         f"={R('Coeficiente de cultivo Kc')}*{R('Evapotranspiración máxima ETo (mm/día)')}/{R('Eficiencia de riego')}", "mm/día"),
        ("Área neta (ha)", f"={R('Area de Riego (Ha)')}", "ha"),
        ("Caudal del emisor (l/h)", f"={R('Caudal del emisor (l/h)')}", "l/h"),
        ("Separación emisores (m)", f"={R('Separación entre emisores (m)')}", "m"),
        ("Separación laterales (m)", f"={R('Separación entre laterales (m)')}", "m"),
        ("Pluviometría del emisor (mm/hr)",
         f"={R('Caudal del emisor (l/h)')}/({R('Separación entre emisores (m)')}*{R('Separación entre laterales (m)')})", "mm/hr"),
        ("Presión de trabajo del emisor (m.c.a.)", f"={R('Presión de trabajo del emisor (m.c.a.)')}", "m.c.a."),
        ("Diámetro humedecido (m)", f"={R('Diámetro humedecido / mojado (m)')}", "m"),
    ]
    r = 4
    fila_pred_ref = {}   # etiqueta -> fila real en DATOS PREDIO
    for etiqueta, formula, unidad in filas_pred:
        wsp.cell(row=r, column=1, value=etiqueta).font = F_LBL
        wsp.cell(row=r, column=3, value=formula)
        wsp.cell(row=r, column=4, value=unidad)
        fila_pred_ref[etiqueta] = r
        r += 1

    row_lam  = fila_pred_ref["Lámina de riego a aplicar (mm/día)"]
    row_area = fila_pred_ref["Área neta (ha)"]

    # Caudales
    wsp.cell(row=r+1, column=1, value="CAUDALES DE DISEÑO").font = F_SEC
    wsp.cell(row=r+1, column=1).fill = FILL_SEC
    wsp.merge_cells(start_row=r+1, start_column=1, end_row=r+1, end_column=4)
    r += 2
    # Caudal total (l/s):
    #   volumen diario (l/día) = lámina(mm/día)/1000 [m] · área(ha)·10000 [m²] · 1000 [l/m³]
    #   caudal (l/s) = volumen / (jornada_h · 3600 s)
    wsp.cell(row=r, column=1, value="Caudal total del área a regar (l/s)").font = F_LBL
    wsp.cell(row=r, column=3,
             value=f"=(C{row_lam}/1000)*(C{row_area}*10000)*1000/"
                   f"({R('Jornada de operación (h/día)')}*3600)")
    row_qtot = r; r += 1
    wsp.cell(row=r, column=1, value="Caudal por sector (l/s)").font = F_LBL
    wsp.cell(row=r, column=3, value=f"=C{row_qtot}/{R('N.º de sectores')}")
    row_qsect = r; r += 1
    wsp.cell(row=r, column=1, value="Caudal por válvula (l/s)").font = F_LBL
    wsp.cell(row=r, column=3, value=f"=C{row_qsect}/{R('N.º de válvulas funcionando')}")
    row_qvalv = r
    ref["Q_sector_lps"] = f"'DATOS PREDIO'!C{row_qsect}"

    # ───────────────────── HOJA 3: HIDRÁULICO ─────────────────────
    wsh = wb.create_sheet("HIDRÁULICO")
    for col, w in {"A": 14, "B": 10, "C": 10, "D": 12, "E": 8, "F": 10,
                   "G": 10, "H": 10, "I": 10, "J": 10, "K": 10, "L": 11,
                   "M": 11, "N": 12, "O": 12, "P": 11, "Q": 11, "R": 11,
                   "S": 11, "T": 11}.items():
        wsh.column_dimensions[col].width = w

    wsh["A1"] = f" SISTEMA DE RIEGO POR {'GOTEO' if es_goteo else 'ASPERSIÓN'}"
    wsh["A1"].font = F_TIT
    wsh["A2"] = "CÁLCULO HIDRÁULICO — CONDUCCIÓN (HAZEN-WILLIAMS), MÚLTIPLE Y LATERAL (DARCY-WEISBACH)"
    wsh["A2"].font = F_HDR; wsh.merge_cells("A2:T2")

    # Constantes
    wsh["R5"] = "C Hazen-Williams (PVC)"; wsh["T5"] = 150
    wsh["R6"] = "Rugosidad e (mm)"; wsh["T6"] = 0.0015 * 1.4
    wsh["R7"] = "Viscosidad cinemática (m²/s)"; wsh["T7"] = 1.074e-6
    wsh["R8"] = "Gravedad g (m/s²)"; wsh["T8"] = 9.81
    wsh["R9"] = "Caudal sector (l/s)"; wsh["T9"] = f"={R('Q_sector_lps')}"
    for rr in range(5, 10):
        wsh.cell(row=rr, column=18).font = F_LBL

    # --- Bloque CONDUCCIÓN (Hazen-Williams)
    wsh["A11"] = "1. CONDUCCIÓN — HAZEN-WILLIAMS"
    wsh["A11"].font = F_SEC; wsh["A11"].fill = FILL_SEC; wsh.merge_cells("A11:T11")
    encabezados_cond = {
        "A12": "NUDO INI", "B12": "NUDO FIN", "C12": "Ø NOM (pulg)",
        "D12": "Ø INT (mm)", "G12": "LONG (m)", "J12": "LONG TOTAL (m)",
        "K12": "Q (l/s)", "L12": "J (m/m)", "M12": "hf (m)",
        "N12": "COTA INI", "O12": "COTA FIN", "R12": "PRES. REQ (m)",
        "S12": "PRES. EST (m)", "T12": "VEL (m/s)",
    }
    for celda, txt in encabezados_cond.items():
        wsh[celda] = txt; wsh[celda].font = F_HDR
        wsh[celda].fill = FILL_HDR; wsh[celda].alignment = CEN

    # Presión inicial en toma
    wsh["R25"] = f"={R('Presión disponible en la toma (m.c.a.)')}"
    wsh["S25"] = "=R25"
    wsh["Q25"] = f"='INGRESO DE DATOS'!E{fila_tramo_ini}"  # cota final del primer nodo

    n_tramos = len(tramos)
    for k in range(n_tramos):
        r = 26 + k
        ws_in_row = fila_tramo_ini + k
        wsh.cell(row=r, column=1, value=f"='INGRESO DE DATOS'!A{ws_in_row}")
        wsh.cell(row=r, column=2, value=f"='INGRESO DE DATOS'!B{ws_in_row}")
        wsh.cell(row=r, column=3, value=1)                 # Ø nominal (pulg) — editable
        wsh.cell(row=r, column=3).fill = FILL_IN
        # Ø interno (mm) según diámetro nominal (tabla RDE 21)
        wsh.cell(row=r, column=4,
                 value=f'=IF(C{r}=4,108.72,IF(C{r}=3,84.56,IF(C{r}=2,57.19,'
                       f'IF(C{r}=1.5,44.56,IF(C{r}=1.25,38.80,IF(C{r}=1,30.40,'
                       f'IF(C{r}=0.75,24.07,IF(C{r}=0.5,18.74,30.40))))))))')
        wsh.cell(row=r, column=7, value=f"='INGRESO DE DATOS'!C{ws_in_row}")  # longitud
        wsh.cell(row=r, column=10, value=f"=ROUND(G{r}*1.01,0)")               # long total (+1%)
        wsh.cell(row=r, column=11, value="=$T$9")                              # Q
        # J Hazen-Williams
        wsh.cell(row=r, column=12,
                 value=f"=ROUND(0.0985*(K{r}*15.85)^1.85/(D{r}/25.4)^4.8666,4)")
        # hf = J(m/m) * L_total(m) + pérdidas localizadas (filtro en tramo 2, accesorios después)
        extra = "+10" if k == 1 else ("+4" if k >= 2 else "")   # filtro/accesorios
        wsh.cell(row=r, column=13, value=f"=ROUND(L{r}*J{r}{extra},2)")
        # Cotas
        wsh.cell(row=r, column=14, value=f"='INGRESO DE DATOS'!D{ws_in_row}")  # cota ini
        wsh.cell(row=r, column=15, value=f"='INGRESO DE DATOS'!E{ws_in_row}")  # cota fin
        # Presión requerida acumulada
        if k == 0:
            wsh.cell(row=r, column=18, value=f"=R25-M{r}-(O{r}-N{r})")
        else:
            wsh.cell(row=r, column=18, value=f"=R{r-1}-M{r}-(O{r}-N{r})")
        wsh.cell(row=r, column=19, value=f"=S25+(N{r}-O{r})")
        # Velocidad
        wsh.cell(row=r, column=20,
                 value=f"=ROUND(4*(K{r}/1000)/3.1416/(D{r}/1000)^2,2)")
        for c in range(1, 21):
            wsh.cell(row=r, column=c).border = BORDER

    r_cond_fin = 26 + n_tramos - 1

    # --- Bloque MÚLTIPLE crítico (Darcy-Weisbach)
    base_m = r_cond_fin + 3
    wsh.cell(row=base_m, column=1, value="2. MÚLTIPLE DE RIEGO CRÍTICO — DARCY-WEISBACH").font = F_SEC
    wsh.cell(row=base_m, column=1).fill = FILL_SEC
    wsh.merge_cells(start_row=base_m, start_column=1, end_row=base_m, end_column=20)
    rm = base_m + 1
    hdr_m = {1: "Ø INT (mm)", 3: "Q (l/s)", 5: "LONG (m)", 7: "Re",
             9: "f", 11: "J (m/m)", 13: "hf (m)", 15: "VEL (m/s)"}
    for col, txt in hdr_m.items():
        wsh.cell(row=rm, column=col, value=txt).font = F_HDR
        wsh.cell(row=rm, column=col).fill = FILL_HDR
    rmv = rm + 1
    wsh.cell(row=rmv, column=1, value=44.56).fill = FILL_IN   # Ø int múltiple (editable)
    wsh.cell(row=rmv, column=3, value=f"={R('Q_sector_lps')}")
    wsh.cell(row=rmv, column=5, value=f"='INGRESO DE DATOS'!C{row_lmult}")
    # Velocidad
    wsh.cell(row=rmv, column=15, value=f"=ROUND(4*(C{rmv}/1000)/3.1416/(A{rmv}/1000)^2,3)")
    # Reynolds = V*D/n
    wsh.cell(row=rmv, column=7, value=f"=ROUND(O{rmv}*(A{rmv}/1000)/$T$7,0)")
    # f (Swamee-Jain / forma del anexo)
    wsh.cell(row=rmv, column=9,
             value=f"=ROUND(1.325/(LN($T$6/(3.7*A{rmv})+5.74/G{rmv}^0.9))^2,4)")
    # J = f * V^2 / (2 g D)
    wsh.cell(row=rmv, column=11,
             value=f"=ROUND(I{rmv}*O{rmv}^2/(2*$T$8*(A{rmv}/1000)),4)")
    # hf = J * L
    wsh.cell(row=rmv, column=13, value=f"=ROUND(K{rmv}*E{rmv},3)")
    for c in [1, 3, 5, 7, 9, 11, 13, 15]:
        wsh.cell(row=rmv, column=c).border = BORDER

    # --- Bloque LATERAL crítico (Darcy-Weisbach)
    base_l = rmv + 3
    wsh.cell(row=base_l, column=1, value="3. LATERAL DE RIEGO CRÍTICO — DARCY-WEISBACH").font = F_SEC
    wsh.cell(row=base_l, column=1).fill = FILL_SEC
    wsh.merge_cells(start_row=base_l, start_column=1, end_row=base_l, end_column=20)
    rl = base_l + 1
    for col, txt in hdr_m.items():
        wsh.cell(row=rl, column=col, value=txt).font = F_HDR
        wsh.cell(row=rl, column=col).fill = FILL_HDR
    rlv = rl + 1
    wsh.cell(row=rlv, column=1, value=15.4 if es_goteo else 26.0).fill = FILL_IN  # Ø int lateral
    wsh.cell(row=rlv, column=3, value=f"={R('Q_sector_lps')}/{R('N.º de válvulas funcionando')}")
    wsh.cell(row=rlv, column=5, value=f"='INGRESO DE DATOS'!C{row_llat}")
    wsh.cell(row=rlv, column=15, value=f"=ROUND(4*(C{rlv}/1000)/3.1416/(A{rlv}/1000)^2,3)")
    wsh.cell(row=rlv, column=7, value=f"=ROUND(O{rlv}*(A{rlv}/1000)/$T$7,0)")
    wsh.cell(row=rlv, column=9,
             value=f"=ROUND(1.325/(LN($T$6/(3.7*A{rlv})+5.74/G{rlv}^0.9))^2,4)")
    wsh.cell(row=rlv, column=11,
             value=f"=ROUND(I{rlv}*O{rlv}^2/(2*$T$8*(A{rlv}/1000)),4)")
    wsh.cell(row=rlv, column=13, value=f"=ROUND(K{rlv}*E{rlv},3)")
    for c in [1, 3, 5, 7, 9, 11, 13, 15]:
        wsh.cell(row=rlv, column=c).border = BORDER

    # --- RESUMEN DE PRESIÓN DE LA BOMBA
    base_r = rlv + 3
    wsh.cell(row=base_r, column=1, value="4. PRESIÓN REQUERIDA POR LA BOMBA").font = F_SEC
    wsh.cell(row=base_r, column=1).fill = FILL_SEC
    wsh.merge_cells(start_row=base_r, start_column=1, end_row=base_r, end_column=6)
    resumen = [
        ("Pérdidas en toma y filtros (m.c.a.)", "10"),
        ("Pérdidas en válvulas y accesorios (m.c.a.)", "4"),
        ("Suma pérdidas conducción hf (m)",
         f"=SUM(M26:M{r_cond_fin})"),
        ("Pérdida múltiple crítico hf (m)", f"=M{rmv}"),
        ("Pérdida lateral crítico hf (m)", f"=M{rlv}"),
        ("Presión de operación del emisor (m.c.a.)",
         f"={R('Presión de trabajo del emisor (m.c.a.)')}"),
    ]
    rr = base_r + 1
    primera = rr
    for etiqueta, val in resumen:
        wsh.cell(row=rr, column=1, value=etiqueta).font = F_LBL
        wsh.cell(row=rr, column=4, value=val)
        rr += 1
    wsh.cell(row=rr, column=1, value="PRESIÓN TOTAL REQUERIDA (m.c.a.)").font = Font(bold=True, size=10)
    wsh.cell(row=rr, column=4, value=f"=SUM(D{primera}:D{rr-1})")
    wsh.cell(row=rr, column=4).fill = FILL_OK
    wsh.cell(row=rr, column=4).font = Font(bold=True)

    # ───────────────────── HOJA 4: RESUMEN ─────────────────────
    wsr = wb.create_sheet("RESUMEN")
    for col, w in {"A": 40, "B": 18, "C": 8}.items():
        wsr.column_dimensions[col].width = w
    wsr["A1"] = "ANEXO 7a — RESUMEN DEL DISEÑO"; wsr["A1"].font = F_TIT
    wsr.merge_cells("A1:C1")
    resumen_final = [
        ("Cultivo", f"={R('Cultivo')}", ""),
        ("Modalidad de riego", f"={R('Modalidad de riego')}", ""),
        ("Área neta de riego", f"={R('Area de Riego (Ha)')}", "ha"),
        ("N.º de sectores", f"={R('N.º de sectores')}", ""),
        ("Caudal por sector", f"={R('Q_sector_lps')}", "l/s"),
        ("Jornada de operación", f"={R('Jornada de operación (h/día)')}", "h/día"),
        ("Presión disponible en toma", f"={R('Presión disponible en la toma (m.c.a.)')}", "m.c.a."),
        ("Presión total requerida", f"=HIDRÁULICO!D{rr}", "m.c.a."),
        ("Velocidad conducción (últ. tramo)", f"=HIDRÁULICO!T{r_cond_fin}", "m/s"),
        ("Velocidad múltiple crítico", f"=HIDRÁULICO!O{rmv}", "m/s"),
        ("Velocidad lateral crítico", f"=HIDRÁULICO!O{rlv}", "m/s"),
    ]
    r = 3
    for etiqueta, formula, unidad in resumen_final:
        wsr.cell(row=r, column=1, value=etiqueta).font = F_LBL
        wsr.cell(row=r, column=2, value=formula)
        wsr.cell(row=r, column=3, value=unidad)
        for c in range(1, 4):
            wsr.cell(row=r, column=c).border = BORDER
        r += 1

    # Guardar en BytesIO
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf
